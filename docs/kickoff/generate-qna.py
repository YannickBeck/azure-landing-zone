# -*- coding: utf-8 -*-
"""Generiert ein Q&A-Vorbereitungsdokument (erwartete Fragen + Antworten) fuer den ALZ-Kickoff."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/home/user/azure-landing-zone/docs/kickoff"
DATE = "16.06.2026"

BLUE = RGBColor(0x00, 0x78, 0xD4)
DARK = RGBColor(0x24, 0x3A, 0x5E)
GREY = RGBColor(0x60, 0x5E, 0x5C)
RED  = RGBColor(0xC0, 0x39, 0x2B)
GREEN= RGBColor(0x2E, 0x7D, 0x32)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9.5, white=False):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color: r.font.color.rgb = color


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, bold=True, white=True); shade_cell(t.rows[0].cells[i], '0078D4')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def body(doc, text, color=None, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.italic = italic
    if color: r.font.color.rgb = color
    return p


def qa(doc, frage, antworten, heikel=False):
    """Eine Frage (fett, blau) + Antwort-Bullets. heikel=True markiert rot."""
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    if heikel:
        m = p.add_run("[heikel]  "); m.bold = True; m.font.color.rgb = RED; m.font.size = Pt(9)
    r = p.add_run("F:  " + frage); r.bold = True; r.font.color.rgb = BLUE; r.font.size = Pt(11)
    for a in antworten:
        bp = doc.add_paragraph(style='List Bullet'); bp.paragraph_format.space_after = Pt(2)
        run = bp.add_run(a); run.font.size = Pt(10.5); run.font.color.rgb = DARK


def section(doc, title):
    doc.add_heading(title, level=1)


def build():
    doc = Document()
    n = doc.styles['Normal']; n.font.name = 'Calibri'; n.font.size = Pt(10.5)
    for lvl, col, sz in [('Heading 1', DARK, 15), ('Heading 2', BLUE, 12.5)]:
        st = doc.styles[lvl]; st.font.color.rgb = col; st.font.name = 'Calibri'; st.font.size = Pt(sz)

    # Titel
    for _ in range(3): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Azure Landing Zone"); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = BLUE
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Q&A-Vorbereitung – Erwartete Fragen & Antworten"); r.font.size = Pt(15); r.font.color.rgb = DARK
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run("Kickoff mit Technikern & Netzwerkern des Kunden"); r.font.size = Pt(11); r.font.color.rgb = GREY
    for _ in range(6): doc.add_paragraph()
    m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = m.add_run(f"Stand: {DATE} · internes Vorbereitungsdokument"); r.font.size = Pt(10); r.font.color.rgb = GREY
    doc.add_page_break()

    # Nutzung
    section(doc, "So nutzt du dieses Dokument")
    body(doc, "Dieses Deck sammelt Fragen, die dir die technischen Ansprechpartner und Netzwerker des Kunden "
              "voraussichtlich stellen – jeweils mit einer kompakten, fachlich belastbaren Antwort. Mit "
              "[heikel] markierte Fragen bergen ein Risiko (Kosten, Versprechen, Fallstrick): hier ehrlich "
              "bleiben und ggf. „klären wir im Detail-Design“ sagen, statt dich festzulegen.")
    body(doc, "Goldene Regel: Was umgesetzt ist, selbstbewusst zeigen. Was Roadmap ist, klar als Roadmap "
              "benennen. Bei IP-/DNS-/Identity-Themen erst ihren Bestand erfragen, dann zusagen.", color=GREY, italic=True)

    # A. Strategie
    section(doc, "A. Strategie & Architektur")
    qa(doc, "Nutzt ihr den offiziellen Microsoft-ALZ-Accelerator?",
       ["Ja. Wir setzen das offizielle ALZ-Bicep-Accelerator-Starter-Modul ein (Spiegel von Azure/alz-bicep-"
        "accelerator) und rollen es über das ALZ-PowerShell-Modul mit Deploy-Accelerator aus.",
        "Governance kommt aus dem offiziellen Modul avm/ptn/alz/empty – also das volle, von Microsoft "
        "gepflegte ALZ-Policy-Set, nicht ein Eigenbau."])
    qa(doc, "Warum Hub-and-Spoke und nicht Virtual WAN?",
       ["Hub-and-Spoke ist transparenter, günstiger im Einstieg und genügt für eine bis wenige Regionen.",
        "Virtual WAN ist als Alternative im Accelerator vorbereitet (network_type=vwanConnectivity) – sinnvoll "
        "bei vielen Regionen/Standorten oder globalem Transit. Wechsel ist eine bewusste Entscheidung."], heikel=True)
    qa(doc, "Warum Bicep und nicht Terraform?",
       ["Die Plattform-Templates sind Bicep (Azure-nativ, ohne State-Management, mit What-If und AVM direkt aus der Registry).",
        "Der Bootstrap des Accelerators ist Terraform-basiert (State-Storage + OIDC + Repo) – das ist Microsoft-Standard und für uns transparent."])
    qa(doc, "Was ist heute schon umgesetzt, was nicht?",
       ["Umgesetzt: 12 Management Groups, volles ALZ-Policy-Set (149 Definitionen, 42 Initiativen, 123 "
        "Assignments, 5 Custom-Rollen), Logging (avm/ptn/alz/ama), Defender-for-Cloud-Onboarding via Policy, "
        "Hub-Networking, Bootstrap + generierte CI/CD-Pipelines mit OIDC.",
        "Roadmap: dedizierte Identity-Ressourcen, Sentinel-Connectors, ggf. WAF-Ingress."])
    qa(doc, "Sind wir nach dem Aufbau an euch gebunden?",
       ["Nein. Es ist das offizielle Microsoft-Modul, vollständig als Code im Git-Repo, dokumentiert und "
        "reproduzierbar. Der Kunde kann es selbst betreiben oder weiterentwickeln."])

    # B. Governance
    section(doc, "B. Governance, Policy & RBAC")
    qa(doc, "Wie viele Policies greifen, und wie tief geht das Set?",
       ["Das volle ALZ-Policy-Set: 149 Custom-Policy-Definitionen, 42 Initiativen und 123 konkrete Assignments "
        "über die MG-Ebenen verteilt (zzgl. Vererbung) – kein schlankes Eigenbau-Set.",
        "Verteilung der Assignments: alz=17, landingzones=53, platform=40, corp=5, identity=4, je 1 auf "
        "connectivity/local/sandbox/decommissioned."])
    qa(doc, "Was greift konkret ab Tag 1?",
       ["Auf alz (Root): Defender-Onboarding (Deploy-MDFC-*), Activity-Log/Diagnose ins LAW, Compute Security "
        "Baseline (Enforce-ACSB), Service Health.",
        "Auf landingzones-corp: keine Public Endpoints/Public-IP an NICs, kein Hybrid-Networking, Private DNS. "
        "Sandbox: gelockerte Guardrails. Decommissioned: gesperrt."])
    qa(doc, "Was, wenn eine Policy ein legitimes Deployment blockiert?",
       ["Gezielte Policy Exemptions auf MG-/Subscription-/RG-Ebene – versioniert im Code, nicht als Klick im "
        "Portal. Zusätzlich sind viele Guardrail-Initiativen im Auslieferungszustand auf DoNotEnforce (auditierend) "
        "und werden erst nach Bedarf scharfgeschaltet."])
    qa(doc, "Müsst ihr das Policy-Set erst noch erweitern?",
       ["Nein – es ist bereits das vollständige, von Microsoft gepflegte ALZ-Set. Compliance-Feintuning heißt hier "
        "eher: DoNotEnforce-Initiativen scharfschalten bzw. Parameter setzen, nicht Policies nachbauen."], heikel=True)
    qa(doc, "Was bedeuten DeployIfNotExists / Modify in den Assignments?",
       ["Viele Assignments korrigieren/ergänzen Ressourcen automatisch (z. B. Diagnose, Monitoring-Agent, "
        "Defender-Onboarding) und legen dafür automatisch Remediation-Identities (Managed Identities) an.",
        "Das ist Microsoft-Standard – nichts, was wir manuell pflegen müssen."])
    qa(doc, "Wie funktioniert RBAC – wer bekommt was?",
       ["Das Accelerator-Modul liefert 5 Custom-Rollen (Subscription-Owner, Security-Operations, Network-"
        "Management, Application-Owners, Network-Subnet-Contributor) und weist Rollen an Entra-ID-Gruppen je MG zu.",
        "Object-IDs werden über die Konfiguration injiziert – keine Hardcodes. Konkrete Gruppen-IDs braucht ihr vom Kunden."])
    qa(doc, "Nutzt ihr PIM / Just-in-Time-Zugriff?",
       ["PIM ist empfohlen und kompatibel (Eligible-Zuweisungen auf die Gruppen). Aktivierung ist eine "
        "Kundenentscheidung im RBAC-Modell."])

    # C. Netzwerk
    section(doc, "C. Netzwerk & Connectivity (Schwerpunkt Netzwerker)")
    qa(doc, "Warum /22 für den Hub – reicht das?",
       ["/22 (1024 Adressen) deckt AzureFirewallSubnet (/26), AzureBastionSubnet (/26), GatewaySubnet (/27) "
        "und Resolver-Subnetze mit Reserve ab.",
        "Bei sehr vielen Spokes/Diensten Adressplan früh größer wählen. Bestand on-prem entscheidet."], heikel=True)
    qa(doc, "Ist das VNet-Peering transitiv?",
       ["Nein. VNet-Peering ist nicht transitiv. Spoke-zu-Spoke läuft per UDR über die Hub-Firewall (Hairpin), "
        "nicht direkt zwischen Spokes."])
    qa(doc, "Standard- oder Premium-Firewall?",
       ["Aktuell Standard. Premium bietet IDPS, TLS-Inspection und URL-Filtering – ca. doppelte Kosten.",
        "Frage zurück: Erwartet ihr Layer-7-Inspection/TLS-Aufbruch? Dann planen wir Premium ein."], heikel=True)
    qa(doc, "Wie sieht das Egress-Regelwerk aus?",
       ["Die Azure Firewall hat eine zugeordnete Firewall Policy (DNS-Proxy aktiv). Grundsatz bleibt Default-Deny "
        "für nicht freigegebenen Verkehr.",
        "Alle Workload-spezifischen Freigaben kommen kontrolliert und versioniert in die Firewall Policy – nicht als Portal-Klick."])
    qa(doc, "Wie bindet ihr On-Prem an – VPN oder ExpressRoute?",
       ["Beide Gateways sind im Microsoft-Default des Accelerators in beiden Hubs aktiviert (VPN Gateway "
        "VpnGw1AZ + ExpressRoute Gateway). VPN für schnellen/günstigen Start, ExpressRoute für Bandbreite/SLA.",
        "Wenn nur eines davon gebraucht wird, schalten wir das andere über die deploy*-Schalter aus – das spart "
        "Kosten. Eure Entscheidung nach Bandbreite, SLA und Budget."], heikel=True)
    qa(doc, "Unterstützt ihr BGP? Welches ASN nutzt Azure?",
       ["Ja, BGP wird unterstützt. Azure VPN Gateway nutzt standardmäßig ASN 65515 – euer On-Prem-ASN darf "
        "nicht kollidieren. ASN bitte vorab nennen."])
    qa(doc, "Wie verhindert ihr asymmetrisches Routing bei ExpressRoute + Firewall?",
       ["Über saubere UDR-/BGP-Propagation und ggf. Forced Tunneling. Das lösen wir im Netzwerk-Detail-Design "
        "sauber – ist ein bekannter Fallstrick, kein Showstopper."], heikel=True)
    qa(doc, "Ist Forced Tunneling (alles über On-Prem) möglich?",
       ["Ja – per UDR/BGP wird der Egress über das Gateway zurück nach on-prem gezwungen. Bewusst zu planen, "
        "da es die Firewall-Egress-Logik und ggf. Internetzugang verändert."])
    qa(doc, "Wie segmentiert ihr – NSG oder Firewall?",
       ["Beides komplementär: Azure Firewall für Nord-Süd und zentrale L3/L4(-L7)-Kontrolle, NSGs/ASGs für "
        "Ost-West auf Subnetz-/NIC-Ebene innerhalb der Spokes."])
    qa(doc, "Verfügbarkeitszonen / Hochverfügbarkeit?",
       ["Azure Firewall und Gateways können zonenredundant betrieben werden; zwei Hub-Regionen (GWC + NE) sind "
        "bidirektional gepeert. Konkrete Zonen-/HA-Stufe nach Verfügbarkeits-Anforderung."])
    qa(doc, "Wie funktioniert Internet-Ingress für öffentliche Workloads?",
       ["Für die Online-Landing-Zone ist Application Gateway/WAF vorgesehen – aktuell Roadmap, noch nicht gebaut.",
        "Egress läuft über die Firewall (SNAT); Ingress-Design je nach Workload (WAF, Front Door)."], heikel=True)
    qa(doc, "Wie hoch ist der Firewall-Durchsatz?",
       ["Azure Firewall skaliert automatisch in den zweistelligen Gbit/s-Bereich; SNAT-Ports sind die übliche "
        "Grenze und werden über mehrere Public IPs erweitert."])
    qa(doc, "DDoS-Schutz?",
       ["DDoS Network Protection ist im Microsoft-Default des Accelerators aktiviert (Hub 1) – mit ~€2.500/Monat "
        "der teuerste Einzelposten.",
        "Für internetseitige Workloads sinnvoll; sonst bewusst über den Schalter deployDdosProtectionPlan "
        "deaktivieren und auf den IP-Basisschutz setzen."], heikel=True)

    # D. DNS
    section(doc, "D. DNS & Namensauflösung")
    qa(doc, "Wie löst ihr Private Endpoints auf?",
       ["Zentrale Private DNS Zones (privatelink.*) am Hub, mit den VNets verknüpft. Private Endpoints "
        "registrieren sich automatisch; die Policy Deploy-Private-DNS-Zones erzwingt das für PaaS-Dienste."])
    qa(doc, "Wie funktioniert hybride Namensauflösung zu On-Prem?",
       ["Über den DNS Private Resolver – im Microsoft-Default in beiden Hubs aktiviert (Inbound-/Outbound-"
        "Endpoints in delegierten Subnetzen): Inbound für Azure→On-Prem-Anfragen, Outbound mit Forwarding-Regeln "
        "zu euren On-Prem-DNS-Servern.",
        "Wir brauchen: autoritative On-Prem-DNS-Server und die Domänen für Conditional Forwarding."], heikel=True)
    qa(doc, "DNS Private Resolver vs. eigene DNS-VMs?",
       ["Resolver ist PaaS, zonenredundant, kein VM-Betrieb – günstiger und wartungsärmer als ein DC/DNS-VM-Paar. "
        "Wir empfehlen den Resolver, sofern keine speziellen On-Prem-Abhängigkeiten dagegensprechen."])
    qa(doc, "Warum DNS-Proxy auf der Firewall?",
       ["Damit FQDN-basierte Firewall-Regeln konsistent auflösen und DNS zentral über die Firewall läuft – "
        "einheitlicher Auflösungspfad für alle Spokes."])

    # E. Security
    section(doc, "E. Security")
    qa(doc, "Wie wird Defender for Cloud aktiviert?",
       ["Automatisch über die Policy-Assignments auf der alz-Root-Ebene (Deploy-MDFC-Config-H224, Deploy-"
        "MDEndpoints, Deploy-MDFC-OssDb, Deploy-MDFC-SqlAtp) – DeployIfNotExists, kein eigenes Security-Template.",
        "Neue Subscriptions unter alz werden so beim Eintritt automatisch onboardet."])
    qa(doc, "Welche Defender-Pläne, welcher Tier?",
       ["Servers/Endpoint, SQL (Server/MI/VM), Open-Source-DB, Storage, Container u. a. – über die MDFC-Initiative gesteuert.",
        "Die Policy hebt die Pläne auf den Standard-Tier (kostenpflichtig, i. d. R. pro Ressource). Der Tier ist über die Konfiguration steuerbar."], heikel=True)
    qa(doc, "Habt ihr Sentinel / ein SIEM?",
       ["Sentinel-Onboarding ist über einen Schalter am Log Analytics Workspace vorbereitet. Data Connectors "
        "und Analyseregeln sind Roadmap und werden bei SIEM-Bedarf ausgestaltet."])
    qa(doc, "Wie werden Activity Logs zentralisiert?",
       ["Pro Subscription per Diagnostic Settings ins zentrale Log Analytics (Administrative, Security, Policy, "
        "ResourceHealth u. a.)."])
    qa(doc, "Wie lange werden Logs aufbewahrt?",
       ["Standard 365 Tage (PerGB2018). Anpassbar nach Compliance – längere Retention/Archiv treibt Kosten."])

    # F. Identity
    section(doc, "F. Identity")
    qa(doc, "Entra-only oder hybrid?",
       ["MG alz-platform-identity und eine Identity-Subscription sind vorgesehen. Ausprägung (Entra-only, "
        "hybrid, AD DS in Azure) ist Roadmap und richtet sich nach euren Anforderungen.",
        "Frage zurück: Habt ihr On-Prem-AD, das hybrid angebunden werden soll?"], heikel=True)
    qa(doc, "Diagnostiziert ihr Entra ID nach Log Analytics?",
       ["Geplant: Entra-Diagnose-Logs (Sign-ins, Audit) ins zentrale Log Analytics. Teil der Identity-Roadmap."])

    # G. Betrieb / IaC
    section(doc, "G. Betrieb, IaC & CI/CD")
    qa(doc, "Wie deployt ihr – manuell oder Pipeline?",
       ["Über das offizielle ALZ-Modul: Deploy-Accelerator führt den Bootstrap aus (erstellt Repo, OIDC, "
        "Pipelines); danach fahren die generierten Pipelines die 18 geordneten Deploy-Stufen.",
        "Reihenfolge: Governance/MGs (1–15) → Core-Logging (16) → Networking Hub bzw. Virtual WAN (17/18)."])
    qa(doc, "Wie testet ihr Änderungen vorab?",
       ["Azure What-If wird über die Pipeline vor dem Apply ausgeführt und zeigt exakt, was sich ändert – ohne "
        "eine Ressource zu erstellen. Apply erst nach Freigabe über das Approval-Gate. Risikofrei und kostenlos."])
    qa(doc, "Wie läuft die Pipeline-Anmeldung an Azure?",
       ["Passwortlos über OIDC Federated Identity – keine gespeicherten Secrets/Zertifikate. Der Bootstrap legt "
        "dafür eine Managed Identity mit Federated Credentials und die GitHub-Environments mit Approval-Gates an."])
    qa(doc, "Was, wenn ein Deployment fehlschlägt – Rollback?",
       ["Bicep ist deklarativ/idempotent: erneutes Deployment des letzten guten Stands stellt den Soll-Zustand "
        "wieder her. What-If vorab minimiert Fehldeployments."])
    qa(doc, "Wie verhindert ihr Configuration Drift?",
       ["Git ist die Quelle der Wahrheit; Re-Deployment korrigiert manuelle Änderungen. Policies (Deny) verhindern "
        "viele Drifts bereits an der Quelle."])

    # H. Kosten
    section(doc, "H. Kosten")
    body(doc, "WICHTIG: Der Microsoft-Default des Accelerators aktiviert alle Netzwerk-Dienste in BEIDEN Regionen "
              "(Summe ≈ €5.800/Monat). Richtwerte (GWC, ohne Workloads) – immer als „circa, abhängig von Nutzung“:")
    add_table(doc, ["Komponente", "Grundkosten/Monat", "Default"], [
        ["Management Groups, volles Policy-Set, RBAC", "kostenlos", "an"],
        ["Log Analytics (geringer Ingress)", "~0 (5 GB frei)", "an"],
        ["Private DNS Zones", "~€15", "an"],
        ["2x Azure Firewall (Standard)", "~€2.200", "an"],
        ["2x Azure Bastion", "~€240", "an"],
        ["2x VPN Gateway (VpnGw1AZ)", "~€280", "an"],
        ["2x ExpressRoute Gateway", "~€560", "an"],
        ["2x DNS Private Resolver", "~€50", "an"],
        ["DDoS Network Protection (Hub 1)", "~€2.500", "an"],
        ["Defender Standard (je Ressource)", "ab ~€13/Server", "via Policy an"],
    ], widths=[3.2, 2.2, 1.6])
    qa(doc, "Was kostet die Plattform-Grundlast im Monat?",
       ["Im vollen Microsoft-Default grob ~€5.800/Monat (doppelte Firewalls/Bastion/Gateways + DDoS).",
        "Kostenarm: network_type=none (nur MGs + volles Policy-Set + Logging) ≈ €0; oder einzelne deploy*-"
        "Schalter aus (z. B. nur VNets + Private DNS Zones ≈ €15/Monat). Ideal für eine risikofreie Startphase."], heikel=True)
    qa(doc, "Was sind die größten Kostentreiber?",
       ["DDoS Network Protection (~€2.500), die doppelten Azure Firewalls (~€2.200) und die Gateways. Alles über "
        "deploy*-Schalter steuerbar – im Default sind sie aber AN, das muss man bewusst entscheiden."])
    qa(doc, "Wie können wir Kosten optimieren?",
       ["Netzwerk-Dienste nur dort/in der Region aktivieren, wo gebraucht (deploy*-Schalter), DDoS bewusst "
        "abwägen, Firewall-SKU Standard statt Premium, Defender selektiv pro Plan, Log-Retention nach Compliance.",
        "Für den Erst-Rollout network_type=none nutzen und Netzwerk später gezielt zuschalten."])

    # I. Migration
    section(doc, "I. Workloads & Migration")
    qa(doc, "Wie kommt unsere erste Workload in die Landing Zone?",
       ["Subscription in die passende MG einhängen (corp/online/local), Spoke-VNet ausrollen (Route Table → "
        "Firewall, Hub-Peering, DNS-Links), dann Workload deployen. Policies greifen automatisch."])
    qa(doc, "Wie platziert der Accelerator Subscriptions?",
       ["Über die Management-Group-Struktur und den Bootstrap: bestehende Subscriptions werden in die Ziel-MG "
        "eingehängt und erben sofort deren Policies und RBAC.",
        "Ein eigenes Subscription-Vending-Template wird nicht mehr aktiv betrieben (liegt archiviert unter legacy-custom/)."])
    qa(doc, "Können wir bestehende Subscriptions einbinden?",
       ["Ja – sie werden in die Hierarchie verschoben und erben sofort Policies und RBAC der Ziel-MG. Keine "
        "Billing-Rechte nötig."])
    qa(doc, "Wir haben aktuell nur eine Subscription – geht das?",
       ["Ja. Im Accelerator setzt man dann alle vier Platform-Subscription-IDs (management/identity/connectivity/"
        "security) auf dieselbe ID; die volle MG-Hierarchie und Policies werden trotzdem erstellt.",
        "Die Trennung kann später erfolgen (Subscriptions in die jeweilige MG verschieben). Multi-Subscription ist "
        "das Produktiv-Ziel, kein Muss zum Start."])

    # Gegenfragen
    section(doc, "J. Gegenfragen, die DU stellen solltest")
    for q in [
        "Welche IP-Bereiche sind on-prem belegt? (Überlappungsfreiheit für Hubs/Spokes)",
        "VPN oder ExpressRoute – Bandbreite, SLA, Provider, Redundanz?",
        "Euer On-Prem-ASN und ist BGP gewünscht?",
        "Wo liegen eure autoritativen DNS-Server und welche Domänen brauchen Forwarding?",
        "Erwartet ihr L7-Inspection/TLS-Aufbruch an der Firewall? (Standard vs. Premium)",
        "Habt ihr On-Prem-AD für hybride Identity?",
        "Welche Entra-Gruppen (Object-IDs) sollen welche Rollen je MG bekommen?",
        "Wollt ihr den vollen Netzwerk-Default (~€5.800/Mon.) oder kostenarm starten (network_type=none / Schalter aus)?",
        "Compliance-Pflichten (ISO/BSI/CIS), für die DoNotEnforce-Guardrails scharfgeschaltet werden sollen?",
        "GitHub-Org und PAT (repo/workflow/admin:org) für den Bootstrap vorhanden?",
        "EA/MCA vorhanden und wer hat Billing-Rechte? (Subscription-Neuanlage)",
        "Wer darf Elevated Access am Tenant Root aktivieren? (MG-Erstellung)",
    ]:
        bp = doc.add_paragraph(style='List Bullet'); bp.paragraph_format.space_after = Pt(3)
        run = bp.add_run(q); run.font.size = Pt(10.5); run.font.color.rgb = DARK

    # Top-Fallen
    section(doc, "K. Fünf Sätze, die du parat haben solltest")
    for q, col in [
        ("„Wir setzen das offizielle Microsoft-Accelerator-Modul ein – volles ALZ-Policy-Set, kein Eigenbau.“", GREEN),
        ("„What-If zeigt jede Änderung vorab – ohne Kosten, ohne Risiko.“", GREEN),
        ("„Der Default aktiviert alles in beiden Regionen (~€5.800/Mon.) – kostenarm starten wir mit network_type=none.“", GREEN),
        ("„Peering ist nicht transitiv – Spoke-zu-Spoke läuft kontrolliert über die Firewall.“", GREEN),
        ("„Das lösen wir sauber im Netzwerk-Detail-Design.“ (bei tiefen Routing-/DNS-Fragen)", DARK),
        ("„Das ist Roadmap, nicht live.“ (bei dedizierten Identity-Ressourcen, Sentinel-Connectors, WAF)", DARK),
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        r = p.add_run("•  " + q); r.font.size = Pt(11); r.bold = True; r.font.color.rgb = col

    # Footer
    sec = doc.sections[0]
    f = sec.footer.paragraphs[0]; f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = f.add_run("Azure Landing Zone – Q&A-Vorbereitung · intern · vertraulich"); fr.font.size = Pt(8); fr.font.color.rgb = GREY

    path = f"{OUT}/Word/Azure-Landing-Zone-QnA-Vorbereitung.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print("DOCX:", build())
