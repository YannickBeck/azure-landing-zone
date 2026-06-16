# -*- coding: utf-8 -*-
"""Generiert eine stark granulare technische Referenz der ALZ aus dem echten Repo.
Parst Policy-JSONs, Assignments je MG, AVM-Versionen, Netzwerk-Settings und
Deploy-Reihenfolge. Rendert Markdown (docs/) + Word (docs/kickoff/).
"""
import json, glob, re, os, subprocess, datetime

REPO = "/home/user/azure-landing-zone"
GOV  = f"{REPO}/templates/core/governance"
LIB  = f"{GOV}/lib/alz"
MGD  = f"{GOV}/mgmt-groups"
CFG  = f"{REPO}/.config/ALZ-Powershell.config.json"
DATE = "16.06.2026"

def short_commit():
    try:
        return subprocess.check_output(["git","-C",REPO,"rev-parse","--short","HEAD"]).decode().strip()
    except Exception:
        return "n/a"

# ============================ DATEN-EXTRAKTION ============================

ASSIGN_PATH_RE = re.compile(r"lib/alz/([A-Za-z0-9/_-]+\.alz_policy_assignment\.json)")

MG_FILES = [
    ("alz (Intermediate Root)", "int-root/main.bicep"),
    ("landingzones", "landingzones/main.bicep"),
    ("landingzones-corp", "landingzones/landingzones-corp/main.bicep"),
    ("landingzones-online", "landingzones/landingzones-online/main.bicep"),
    ("landingzones-local", "landingzones/landingzones-local/main.bicep"),
    ("platform", "platform/main.bicep"),
    ("platform-connectivity", "platform/platform-connectivity/main.bicep"),
    ("platform-identity", "platform/platform-identity/main.bicep"),
    ("platform-management", "platform/platform-management/main.bicep"),
    ("platform-security", "platform/platform-security/main.bicep"),
    ("sandbox", "sandbox/main.bicep"),
    ("decommissioned", "decommissioned/main.bicep"),
]

def assignments_per_mg():
    """Liefert je MG eine Liste (name, displayName, enforcementMode) – aus den
    loadJsonContent-Pfaden der main.bicep + den geladenen Assignment-JSONs."""
    out = []
    for label, rel in MG_FILES:
        path = f"{MGD}/{rel}"
        items = []
        if os.path.exists(path):
            txt = open(path, encoding="utf-8").read()
            for relp in sorted(set(ASSIGN_PATH_RE.findall(txt))):
                name = os.path.basename(relp).replace(".alz_policy_assignment.json", "")
                disp, enf = "", "Default"
                jp = f"{LIB}/{relp.split('lib/alz/')[-1]}" if "lib/alz/" in relp else f"{LIB}/{relp}"
                if os.path.exists(jp):
                    try:
                        d = json.load(open(jp, encoding="utf-8")); pr = d.get("properties", {})
                        disp = pr.get("displayName", "") or ""
                        enf = pr.get("enforcementMode", "Default") or "Default"
                    except Exception:
                        pass
                items.append((name, disp, enf))
        out.append((label, items))
    return out

def effect_of(p):
    params = p.get("parameters") or {}
    for key in ("effect", "Effect"):
        if key in params and isinstance(params[key], dict):
            dv = params[key].get("defaultValue")
            if dv: return dv
            av = params[key].get("allowedValues")
            if av: return "/".join(map(str, av))
    then = (p.get("policyRule") or {}).get("then") or {}
    e = then.get("effect", "")
    return e if isinstance(e, str) and not e.startswith("[") else "—"

def load_defs():
    rows = []
    for f in sorted(glob.glob(f"{LIB}/*.alz_policy_definition.json")):
        try:
            d = json.load(open(f, encoding="utf-8")); p = d.get("properties", {})
            rows.append((d.get("name") or os.path.basename(f),
                         p.get("displayName", ""),
                         (p.get("metadata") or {}).get("category", ""),
                         effect_of(p)))
        except Exception:
            pass
    return rows

def load_sets():
    rows = []
    for f in sorted(glob.glob(f"{LIB}/*.alz_policy_set_definition.json")):
        try:
            d = json.load(open(f, encoding="utf-8")); p = d.get("properties", {})
            rows.append((d.get("name") or os.path.basename(f),
                         p.get("displayName", ""),
                         (p.get("metadata") or {}).get("category", ""),
                         str(len(p.get("policyDefinitions", [])))))
        except Exception:
            pass
    return rows

def load_roles():
    rows = []
    for f in sorted(glob.glob(f"{LIB}/*.alz_role_definition.json")):
        try:
            d = json.load(open(f, encoding="utf-8")); p = d.get("properties", {})
            perms = p.get("permissions", [{}])
            actions = perms[0].get("actions", []) if perms else []
            rows.append((p.get("roleName", d.get("name", os.path.basename(f))),
                         p.get("description", "")[:80],
                         str(len(actions))))
        except Exception:
            pass
    return rows

def avm_modules():
    out = subprocess.check_output(
        ["bash","-lc", f"grep -rhoE 'br/public:avm/[a-z/-]+:[0-9.]+' {REPO}/templates | sort | uniq -c | sort -rn"]
    ).decode().strip().splitlines()
    rows = []
    for line in out:
        m = re.match(r"\s*(\d+)\s+br/public:(.+):([0-9.]+)", line)
        if m:
            rows.append((m.group(2), m.group(3), m.group(1)))
    return rows

def deployment_order():
    d = json.load(open(CFG, encoding="utf-8"))
    plz = d["starter_modules"]["platform_landing_zone"]["deployment_files"]
    rows = []
    for f in plz:
        rows.append((str(f["order"]), f["displayName"], f["deploymentType"],
                     f.get("templateFilePath","").replace("templates/","")))
    return rows

def cat_summary(defs):
    from collections import Counter
    c = Counter(r[2] for r in defs)
    return sorted(c.items(), key=lambda x: (-x[1], x[0]))

# Statische, aus den bicepparam gelesene Netzwerk-Struktur (Hub-and-Spoke Default)
HUB_SUBNETS = [
    ("AzureFirewallSubnet", "10.0.0.0/26", "10.1.0.0/26"),
    ("AzureFirewallManagementSubnet", "10.0.0.192/26", "10.1.0.192/26"),
    ("AzureBastionSubnet", "10.0.0.64/26", "10.1.0.64/26"),
    ("GatewaySubnet", "10.0.0.128/27", "10.1.0.128/27"),
    ("DNSPrivateResolverInboundSubnet", "10.0.0.160/28", "10.1.0.160/28"),
    ("DNSPrivateResolverOutboundSubnet", "10.0.0.176/28", "10.1.0.176/28"),
]
HUB_SERVICES = [
    ("Azure Firewall (Standard)", "deployAzureFirewall", "true", "true", "~€1.100/Hub/Monat"),
    ("Azure Bastion (Standard)", "deployBastion", "true", "true", "~€120/Hub/Monat"),
    ("VPN Gateway (VpnGw1AZ, activeActiveBgp, ASN 65515)", "deployVpnGateway", "true", "true", "~€140/Hub/Monat"),
    ("ExpressRoute Gateway", "deployExpressRouteGateway", "true", "true", "~€280/Hub/Monat"),
    ("DDoS Network Protection", "deployDdosProtectionPlan", "true", "false", "~€2.500/Monat (nur Hub 1)"),
    ("Private DNS Zones", "deployPrivateDnsZones", "true", "true", "~€15/Monat"),
    ("DNS Private Resolver", "deployDnsPrivateResolver", "true", "true", "~€25/Hub/Monat"),
]
LOGGING_RES = [
    ("Resource Group", "rg-alz-logging-<region>", "avm/res/resources/resource-group:0.4.3"),
    ("Log Analytics Workspace", "law-alz-<region> (PerGB2018, 365 Tage Retention)", "avm/res/operational-insights/workspace:0.14.2"),
    ("LAW Solution", "ChangeTracking", "(im Workspace)"),
    ("User-Assigned Managed Identity", "mi-alz-<region>", "avm/ptn/alz/ama:0.2.0"),
    ("Data Collection Rule – VM Insights", "dcr-vmi-alz-<region>", "avm/ptn/alz/ama:0.2.0"),
    ("Data Collection Rule – Change Tracking", "dcr-ct-alz-<region>", "avm/ptn/alz/ama:0.2.0"),
    ("Data Collection Rule – Defender SQL", "dcr-mdfcsql-alz-<region>", "avm/ptn/alz/ama:0.2.0"),
    ("Automation Account (optional, default aus)", "aa-alz-<region> (Basic)", "avm/res/automation/automation-account:0.17.1"),
]
MG_HIER = [
    ("alz", "Azure Landing Zones", "Tenant Root", "Intermediate Root"),
    ("platform", "Platform", "alz", "Plattform-Dienste"),
    ("connectivity", "Connectivity", "platform", "Hub-Netzwerk, Firewall, DNS"),
    ("identity", "Identity", "platform", "Identity-Dienste"),
    ("management", "Management", "platform", "Logging, Monitoring"),
    ("security", "Security", "platform", "Security-Tooling"),
    ("landingzones", "Landing Zones", "alz", "Workload-Container"),
    ("corp", "Corp", "landingzones", "interne Workloads (keine Public Endpoints)"),
    ("online", "Online", "landingzones", "internetseitige Workloads"),
    ("local", "Local", "landingzones", "souveräne/lokale Workloads"),
    ("sandbox", "Sandbox", "alz", "Experimente (gelockerte Policies)"),
    ("decommissioned", "Decommissioned", "alz", "Stilllegung (gesperrt)"),
]

# ============================ MARKDOWN-RENDERER ============================

def md_table(headers, rows):
    out = ["| " + " | ".join(headers) + " |", "|" + "|".join(["---"]*len(headers)) + "|"]
    for r in rows:
        out.append("| " + " | ".join(str(c).replace("|","\\|") for c in r) + " |")
    return "\n".join(out) + "\n"

def build_markdown(data):
    C = short_commit()
    L = []
    L.append(f"# Azure Landing Zone – Technische Referenz\n")
    L.append(f"> Granulare, datengetriebene Dokumentation aller deployten Ressourcen, Policies, "
             f"Management Groups und Module. **Automatisch generiert** aus dem Repository-Stand "
             f"`{C}` am {DATE} via `docs/kickoff/generate-techref.py`.\n")
    L.append("Basis: offizielles **Azure ALZ Bicep Accelerator** Starter-Modul "
             "(`avm/ptn/alz/empty`, volles ALZ-Policy-Set).\n")

    L.append("## 1. Deployment-Topologie & Reihenfolge\n")
    L.append("Die Plattform wird in dieser festen Reihenfolge deployt "
             "(`.config/ALZ-Powershell.config.json`):\n")
    L.append(md_table(["#","Stufe","Scope","Template"], data["order"]))

    L.append("## 2. Management Groups\n")
    L.append(f"{len(MG_HIER)} Management Groups (IDs ggf. mit Präfix/Postfix gemäß `config/`):\n")
    L.append(md_table(["ID","Anzeigename","Parent","Zweck"], MG_HIER))

    L.append("## 3. Azure Verified Modules (Versionen)\n")
    L.append("Alle verwendeten AVM-Module aus der öffentlichen MCR:\n")
    L.append(md_table(["Modul","Version","Verwendungen"], data["avm"]))

    L.append("## 4. Governance – Policy-Assignments je Management-Group-Ebene\n")
    L.append("Konkret zugewiesene Policies/Guardrails (extrahiert aus den Templates; "
             "Vererbung an Kind-MGs greift zusätzlich):\n")
    total = 0
    for label, items in data["assign"]:
        L.append(f"### {label} — {len(items)} Assignment(s)\n")
        total += len(items)
        if items:
            L.append(md_table(["Assignment","Anzeigename","Enforcement"], items))
        else:
            L.append("_keine eigenen Assignments (erbt von übergeordneter MG)_\n")
    L.append(f"\n**Summe direkt zugewiesener Assignments: {total}** (zzgl. Vererbung an Kind-MGs).\n")

    L.append("## 5. Governance – Policy-Definitionen (Custom)\n")
    L.append(f"**{len(data['defs'])}** Custom-Policy-Definitionen werden via `loadJsonContent` "
             f"aus `lib/alz/` geladen. Verteilung nach Kategorie:\n")
    L.append(md_table(["Kategorie","Anzahl"], data["catsum"]))
    L.append("\nVollständige Liste:\n")
    L.append(md_table(["Name","Anzeigename","Kategorie","Default-Effekt"], data["defs"]))

    L.append("## 6. Governance – Initiativen (Policy-Set-Definitionen)\n")
    L.append(f"**{len(data['sets'])}** Custom-Initiativen:\n")
    L.append(md_table(["Name","Anzeigename","Kategorie","# Policies"], data["sets"]))

    L.append("## 7. Governance – Custom Role-Definitionen\n")
    L.append(f"**{len(data['roles'])}** Custom-Rollen:\n")
    L.append(md_table(["Rolle","Beschreibung","# Actions"], data["roles"]))

    L.append("## 8. Logging-Ressourcen (Management-Subscription)\n")
    L.append(md_table(["Ressource","Name / Konvention","AVM-Modul"], LOGGING_RES))

    L.append("## 9. Hub-Networking-Ressourcen (Connectivity-Subscription)\n")
    L.append("Zwei Hub-VNets, bidirektional gepeert: **Hub 1** `10.0.0.0/22` (primär), "
             "**Hub 2** `10.1.0.0/22` (sekundär). Resource Group `rg-alz-conn-<region>`.\n")
    L.append("### Subnetze je Hub\n")
    L.append(md_table(["Subnetz","Hub 1 (primär)","Hub 2 (sekundär)"], HUB_SUBNETS))
    L.append("### Dienste & Default-Schalter (Microsoft-Standard)\n")
    L.append(md_table(["Dienst","Schalter","Hub 1 Default","Hub 2 Default","Kosten"], HUB_SERVICES))
    L.append("> ⚠️ **Alle Dienste sind im Microsoft-Default aktiviert** (Summe ≈ €5.800/Monat). "
             "Für kostenarme Rollouts siehe `docs/ACCELERATOR-BOOTSTRAP.md` (kostenarme Variante).\n")

    L.append("## 10. Virtual WAN (Alternative zu Hub-and-Spoke)\n")
    L.append("Aktiv nur bei `network_type: vwanConnectivity`. Ressourcen: Virtual WAN "
             "(`vwan-alz-<region>`, Standard), Virtual Hub (`vhub-alz-<region>`, `10.0.0.0/22`), "
             "Azure Firewall, ExpressRoute Gateway (default an), S2S/P2S VPN Gateway (default aus). "
             "Module: `avm/res/network/virtual-wan:0.4.3`, `virtual-hub:0.4.3`.\n")

    return "\n".join(L)

# ============================ WORD-RENDERER ============================

def build_docx(data, out_path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    BLUE=RGBColor(0x00,0x78,0xD4); DARK=RGBColor(0x24,0x3A,0x5E); GREY=RGBColor(0x60,0x5E,0x5C)

    def shade(cell,hexc):
        tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc); tcPr.append(shd)
    def sct(cell,text,bold=False,white=False,size=8):
        cell.text=""; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(1); p.paragraph_format.space_before=Pt(1)
        r=p.add_run(str(text)); r.bold=bold; r.font.size=Pt(size)
        if white: r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        else: r.font.color.rgb=DARK
    def table(headers,rows,widths=None,size=8):
        t=doc.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'
        for i,h in enumerate(headers): sct(t.rows[0].cells[i],h,bold=True,white=True,size=size); shade(t.rows[0].cells[i],'0078D4')
        for row in rows:
            cells=t.add_row().cells
            for i,v in enumerate(row): sct(cells[i],v,size=size)
        if widths:
            for i,w in enumerate(widths):
                for row in t.rows: row.cells[i].width=Inches(w)
        doc.add_paragraph()
    def h1(t): doc.add_heading(t,level=1)
    def h2(t): doc.add_heading(t,level=2)
    def body(t,color=None,italic=False):
        p=doc.add_paragraph(); r=p.add_run(t); r.italic=italic
        if color: r.font.color.rgb=color
        p.paragraph_format.space_after=Pt(6)

    doc=Document()
    n=doc.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(9.5)
    for lvl,col,sz in [('Heading 1',DARK,15),('Heading 2',BLUE,12)]:
        st=doc.styles[lvl]; st.font.color.rgb=col; st.font.name='Calibri'; st.font.size=Pt(sz)

    for _ in range(3): doc.add_paragraph()
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("Azure Landing Zone"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=BLUE
    s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=s.add_run("Technische Referenz – Ressourcen, Policies, Module"); r.font.size=Pt(15); r.font.color.rgb=DARK
    s2=doc.add_paragraph(); s2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=s2.add_run(f"Automatisch generiert aus Repo-Stand {short_commit()} · {DATE}"); r.font.size=Pt(10); r.font.color.rgb=GREY
    doc.add_page_break()

    h1("1. Deployment-Topologie & Reihenfolge")
    body("Feste Deployment-Reihenfolge gemäß .config/ALZ-Powershell.config.json:")
    table(["#","Stufe","Scope","Template"], data["order"], widths=[0.4,2.8,1.3,2.5])

    h1("2. Management Groups")
    table(["ID","Anzeigename","Parent","Zweck"], MG_HIER, widths=[1.4,1.6,1.4,2.6])

    h1("3. Azure Verified Modules (Versionen)")
    table(["Modul","Version","Verw."], data["avm"], widths=[4.6,1.2,0.8])

    h1("4. Policy-Assignments je Management-Group-Ebene")
    body("Konkret zugewiesene Guardrails (extrahiert aus den Templates). Vererbung greift zusätzlich.")
    for label,items in data["assign"]:
        h2(f"{label} — {len(items)} Assignment(s)")
        if items: table(["Assignment","Anzeigename","Enforce"], items, widths=[2.1,4.0,0.7], size=7.5)
        else: body("keine eigenen Assignments (erbt von übergeordneter MG)", color=GREY, italic=True)

    h1("5. Policy-Definitionen (Custom)")
    body(f"{len(data['defs'])} Custom-Definitionen. Verteilung nach Kategorie:")
    table(["Kategorie","Anzahl"], data["catsum"], widths=[5.0,1.0])
    body("Vollständige Liste:")
    table(["Name","Anzeigename","Kategorie","Effekt"], data["defs"], widths=[1.9,2.8,1.3,0.8], size=7)

    h1("6. Initiativen (Policy-Set-Definitionen)")
    table(["Name","Anzeigename","Kategorie","#Pol"], data["sets"], widths=[1.9,2.9,1.3,0.6], size=7)

    h1("7. Custom Role-Definitionen")
    table(["Rolle","Beschreibung","#Act"], data["roles"], widths=[2.2,4.0,0.6])

    h1("8. Logging-Ressourcen")
    table(["Ressource","Name / Konvention","AVM-Modul"], LOGGING_RES, widths=[2.3,2.6,2.0])

    h1("9. Hub-Networking-Ressourcen")
    body("Zwei Hub-VNets (10.0.0.0/22 primär, 10.1.0.0/22 sekundär), bidirektional gepeert.")
    h2("Subnetze je Hub")
    table(["Subnetz","Hub 1","Hub 2"], HUB_SUBNETS, widths=[3.4,1.7,1.7])
    h2("Dienste & Default-Schalter (Microsoft-Standard – alle an)")
    table(["Dienst","Schalter","H1","H2","Kosten"], HUB_SERVICES, widths=[2.8,1.9,0.5,0.5,1.5], size=7.5)
    body("Achtung: Microsoft-Default aktiviert alle Dienste (Summe ~€5.800/Monat). "
         "Kostenarme Variante siehe docs/ACCELERATOR-BOOTSTRAP.md.", color=RGBColor(0xC0,0x39,0x2B))

    h1("10. Virtual WAN (Alternative)")
    body("Aktiv bei network_type=vwanConnectivity: Virtual WAN (Standard), Virtual Hub (10.0.0.0/22), "
         "Azure Firewall, ExpressRoute Gateway (an), S2S/P2S VPN (aus).")

    sec=doc.sections[0]; f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fr=f.add_run("Azure Landing Zone – Technische Referenz · automatisch generiert · vertraulich")
    fr.font.size=Pt(8); fr.font.color.rgb=GREY
    doc.save(out_path)

# ============================ MAIN ============================

if __name__ == "__main__":
    defs = load_defs()
    data = {
        "order": deployment_order(),
        "avm": avm_modules(),
        "assign": assignments_per_mg(),
        "defs": defs,
        "sets": load_sets(),
        "roles": load_roles(),
        "catsum": cat_summary(defs),
    }
    md = build_markdown(data)
    md_path = f"{REPO}/docs/TECHNICAL-REFERENCE.md"
    open(md_path, "w", encoding="utf-8").write(md)
    print("MARKDOWN:", md_path, f"({len(md.splitlines())} Zeilen)")

    docx_path = f"{REPO}/docs/kickoff/Word/Azure-Landing-Zone-Technische-Referenz.docx"
    build_docx(data, docx_path)
    print("DOCX:", docx_path)
    print(f"Defs={len(data['defs'])} Sets={len(data['sets'])} Roles={len(data['roles'])} "
          f"AVM={len(data['avm'])} Stufen={len(data['order'])}")
