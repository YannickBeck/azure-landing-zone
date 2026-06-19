#!/usr/bin/env python3
"""
Generates: Azure-Landing-Zone-ALZ-Referenz.docx
Vollständige Ressourcen- und Richtlinien-Referenz für das ALZ Bicep Accelerator Setup.
Referenced from Azure-Landing-Zone-Konzept.docx §4.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE  = os.path.join(BASE_DIR, "bechtle-brand", "Word", "VORLAGE_Bechtle_Management_Summary.docx")
OUTPUT    = os.path.join(BASE_DIR, "Word", "Azure-Landing-Zone-ALZ-Referenz.docx")
DATE      = "17.06.2026"

# ─── COLORS ──────────────────────────────────────────────────────────────────
BECHTLE_DARK  = RGBColor(0x05, 0x3B, 0x25)
BECHTLE_MID   = RGBColor(0x07, 0x50, 0x33)
BECHTLE_GREEN = RGBColor(0x23, 0xA9, 0x6A)
BECHTLE_LIME  = RGBColor(0xAA, 0xDE, 0x0C)
GREY_TEXT     = RGBColor(0x59, 0x59, 0x59)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

# ─── HELPERS ─────────────────────────────────────────────────────────────────
def clear_body(doc):
    for el in list(doc.element.body):
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag != "sectPr":
            doc.element.body.remove(el)

def add_body(doc, text, size=11):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = GREY_TEXT
    return p

def add_bullet(doc, text, level=0, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = GREY_TEXT
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = WHITE
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "053B25")
        tcPr.append(shd)
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        fill = "F2F9F5" if r_idx % 2 == 0 else "FFFFFF"
        for i, cell_text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(cell_text)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = GREY_TEXT
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def add_section_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    border_el = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "23A96A")
    border_el.append(left)
    p._p.get_or_add_pPr().append(border_el)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = GREY_TEXT
    run.italic = True
    return p

# ─── DOCUMENT ────────────────────────────────────────────────────────────────
doc = Document(TEMPLATE)
clear_body(doc)

doc.add_heading("Azure Landing Zone – Governance- und Ressourcen-Referenz", level=1)
add_body(doc,
    "Dieses Dokument ist die vollständige technische Referenz zum ALZ Bicep Accelerator Setup "
    "(avm/ptn/alz/empty:0.3.6). Es listet alle Management Groups, Policy-Definitionen, "
    "Initiativen, Assignments, RBAC-Rollen und Bicep-Module konkret und namentlich auf. "
    "Das Konzeptdokument 'Azure-Landing-Zone-Konzept.docx' verweist auf diese Referenz für Details."
)
add_body(doc, f"Stand: {DATE}  |  Quelle: Microsoft ALZ Bicep Accelerator (GitHub: Azure/ALZ-Bicep)")
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
# 1. MANAGEMENT GROUPS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("1. Management Groups (12)", level=1)
add_body(doc,
    "Der ALZ Bicep Accelerator erstellt eine standardisierte Hierarchie von 12 Management Groups "
    "unterhalb des Tenant Root. Jede MG definiert einen Vertrauensbereich mit eigenem "
    "Policy- und RBAC-Scope. Kind-MGs erben alle Assignments der Eltern-MG."
)

add_table(doc,
    ["Management Group", "Display Name", "Parent", "Zweck / Scope"],
    [
        ["Tenant Root Group",              "(vorhanden)",                        "–",                 "Azure-Tenant-Root; kein ALZ-Scope"],
        ["alz",                            "ALZ Intermediate Root",              "Tenant Root",        "Einstieg für alle ALZ-Policies; alle MGs erben von hier"],
        ["alz-platform",                   "Platform",                           "alz",               "Container für Plattform-Dienste-Subscriptions"],
        ["alz-platform-connectivity",      "Connectivity",                       "alz-platform",      "Hub-Netzwerk, Azure Firewall, DNS, VPN/ExpressRoute"],
        ["alz-platform-identity",          "Identity",                           "alz-platform",      "Identity-Dienste (AD DS, Entra Connect – Roadmap)"],
        ["alz-platform-management",        "Management",                         "alz-platform",      "Log Analytics, Azure Monitor, Operations-Tooling"],
        ["alz-platform-security",          "Security",                           "alz-platform",      "Security-Tooling, Defender for Cloud Konfiguration"],
        ["alz-landingzones",               "Landing Zones",                      "alz",               "Container für alle Workload-Subscriptions"],
        ["alz-landingzones-corp",          "Corp",                               "alz-landingzones",  "Interne Workloads; Deny Public Endpoints/IP; Private DNS"],
        ["alz-landingzones-online",        "Online",                             "alz-landingzones",  "Internetseitige Workloads (Public APIs, Websites)"],
        ["alz-landingzones-local",         "Local (Sovereign)",                  "alz-landingzones",  "Souveräne / vertrauliche Workloads; Azure Local Support"],
        ["alz-sandbox",                    "Sandbox",                            "alz",               "Experimente; gelockerte Policies; kein Prod-Scope"],
        ["alz-decommissioned",             "Decommissioned",                     "alz",               "Stilllegungszone; Resource-Erstellung gesperrt"],
    ],
    col_widths=[1.9, 1.7, 1.6, 2.8]
)
add_section_note(doc,
    "Policy-Vererbung: Assignments auf 'alz' gelten für alle 11 Kind-MGs. "
    "Assignments auf 'alz-platform' gelten für connectivity, identity, management und security. "
    "Assignments auf 'alz-landingzones' gelten für corp, online und local."
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 2. POLICY ASSIGNMENTS (123) je MG
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("2. Policy Assignments je Management Group (123 gesamt)", level=1)
add_body(doc,
    "Die 123 Policy Assignments sind konkret auf 9 der 12 Management Groups verteilt. "
    "Assignments bestehen aus einer Kombination von Policy-Definitionen und Initiativen (Policy Sets). "
    "Im Auslieferungszustand sind Enforce-Guardrails-*-Initiativen auf 'DoNotEnforce' gesetzt "
    "(Audit-only); Deploy-* Assignments sind auf 'Enabled' gesetzt."
)

doc.add_heading("2.1 alz (Intermediate Root) – 17 Assignments", level=2)
add_table(doc,
    ["Assignment-Name", "Typ", "Effekt", "Zweck"],
    [
        ["Deploy-MDFC-Config-H224",         "Initiative", "DeployIfNotExists", "Defender for Cloud aktivieren + Security Contacts setzen"],
        ["Deploy-MDEndpoints",              "Initiative", "DeployIfNotExists", "Defender for Endpoint Agent auf VMs (MDE)"],
        ["Deploy-MDEndpointsAMA",           "Initiative", "DeployIfNotExists", "MDE-Integration mit Azure Monitor Agent"],
        ["Deploy-MDFC-OssDb",               "Initiative", "DeployIfNotExists", "Defender for OSS Databases (PostgreSQL, MySQL, MariaDB)"],
        ["Deploy-MDFC-SqlAtp",              "Policy",     "DeployIfNotExists", "Defender für SQL Server (ATP)"],
        ["Deploy-AzActivity-Log",           "Initiative", "DeployIfNotExists", "Activity Logs aller Subscriptions → Log Analytics leiten"],
        ["Deploy-Diagnostics-LogAnalytics", "Initiative", "DeployIfNotExists", "Diagnose-Einstellungen für alle Azure-Ressourcentypen"],
        ["Deploy-VM-Monitoring",            "Initiative", "DeployIfNotExists", "VM Insights + Azure Monitor Agent auf alle VMs"],
        ["Deploy-VMSS-Monitoring",          "Initiative", "DeployIfNotExists", "VM Insights auf VM Scale Sets"],
        ["DenyAction-DeleteUAMIAMA",        "Policy",     "DenyAction",        "Verhindert versehentliches Löschen der User-Assigned Managed Identity für AMA"],
        ["Enforce-ACSB",                    "Initiative", "Audit/Deny",        "Azure Cloud Security Baseline (CIS/NIST-Alignment)"],
        ["Deploy-VM-Backup",                "Initiative", "DeployIfNotExists", "VMs ohne Backup-Tag automatisch in Recovery Vault sichern"],
        ["Audit-Tags-Mandatory-Rg",         "Policy",     "Audit",             "Pflicht-Tags auf Resource Groups prüfen"],
        ["Deny-Classic-Resources",          "Policy",     "Deny",              "Klassische (ASM) Ressourcen verbieten"],
        ["Deny-IP-Forwarding-On-NIC",       "Policy",     "Deny",              "IP-Forwarding auf NICs verbieten (Sicherheitsrisiko)"],
        ["Deny-Subnet-Without-Nsg",         "Policy",     "Deny",              "Subnetze ohne Network Security Group verbieten"],
        ["Deploy-RG-ActivityLog",           "Policy",     "DeployIfNotExists", "Activity Logs auf Resource-Group-Ebene aktivieren"],
    ],
    col_widths=[2.5, 0.8, 1.4, 3.3]
)

doc.add_heading("2.2 alz-landingzones – 53 Assignments", level=2)
add_table(doc,
    ["Assignment-Name", "Typ", "Effekt / DoNotEnforce"],
    [
        ["Deny-Storage-http",               "Policy",     "Deny – Storage-Accounts ohne HTTPS verbieten"],
        ["Deny-Storage-MinTls",             "Policy",     "Deny – TLS < 1.2 verbieten"],
        ["Deny-Storage-PublicAccess",       "Policy",     "Deny – Public Blob Access verbieten"],
        ["Deny-Storage-SharedKey",          "Policy",     "Deny – Shared Key Authentication verbieten (Entra-only)"],
        ["Deny-MgmtPorts-From-Internet",    "Policy",     "Deny – RDP (3389) / SSH (22) aus Internet sperren"],
        ["Deny-Subnet-Without-Nsg",         "Policy",     "Deny – NSG-Pflicht (auch auf LZ-Ebene)"],
        ["Deny-Public-AKS-APIServer-Access","Policy",     "Deny – AKS API-Server nicht öffentlich"],
        ["Deny-MachineLearning-PublicNetworkAccess", "Policy", "Deny – Azure ML ohne Public Network"],
        ["Deny-APIM-Without-SSL",           "Policy",     "Deny – API Management ohne TLS"],
        ["Deny-AppGW-Without-WAF",          "Policy",     "Deny – Application Gateway ohne WAF"],
        ["Deploy-VM-Monitoring",            "Initiative", "DeployIfNotExists – VM Insights auf LZ-VMs"],
        ["Deploy-VMSS-Monitoring",          "Initiative", "DeployIfNotExists – VMSS Monitoring"],
        ["Deploy-VM-Backup",                "Initiative", "DeployIfNotExists – VM Backup Enforcement"],
        ["Enforce-Guardrails-APIM",         "Initiative", "DoNotEnforce – API Management Guardrails"],
        ["Enforce-Guardrails-AppServices",  "Initiative", "DoNotEnforce – App Services Guardrails"],
        ["Enforce-Guardrails-Automation",   "Initiative", "DoNotEnforce – Automation Accounts"],
        ["Enforce-Guardrails-CognitiveServices","Initiative","DoNotEnforce – Cognitive Services / AI"],
        ["Enforce-Guardrails-ContainerApps","Initiative", "DoNotEnforce – Container Apps"],
        ["Enforce-Guardrails-ContainerInstance","Initiative","DoNotEnforce – Azure Container Instances"],
        ["Enforce-Guardrails-ContainerRegistry","Initiative","DoNotEnforce – Container Registry"],
        ["Enforce-Guardrails-CosmosDb",     "Initiative", "DoNotEnforce – Cosmos DB"],
        ["Enforce-Guardrails-DataExplorer", "Initiative", "DoNotEnforce – Azure Data Explorer"],
        ["Enforce-Guardrails-DataFactory",  "Initiative", "DoNotEnforce – Data Factory"],
        ["Enforce-Guardrails-EventGrid",    "Initiative", "DoNotEnforce – Event Grid"],
        ["Enforce-Guardrails-EventHub",     "Initiative", "DoNotEnforce – Event Hub"],
        ["Enforce-Guardrails-KeyVault",     "Initiative", "DoNotEnforce – Key Vault"],
        ["Enforce-Guardrails-Kubernetes",   "Initiative", "DoNotEnforce – AKS Kubernetes"],
        ["Enforce-Guardrails-MachineLearning","Initiative","DoNotEnforce – Azure Machine Learning"],
        ["Enforce-Guardrails-MariaDb",      "Initiative", "DoNotEnforce – MariaDB"],
        ["Enforce-Guardrails-MySQL",        "Initiative", "DoNotEnforce – MySQL Flexible Server"],
        ["Enforce-Guardrails-Network",      "Initiative", "DoNotEnforce – Netzwerk-Sicherheit"],
        ["Enforce-Guardrails-OpenAI",       "Initiative", "DoNotEnforce – Azure OpenAI Service"],
        ["Enforce-Guardrails-PostgreSQL",   "Initiative", "DoNotEnforce – PostgreSQL Flexible Server"],
        ["Enforce-Guardrails-ServiceBus",   "Initiative", "DoNotEnforce – Service Bus"],
        ["Enforce-Guardrails-SQL",          "Initiative", "DoNotEnforce – SQL Server / SQL Databases"],
        ["Enforce-Guardrails-Storage",      "Initiative", "DoNotEnforce – Storage Accounts"],
        ["Enforce-Guardrails-Synapse",      "Initiative", "DoNotEnforce – Azure Synapse Analytics"],
        ["Enforce-Guardrails-VirtualDesktop","Initiative","DoNotEnforce – Azure Virtual Desktop"],
        ["Enforce-Guardrails-AADB2C",       "Initiative", "DoNotEnforce – Azure AD B2C"],
        ["Deploy-Private-DNS-Zones",        "Initiative", "DeployIfNotExists – 59 Private DNS Zonen für PaaS-Dienste"],
        ["Deny-MgmtPorts-Internet",         "Policy",     "Deny – Verwaltungsports aus Internet (alz-weit)"],
        ["DenyAction-DeleteUAMIAMA",        "Policy",     "DenyAction – AMA User-Assigned MI schützen"],
        ["Deny-ERPeering-Without-ER-Direct","Policy",     "Deny – ER-Circuit-Peering ohne ER-Direct"],
        ["Deny-AINE-Pip-OnNic",             "Policy",     "Deny – Public IP auf NIC in AINE-Szenarien"],
        ["Deny-IP-Forwarding-On-NIC",       "Policy",     "Deny – IP-Forwarding auf NICs"],
        ["Deny-SQL-FWRules-For-All-IPs",    "Policy",     "Deny – SQL-FW-Regel 0.0.0.0-255.255.255.255"],
        ["Deny-Sql-MinTLS",                 "Policy",     "Deny – SQL Server TLS < 1.2"],
        ["Deny-SQLDBs-TLS",                 "Policy",     "Deny – SQL DB TLS-Enforcement"],
        ["Deny-Sql-Managed-Instance-Public","Policy",     "Deny – SQL Managed Instance Public Endpoint"],
        ["Deny-Storage-CopyScope",          "Policy",     "Deny – Storage Copy Scope"],
        ["Deny-Storage-SFTP",               "Policy",     "Deny – SFTP auf Storage Accounts"],
        ["Deny-Blob-Soft-Delete",           "Policy",     "Audit – Blob Soft Delete aktiviert"],
        ["Audit-AppGW-WAF",                 "Policy",     "Audit – Application Gateway WAF Mode"],
    ],
    col_widths=[2.7, 0.8, 4.5]
)

doc.add_heading("2.3 alz-landingzones-corp – 5 Assignments", level=2)
add_table(doc,
    ["Assignment-Name", "Typ", "Effekt"],
    [
        ["Deny-Public-Endpoints",           "Initiative", "Deny – Alle PaaS-Dienste ohne Public Endpoint (59 Policies)"],
        ["Deny-Public-IP-On-NIC",           "Policy",     "Deny – Keine Public IP direkt an NIC"],
        ["Deny-HybridNetworking",           "Policy",     "Deny – Kein on-prem Netzwerk-Peering (nur über Hub)"],
        ["Deploy-Private-DNS-Zones",        "Initiative", "DeployIfNotExists – Private DNS für alle PaaS (corp-spezifisch)"],
        ["Deny-Subnet-Without-Udr",         "Policy",     "Deny – Subnetze ohne User-Defined Route Table (erzwingt Hub-Routing)"],
    ],
    col_widths=[2.5, 0.8, 4.7]
)

doc.add_heading("2.4 alz-landingzones-local – 1 Assignment", level=2)
add_table(doc,
    ["Assignment-Name", "Typ", "Effekt"],
    [
        ["Enforce-ALDO-Services", "Initiative", "Enforce – Azure Local (Sovereign) spezifische Dienste und Einschränkungen"],
    ],
    col_widths=[2.5, 0.8, 4.7]
)

doc.add_heading("2.5 alz-platform – 40 Assignments", level=2)
add_body(doc, "Gelten für alle Plattform-Subscriptions (connectivity, identity, management, security).", size=10)
add_table(doc,
    ["Assignment-Name", "Typ", "Effekt"],
    [
        ["Deploy-VM-Monitoring",            "Initiative", "DeployIfNotExists – VM Insights auf Plattform-VMs"],
        ["Deploy-VMSS-Monitoring",          "Initiative", "DeployIfNotExists – VMSS Monitoring"],
        ["Deploy-VM-Backup",                "Initiative", "DeployIfNotExists – VM Backup"],
        ["DenyAction-DeleteUAMIAMA",        "Policy",     "DenyAction – AMA Managed Identity schützen"],
        ["Enforce-Guardrails-APIM",         "Initiative", "DoNotEnforce – APIM"],
        ["Enforce-Guardrails-AppServices",  "Initiative", "DoNotEnforce – App Services"],
        ["Enforce-Guardrails-Automation",   "Initiative", "DoNotEnforce – Automation"],
        ["Enforce-Guardrails-CognitiveServices","Initiative","DoNotEnforce – Cognitive Services"],
        ["Enforce-Guardrails-ContainerApps","Initiative", "DoNotEnforce – Container Apps"],
        ["Enforce-Guardrails-ContainerInstance","Initiative","DoNotEnforce – Container Instances"],
        ["Enforce-Guardrails-ContainerRegistry","Initiative","DoNotEnforce – Container Registry"],
        ["Enforce-Guardrails-CosmosDb",     "Initiative", "DoNotEnforce – Cosmos DB"],
        ["Enforce-Guardrails-DataExplorer", "Initiative", "DoNotEnforce – Data Explorer"],
        ["Enforce-Guardrails-DataFactory",  "Initiative", "DoNotEnforce – Data Factory"],
        ["Enforce-Guardrails-EventGrid",    "Initiative", "DoNotEnforce – Event Grid"],
        ["Enforce-Guardrails-EventHub",     "Initiative", "DoNotEnforce – Event Hub"],
        ["Enforce-Guardrails-KeyVault",     "Initiative", "DoNotEnforce – Key Vault"],
        ["Enforce-Guardrails-Kubernetes",   "Initiative", "DoNotEnforce – AKS"],
        ["Enforce-Guardrails-MachineLearning","Initiative","DoNotEnforce – Machine Learning"],
        ["Enforce-Guardrails-MariaDb",      "Initiative", "DoNotEnforce – MariaDB"],
        ["Enforce-Guardrails-MySQL",        "Initiative", "DoNotEnforce – MySQL"],
        ["Enforce-Guardrails-Network",      "Initiative", "DoNotEnforce – Netzwerk"],
        ["Enforce-Guardrails-OpenAI",       "Initiative", "DoNotEnforce – OpenAI"],
        ["Enforce-Guardrails-PostgreSQL",   "Initiative", "DoNotEnforce – PostgreSQL"],
        ["Enforce-Guardrails-ServiceBus",   "Initiative", "DoNotEnforce – Service Bus"],
        ["Enforce-Guardrails-SQL",          "Initiative", "DoNotEnforce – SQL"],
        ["Enforce-Guardrails-Storage",      "Initiative", "DoNotEnforce – Storage"],
        ["Enforce-Guardrails-Synapse",      "Initiative", "DoNotEnforce – Synapse Analytics"],
        ["Enforce-Guardrails-VirtualDesktop","Initiative","DoNotEnforce – WVD/AVD"],
        ["Enforce-Guardrails-AADB2C",       "Initiative", "DoNotEnforce – AAD B2C"],
        ["Deploy-Private-DNS-Zones",        "Initiative", "DeployIfNotExists – 59 Private DNS Zonen"],
        ["Deny-MgmtPorts-From-Internet",    "Policy",     "Deny – RDP/SSH aus Internet"],
        ["Deny-Subnet-Without-Nsg",         "Policy",     "Deny – NSG-Pflicht"],
        ["Deny-Storage-http",               "Policy",     "Deny – Storage ohne HTTPS"],
        ["Deny-IP-Forwarding-On-NIC",       "Policy",     "Deny – IP-Forwarding"],
        ["Deny-Classic-Resources",          "Policy",     "Deny – Klassische ARM-Ressourcen"],
        ["Deny-APIM-Without-SSL",           "Policy",     "Deny – APIM ohne TLS"],
        ["Deny-SQL-FWRules-For-All-IPs",    "Policy",     "Deny – SQL-FW 0.0.0.0/0"],
        ["Deny-Sql-MinTLS",                 "Policy",     "Deny – SQL TLS < 1.2"],
        ["Deny-MachineLearning-PublicNetworkAccess","Policy","Deny – ML ohne Public Network"],
    ],
    col_widths=[2.7, 0.8, 4.5]
)

doc.add_heading("2.6 alz-platform-connectivity – 1 Assignment", level=2)
add_table(doc,
    ["Assignment-Name", "Typ", "Effekt"],
    [["Enable-DDoS-VNET", "Policy", "Audit/Modify – DDoS Network Protection auf allen VNets in connectivity aktivieren"]],
    col_widths=[2.5, 0.8, 4.7]
)

doc.add_heading("2.7 alz-platform-identity – 4 Assignments", level=2)
add_table(doc,
    ["Assignment-Name", "Typ", "Effekt"],
    [
        ["Deny-MgmtPorts-From-Internet", "Policy", "Deny – RDP/SSH aus Internet (besonders kritisch für DC)"],
        ["Deny-Public-IP",               "Policy", "Deny – Keine Public IPs in der Identity-Subscription"],
        ["Deny-Subnet-Without-Nsg",      "Policy", "Deny – NSG-Pflicht für alle Subnetze"],
        ["Deploy-VM-Backup",             "Initiative","DeployIfNotExists – VMs (Domain Controller) zwingend sichern"],
    ],
    col_widths=[2.5, 0.8, 4.7]
)

doc.add_heading("2.8 alz-sandbox – 1 Assignment", level=2)
add_table(doc,
    ["Assignment-Name", "Typ", "Effekt"],
    [["Enforce-ALZ-Sandbox", "Initiative", "Audit – Sandbox-spezifische Guardrails; kein Enforce; kein ExpressRoute/VPN erlaubt"]],
    col_widths=[2.5, 0.8, 4.7]
)

doc.add_heading("2.9 alz-decommissioned – 1 Assignment", level=2)
add_table(doc,
    ["Assignment-Name", "Typ", "Effekt"],
    [["Enforce-ALZ-Decomm", "Initiative", "Deny – Keine neue Ressourcen-Erstellung; vorhandene Ressourcen bleiben bestehen"]],
    col_widths=[2.5, 0.8, 4.7]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 3. POLICY DEFINITIONEN (149)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("3. Custom Policy Definitionen (149)", level=1)
add_body(doc,
    "Die 149 Custom Policy-Definitionen sind unter dem 'alz' Management Group Scope definiert "
    "und durch den ALZ Bicep Accelerator verwaltet. Sie werden in Initiativen gebündelt und "
    "per Assignment aktiviert. Kategorisierung nach Dienst/Funktion:"
)

doc.add_heading("3.1 Monitoring & Diagnose (67 Definitionen)", level=2)
add_body(doc,
    "Alle Deploy-Diagnostics-* Policies aktivieren automatisch Diagnose-Einstellungen und leiten "
    "Logs/Metriken in den zentralen Log Analytics Workspace (Effekt: DeployIfNotExists).", size=10
)
add_table(doc,
    ["Policy-Name", "Ressourcentyp"],
    [
        ["Deploy-AzActivity-Log",                    "Subscription Activity Log → Log Analytics"],
        ["Deploy-RG-ActivityLog",                    "Resource Group Activity Log"],
        ["Deploy-Diagnostics-AA",                    "Automation Account"],
        ["Deploy-Diagnostics-ACI",                   "Container Instances"],
        ["Deploy-Diagnostics-ACR",                   "Container Registry"],
        ["Deploy-Diagnostics-AnalysisService",       "Analysis Services"],
        ["Deploy-Diagnostics-ApiForFHIR",            "Azure API for FHIR"],
        ["Deploy-Diagnostics-APIMgmt",               "API Management"],
        ["Deploy-Diagnostics-ApplicationGW",         "Application Gateway"],
        ["Deploy-Diagnostics-AppServicePlan",        "App Service Plan"],
        ["Deploy-Diagnostics-AppServices",           "App Service (Web Apps, Functions)"],
        ["Deploy-Diagnostics-Bastion",               "Azure Bastion"],
        ["Deploy-Diagnostics-CDNEndpoints",          "CDN Endpoints"],
        ["Deploy-Diagnostics-CognitiveServices",     "Cognitive Services / AI Services"],
        ["Deploy-Diagnostics-CosmosDbs",             "Cosmos DB Accounts"],
        ["Deploy-Diagnostics-DatabricksWs",          "Databricks Workspaces"],
        ["Deploy-Diagnostics-DataFactory",           "Data Factory v2"],
        ["Deploy-Diagnostics-DLStore",               "Data Lake Store (Gen1)"],
        ["Deploy-Diagnostics-EventGridSub",          "Event Grid Subscriptions"],
        ["Deploy-Diagnostics-EventGridSystemTopic",  "Event Grid System Topics"],
        ["Deploy-Diagnostics-EventGridTopic",        "Event Grid Topics"],
        ["Deploy-Diagnostics-EventHub",              "Event Hub Namespaces"],
        ["Deploy-Diagnostics-ExpressRoute",          "ExpressRoute Circuits"],
        ["Deploy-Diagnostics-Firewall",              "Azure Firewall"],
        ["Deploy-Diagnostics-FrontDoor",             "Azure Front Door"],
        ["Deploy-Diagnostics-Function",              "Azure Functions"],
        ["Deploy-Diagnostics-HDInsight",             "HDInsight Clusters"],
        ["Deploy-Diagnostics-HealthcareApis",        "Healthcare APIs / FHIR Service"],
        ["Deploy-Diagnostics-KeyVault",              "Key Vault"],
        ["Deploy-Diagnostics-LoadBalancer",          "Load Balancer"],
        ["Deploy-Diagnostics-LogAnalytics",          "Log Analytics Workspaces"],
        ["Deploy-Diagnostics-LogicAppsISE",          "Logic Apps Integration Service Environment"],
        ["Deploy-Diagnostics-LogicAppsWF",           "Logic Apps Workflows"],
        ["Deploy-Diagnostics-MariaDB",               "MariaDB Server"],
        ["Deploy-Diagnostics-MediaService",          "Media Services"],
        ["Deploy-Diagnostics-MlWorkspace",           "Machine Learning Workspaces"],
        ["Deploy-Diagnostics-MySQL",                 "MySQL Server / Flexible Server"],
        ["Deploy-Diagnostics-NetworkSecurityGroups", "Network Security Groups (NSG)"],
        ["Deploy-Diagnostics-NIC",                   "Network Interfaces (NIC)"],
        ["Deploy-Diagnostics-NotificationHubs",      "Notification Hub Namespaces"],
        ["Deploy-Diagnostics-PostgreSQL",            "PostgreSQL Server / Flexible Server"],
        ["Deploy-Diagnostics-PowerBIEmbedded",       "Power BI Embedded"],
        ["Deploy-Diagnostics-PublicIP",              "Public IP Adressen"],
        ["Deploy-Diagnostics-RecoveryVault",         "Recovery Services Vault"],
        ["Deploy-Diagnostics-RedisCache",            "Azure Cache for Redis"],
        ["Deploy-Diagnostics-Relay",                 "Azure Relay"],
        ["Deploy-Diagnostics-SearchDomains",         "Azure Cognitive Search"],
        ["Deploy-Diagnostics-ServiceBus",            "Service Bus Namespaces"],
        ["Deploy-Diagnostics-SignalR",               "SignalR Service"],
        ["Deploy-Diagnostics-SQLDBs",                "SQL Databases"],
        ["Deploy-Diagnostics-SQLManagedInstances",   "SQL Managed Instances"],
        ["Deploy-Diagnostics-StorageAccounts",       "Storage Accounts (Blob, File, Queue, Table)"],
        ["Deploy-Diagnostics-StreamAnalytics",       "Stream Analytics Jobs"],
        ["Deploy-Diagnostics-TimeSeriesInsights",    "Time Series Insights"],
        ["Deploy-Diagnostics-TrafficManager",        "Traffic Manager Profiles"],
        ["Deploy-Diagnostics-VirtualNetwork",        "Virtual Networks (VNet)"],
        ["Deploy-Diagnostics-VM",                    "Virtual Machines"],
        ["Deploy-Diagnostics-VMSS",                  "Virtual Machine Scale Sets"],
        ["Deploy-Diagnostics-VNetGW",                "Virtual Network Gateways"],
        ["Deploy-Diagnostics-WVDAppGroup",           "Azure Virtual Desktop App Groups"],
        ["Deploy-Diagnostics-WVDHostPools",          "Azure Virtual Desktop Host Pools"],
        ["Deploy-Diagnostics-WVDWorkspace",          "Azure Virtual Desktop Workspaces"],
        ["Deploy-Diagnostics-WebServerFarm",         "App Service Plans"],
        ["Deploy-Diagnostics-Website",               "Websites / Static Web Apps"],
        ["Deploy-VM-Monitoring-DCR",                 "VM Insights Data Collection Rule"],
        ["Deploy-VMSS-Monitoring-DCR",               "VMSS Insights Data Collection Rule"],
        ["DenyAction-DeleteUAMIAMA",                 "Managed Identity für Azure Monitor Agent"],
    ],
    col_widths=[3.5, 4.5]
)

doc.add_heading("3.2 Netzwerk (20 Definitionen)", level=2)
add_table(doc,
    ["Policy-Name", "Effekt", "Beschreibung"],
    [
        ["Deny-AINE-Pip-OnNic",               "Deny",              "Public IP auf NIC in Azure-in-Network-Environments verbieten"],
        ["Deny-Classic-Resources",             "Deny",              "Klassische (ASM) Netzwerkressourcen verbieten"],
        ["Deny-ERPeering-Without-ER-Direct",   "Deny",              "ExpressRoute-Peering ohne ER-Direct verbieten"],
        ["Deny-HybridNetworking",              "Deny",              "Direktes On-Prem-Networking (kein VPN/ER über Hub)"],
        ["Deny-IP-Forwarding-On-NIC",          "Deny",              "IP-Forwarding auf NICs (NVA-Bypass-Schutz)"],
        ["Deny-MgmtPorts-From-Internet",       "Deny",              "RDP (3389) und SSH (22) aus Internet sperren"],
        ["Deny-NetworkInterfaces-PrivateIP",   "Deny",              "Statische Private IPs auf NICs sperren"],
        ["Deny-Private-DNS-Zones",             "Deny",              "Private DNS Zonen nur im connectivity-Scope erlaubt"],
        ["Deny-Public-AKS-APIServer-Access",   "Deny",              "AKS API-Server nicht öffentlich erreichbar"],
        ["Deny-Public-IP",                     "Deny",              "Keine Public IP Adressen in geschützten MGs"],
        ["Deny-Public-IP-On-NIC",              "Deny",              "Public IP direkt an NIC verboten"],
        ["Deny-Subnet-Without-Nsg",            "Deny",              "Subnetz muss eine NSG haben"],
        ["Deny-Subnet-Without-Udr",            "Deny",              "Subnetz muss UDR (Route Table) haben"],
        ["Deny-VNET-Peer-Cross-Sub",           "Deny",              "VNet-Peering nur innerhalb derselben Sub erlaubt"],
        ["Deploy-Custom-Route-Table",          "DeployIfNotExists", "Standard-Route-Table auf neue Subnetze anwenden"],
        ["Deploy-DDoSProtection",              "DeployIfNotExists", "DDoS Network Protection auf VNets in connectivity"],
        ["Deploy-Private-DNS-Zones",           "DeployIfNotExists", "Private DNS Zone-Links automatisch erstellen"],
        ["Enable-DDoS-VNET",                   "Audit/Modify",      "DDoS-Schutz auf VNets aktivieren"],
        ["Modify-NSG",                         "Modify",            "NSG Flow Logs aktivieren (Traffic Analytics)"],
        ["Deny-ERPeering-Without-ER-Direct",   "Deny",              "ER-Circuit darf nur via ER-Direct gepeert werden"],
    ],
    col_widths=[2.7, 1.4, 3.9]
)

doc.add_heading("3.3 Storage (16 Definitionen)", level=2)
add_table(doc,
    ["Policy-Name", "Effekt", "Beschreibung"],
    [
        ["Deny-Blob-Soft-Delete",              "Audit",  "Soft Delete für Blobs muss aktiviert sein"],
        ["Deny-Storage-CopyScope",             "Deny",   "Kopier-Scope auf eigenen Tenant einschränken"],
        ["Deny-Storage-http",                  "Deny",   "HTTP-Zugriff auf Storage Accounts verbieten (HTTPS only)"],
        ["Deny-Storage-MinTls",                "Deny",   "TLS-Version < 1.2 auf Storage verbieten"],
        ["Deny-Storage-NetworkAcls",           "Deny",   "Storage Accounts ohne Netzwerk-ACLs verbieten"],
        ["Deny-Storage-PublicAccess",          "Deny",   "Anonymer Public Blob Zugriff verbieten"],
        ["Deny-Storage-SFTP",                  "Deny",   "SFTP-Protokoll auf Storage Accounts verbieten"],
        ["Deny-Storage-SharedKey",             "Deny",   "Shared-Key-Authentifizierung verbieten (Entra-only)"],
        ["Deny-StorageAccount-CustomDomain",   "Deny",   "Custom-Domain auf Storage Accounts verbieten"],
        ["Deploy-Storage-SSE",                 "DeployIfNotExists", "Server-Side-Encryption auf Storage aktivieren"],
        ["Modify-StorageAccount-MinimumTls",   "Modify", "TLS-Mindestversion automatisch auf 1.2 setzen"],
        ["Deny-Storage-CorsRules",             "Deny",   "CORS-Regeln auf Storage Accounts einschränken"],
        ["Deploy-Diagnostics-StorageAccounts", "DeployIfNotExists", "Diagnose für Storage (Blob/File/Queue/Table)"],
        ["Audit-StorageAccount-SoftDelete",    "Audit",  "Soft Delete auf Speicherkonto prüfen"],
        ["Deny-StorageAccount-Infrastructure-Encryption", "Deny", "Doppelte Verschlüsselung (Infra-Ebene) erzwingen"],
        ["Deploy-StorageAccount-EncryptionScopes","DeployIfNotExists","Encryption Scopes auf Storage erstellen"],
    ],
    col_widths=[2.8, 1.4, 3.8]
)

doc.add_heading("3.4 SQL & Datenbanken (13 Definitionen)", level=2)
add_table(doc,
    ["Policy-Name", "Effekt", "Beschreibung"],
    [
        ["Deploy-MDFC-SqlAtp",                     "DeployIfNotExists","Defender for SQL (Advanced Threat Protection)"],
        ["Deploy-SQL-AuditingSettings",             "DeployIfNotExists","SQL Server Auditing → Log Analytics"],
        ["Deploy-SQL-SecurityAlertPolicies",        "DeployIfNotExists","Security Alert Policies auf SQL Server"],
        ["Deploy-SQL-Tde",                          "DeployIfNotExists","Transparent Data Encryption auf SQL DBs"],
        ["Deploy-SQL-VulnerabilityAssessments",     "DeployIfNotExists","SQL Vulnerability Assessment aktivieren"],
        ["Deploy-SqlMi-AuditingSettings",           "DeployIfNotExists","SQL Managed Instance Auditing"],
        ["Deploy-SqlMi-SecurityAlertPolicies",      "DeployIfNotExists","SQL MI Security Alert Policies"],
        ["Deploy-SqlMi-Tde",                        "DeployIfNotExists","TDE auf SQL Managed Instance"],
        ["Deploy-SqlMi-VulnerabilityAssessments",   "DeployIfNotExists","Vulnerability Assessment auf SQL MI"],
        ["Deny-Sql-Managed-Instance-Public",        "Deny",             "SQL MI Public Endpoint verbieten"],
        ["Deny-SQL-FWRules-For-All-IPs",            "Deny",             "SQL-FW-Regel 0.0.0.0-255.255.255.255 verbieten"],
        ["Deny-Sql-MinTLS",                         "Deny",             "SQL Server TLS < 1.2 verbieten"],
        ["Deny-SQLDBs-TLS",                         "Deny",             "SQL Databases TLS-Enforcement"],
    ],
    col_widths=[2.8, 1.4, 3.8]
)

doc.add_heading("3.5 Security & Defender (12 Definitionen)", level=2)
add_table(doc,
    ["Policy-Name", "Effekt", "Beschreibung"],
    [
        ["Deploy-MDFC-Config-H224",         "DeployIfNotExists","Defender for Cloud Config + Security Contacts"],
        ["Deploy-MDEndpoints",              "DeployIfNotExists","Defender for Endpoint Agent (MDE) auf VMs"],
        ["Deploy-MDEndpointsAMA",           "DeployIfNotExists","MDE + Azure Monitor Agent Integration"],
        ["Deploy-MDFC-OssDb",               "DeployIfNotExists","Defender for OSS Databases (PG/MySQL/MariaDB)"],
        ["Deploy-Defender-StorageAccounts", "DeployIfNotExists","Defender for Storage auf alle Storage Accounts"],
        ["Deploy-Defender-OssDb",           "DeployIfNotExists","Defender für Open-Source Datenbanken"],
        ["Deploy-ASC-Monitoring",           "AuditIfNotExists", "Azure Security Center Monitoring aktivieren"],
        ["Deny-AppGW-Without-WAF",          "Deny",             "Application Gateway ohne WAF Deny"],
        ["Audit-AppGW-WAF",                 "Audit",            "WAF-Mode (Prevention) auf AppGW prüfen"],
        ["Deny-APIM-Without-SSL",           "Deny",             "API Management ohne TLS/SSL verbieten"],
        ["Deny-MachineLearning-PublicNetworkAccess","Deny",     "Azure ML ohne Public Network"],
        ["Deploy-MDFC-SqlAtp",              "DeployIfNotExists","Defender für SQL (auch in Security-Kategorie)"],
    ],
    col_widths=[2.8, 1.4, 3.8]
)

doc.add_heading("3.6 Compute & Backup (11 Definitionen)", level=2)
add_table(doc,
    ["Policy-Name", "Effekt", "Beschreibung"],
    [
        ["Deploy-VM-Backup",                    "DeployIfNotExists","VMs ohne Backup-Tag automatisch sichern"],
        ["Deploy-VM-Monitoring",                "DeployIfNotExists","VM Insights + Azure Monitor Agent auf VMs"],
        ["Deploy-VMSS-Monitoring",              "DeployIfNotExists","VM Insights auf VM Scale Sets"],
        ["Deploy-AUMConfigure-VMLinuxHybrid",   "DeployIfNotExists","Azure Update Manager für Linux Hybrid (Arc)"],
        ["Deploy-AUMConfigure-VMWinHybrid",     "DeployIfNotExists","Azure Update Manager für Windows Hybrid (Arc)"],
        ["Deploy-AUMConfigure-VMLinuxAzure",    "DeployIfNotExists","Azure Update Manager für Linux Azure VMs"],
        ["Deploy-AUMConfigure-VMWinAzure",      "DeployIfNotExists","Azure Update Manager für Windows Azure VMs"],
        ["Deny-IP-Forwarding-On-NIC",           "Deny",             "IP-Forwarding auf VM-NICs"],
        ["Deploy-VM-Monitoring-DCR",            "DeployIfNotExists","Data Collection Rule für VM Insights"],
        ["Deploy-VMSS-Monitoring-DCR",          "DeployIfNotExists","Data Collection Rule für VMSS Insights"],
        ["DenyAction-DeleteUAMIAMA",            "DenyAction",       "Managed Identity für AMA schützen"],
    ],
    col_widths=[2.8, 1.4, 3.8]
)

doc.add_heading("3.7 Governance & Tagging (10 Definitionen)", level=2)
add_table(doc,
    ["Policy-Name", "Effekt", "Beschreibung"],
    [
        ["Audit-Tags-Mandatory-Rg",             "Audit",  "Pflicht-Tags (Environment, Owner, CostCenter) auf RGs prüfen"],
        ["Deny-Classic-Resources",              "Deny",   "Alle klassischen ASM-Ressourcen verbieten"],
        ["Deny-ERPeering-Without-ER-Direct",    "Deny",   "ER-Peering nur via ER-Direct"],
        ["Deny-VNET-Peer-Cross-Sub",            "Deny",   "VNet-Peering Subscription-übergreifend verbieten"],
        ["Deny-Private-DNS-Zones",              "Deny",   "Private DNS Zonen außerhalb connectivity verbieten"],
        ["Enforce-ACSB",                        "Audit",  "Azure Cloud Security Baseline (CIS/NIST Benchmark)"],
        ["Enforce-ALZ-Decomm",                  "Deny",   "Decommissioned MG: Keine neuen Ressourcen"],
        ["Enforce-ALZ-Sandbox",                 "Audit",  "Sandbox: Keine Prod-Ressourcen (ExpressRoute/VPN)"],
        ["Enforce-ALDO-Services",               "Enforce","Azure Local / Sovereign Workload Richtlinien"],
        ["Deny-StorageAccount-CustomDomain",    "Deny",   "Custom Domains auf Storage verbieten"],
    ],
    col_widths=[2.8, 1.0, 4.2]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 4. POLICY INITIATIVEN (42)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("4. Policy-Set-Definitionen / Initiativen (42)", level=1)
add_body(doc,
    "Initiativen (Policy Sets) bündeln mehrere Policy-Definitionen zu einem logischen Deployment-Block. "
    "Sie werden als einzelnes Assignment zugewiesen und steuern den Enforcement-Status gemeinsam."
)

add_table(doc,
    ["Initiativenname", "Enthält (Policies)", "Scope / MG"],
    [
        ["Deploy-MDFC-Config",             "~8 Policies (Defender Config, Security Contacts, Auto-Provisioning)", "alz"],
        ["Deploy-Private-DNS-Zones",       "59 Private DNS Zone-Policies (je PaaS-Dienst eine Zone)",            "alz-landingzones, alz-platform"],
        ["Deploy-Diagnostics-LogAnalytics","67 Diagnose-Policies (alle Ressourcentypen → LAW)",                   "alz"],
        ["Deploy-VM-Monitoring",           "4 Policies (AMA, VM Insights, DCR, Linux/Windows)",                   "alz, alz-platform, alz-landingzones"],
        ["Deploy-VMSS-Monitoring",         "3 Policies (VMSS AMA, VMSS Insights, DCR)",                           "alz, alz-platform, alz-landingzones"],
        ["Deploy-VM-Backup",               "2 Policies (Windows VM Backup, Linux VM Backup → Recovery Vault)",    "alz, alz-platform, alz-landingzones"],
        ["Deploy-MDEndpoints",             "3 Policies (MDE Windows, MDE Linux, MDE ARC)",                        "alz"],
        ["Deploy-MDEndpointsAMA",          "3 Policies (MDE+AMA Windows, Linux, ARC)",                            "alz"],
        ["Deploy-MDFC-OssDb",              "3 Policies (Defender PostgreSQL, MySQL, MariaDB)",                     "alz"],
        ["Deploy-MDFC-SqlAtp",             "2 Policies (SQL Server ATP, SQL MI ATP)",                              "alz, alz-landingzones"],
        ["Deny-Public-Endpoints",          "59 Policies (Deny Public Endpoint je PaaS-Dienst)",                   "alz-landingzones-corp"],
        ["Enforce-ACSB",                   "~15 Policies (CIS/NIST Baseline Audit-Policies)",                     "alz"],
        ["Enforce-ALZ-Decomm",             "2 Policies (Deny Compute, Deny Resource Creation)",                   "alz-decommissioned"],
        ["Enforce-ALZ-Sandbox",            "2 Policies (Deny ER/VPN, Deny Prod-Ressourcen)",                      "alz-sandbox"],
        ["Enforce-ALDO-Services",          "~5 Policies (Azure Local / Sovereign Services)",                      "alz-landingzones-local"],
        ["Enforce-Guardrails-APIM",        "~6 Policies (TLS, Private, WAF, SKU-Check)",                          "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-AppServices", "~8 Policies (TLS, HTTPS, Auth, Remote Debug, Slots)",                 "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-Automation",  "~4 Policies (Public Access, Managed Identity, TLS)",                  "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-CognitiveServices","~6 Policies (Public Network, CMK, SKU, Outbound)",               "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-ContainerApps","~5 Policies (Auth, Managed Identity, Ingress)",                      "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-ContainerInstance","~3 Policies (Public IP, Managed Identity)",                      "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-ContainerRegistry","~6 Policies (Public Access, Admin, Geo-Replication, CMK)",       "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-CosmosDb",    "~5 Policies (Public Access, CMK, CORS, TLS)",                         "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-DataExplorer","~4 Policies (Public Access, SKU, Disk Encryption)",                   "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-DataFactory", "~5 Policies (Public Access, Managed Identity, CMK, Git)",             "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-EventGrid",   "~4 Policies (Public Access, Managed Identity, SKU)",                  "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-EventHub",    "~5 Policies (CMK, TLS, Public Access, Double Encryption)",            "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-KeyVault",    "~6 Policies (Soft Delete, Purge Protection, Firewall, RBAC)",         "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-Kubernetes",  "~8 Policies (API Server, Ingress, Privileged Pods, Audit)",           "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-MachineLearning","~5 Policies (Public Network, CMK, Managed Identity)",              "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-MariaDb",     "~4 Policies (SSL, Firewall, Geo-Backup, TLS)",                        "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-MySQL",       "~5 Policies (SSL, Public Access, CMK, Audit, Backup)",                "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-Network",     "~6 Policies (NSG Flow Logs, Bastion, DDoS, FW Policy)",               "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-OpenAI",      "~4 Policies (Public Access, Managed Identity, CMK)",                  "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-PostgreSQL",  "~5 Policies (SSL, Public Access, CMK, Geo-Backup)",                   "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-ServiceBus",  "~5 Policies (Public Access, CMK, TLS, Double Encryption)",            "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-SQL",         "~8 Policies (TDE, ATP, Auditing, AAD-Admin, FW, TLS)",               "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-Storage",     "~8 Policies (HTTPS, TLS, PublicAccess, SharedKey, CMK, SoftDelete)",  "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-Synapse",     "~5 Policies (CMK, Managed Identity, Public Access, Data Exfiltration)","alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-VirtualDesktop","~4 Policies (Managed Identity, Session Type, Disk Encryption)",     "alz-landingzones, alz-platform"],
        ["Enforce-Guardrails-AADB2C",      "~3 Policies (Tenant Isolation, Logging, GDPR)",                       "alz-landingzones, alz-platform"],
        ["DenyAction-DeleteUAMIAMA",       "1 Policy (DenyAction Delete auf UAMI für Azure Monitor Agent)",        "alz, alz-platform"],
    ],
    col_widths=[2.5, 3.3, 2.2]
)

add_section_note(doc,
    "Private DNS Zones Initiative: Die 59 Private DNS Zonen decken alle gängigen PaaS-Dienste ab, u.a.: "
    "Azure SQL, Storage (Blob/File/Queue/Table/Web/DFS), Key Vault, App Services, Container Registry, "
    "Cosmos DB, Service Bus, Event Hub, Azure ML, Cognitive Services, IoT Hub, Backup, Kubernetes, "
    "Databricks, API Management, Synapse, Azure Monitor, Batch, HDInsight, Redis Cache, u.v.m."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 5. CUSTOM RBAC ROLLEN
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("5. Custom RBAC-Rollen (5)", level=1)
add_body(doc,
    "Der ALZ Bicep Accelerator erstellt 5 benutzerdefinierte RBAC-Rollen als Ergänzung zu den "
    "Built-in Azure-Rollen. Sie werden auf Management-Group-Ebene definiert und über "
    "Entra-ID-Gruppen-Object-IDs zugewiesen."
)

add_table(doc,
    ["Rollenname", "Scope", "Berechtigungen (Zusammenfassung)"],
    [
        ["ALZ Subscription-Owner",
         "alz-landingzones / alz-platform",
         "Vollzugriff auf Subscriptions und Ressourcen; keine Policy/RBAC-Verwaltung auf MG-Ebene; "
         "entspricht Owner ohne Tenant-Rechte. Für Subscription-Verantwortliche."],
        ["ALZ Security-Operations",
         "alz (Tenant-weit)",
         "Lesen aller Security-Ressourcen; Schreiben in Defender for Cloud, Security Center, Sentinel; "
         "kein Zugriff auf Produktionsdaten. Für das Security Operations Center (SOC)."],
        ["ALZ Network-Management",
         "alz-platform-connectivity",
         "Vollzugriff auf Netzwerkressourcen (VNet, Firewall, NSG, Route Tables, DNS, Gateway); "
         "kein Zugriff auf Compute/Storage. Für das Netzwerkteam."],
        ["ALZ Application-Owners",
         "alz-landingzones",
         "Contributor-Rechte auf zugewiesene Ressourcengruppen; kein Netzwerk-/Policy-Zugriff; "
         "darf RBAC-Zuweisungen innerhalb der eigenen RGs verwalten. Für Workload-Teams."],
        ["ALZ Network-Subnet-Contributor",
         "alz-landingzones",
         "Nur Verwaltung von Subnetzen in bestehenden VNets (kein VNet-Create/Delete); "
         "darf NSGs und Route Tables auf Subnetzen setzen. Für Netzwerk-Delegierung."],
    ],
    col_widths=[1.9, 1.7, 4.4]
)

add_table(doc,
    ["RBAC-Deployment", "Management Group", "Bicep-Modul"],
    [
        ["RBAC Subscription-Owner",       "alz",                       "templates/core/governance/rbac/main.bicep"],
        ["RBAC Security-Operations",      "alz",                       "templates/core/governance/rbac/modules/roleAssignment-mg.bicep"],
        ["RBAC Network-Management",       "alz-platform-connectivity", "templates/core/governance/rbac/modules/roleAssignment-mg.bicep"],
        ["RBAC Application-Owners",       "alz-landingzones",          "templates/core/governance/rbac/modules/roleAssignment-mg.bicep"],
        ["RBAC Network-Subnet-Contributor","alz-landingzones",         "templates/core/governance/rbac/modules/roleAssignment-mg.bicep"],
    ],
    col_widths=[2.2, 2.0, 3.8]
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 6. BICEP MODULE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("6. Bicep-Module und Deploy-Stufen (20)", level=1)
add_body(doc,
    "Das ALZ Setup verwendet ausschließlich Azure Verified Modules (AVM) und offizielle ALZ-Pattern-Module. "
    "Alle Templates liegen unter /templates/ und werden in 18 geordneten Stufen deployed. "
    "Die Pipeline (GitHub Actions: deploy-alz.yml) validiert alle Templates mit Bicep 0.44.1 "
    "und dem Azure What-If-Deployment vor jedem Apply."
)

doc.add_heading("6.1 Governance-Module", level=2)
add_table(doc,
    ["Deploy-Stufe", "Template-Pfad", "Erstellt", "AVM-Modul"],
    [
        ["Stufe 1",  "templates/core/governance/mgmt-groups/int-root/main.bicep",
         "alz MG + alle 11 Kind-MGs",                                           "avm/ptn/alz/empty:0.3.6"],
        ["Stufe 2",  "templates/core/governance/mgmt-groups/landingzones/main.bicep",
         "alz-landingzones Sub-MGs (corp, online, local)",                       "Az. Management Module"],
        ["Stufe 3-9","templates/core/governance/mgmt-groups/int-root/main.bicep (Policy-Blöcke)",
         "149 Policy-Definitionen, 42 Initiativen",                              "avm/ptn/alz/empty:0.3.6"],
        ["Stufe 10-17","(Policy Assignment Templates)",
         "123 Policy Assignments auf 9 MG-Ebenen",                              "avm/ptn/alz/empty:0.3.6"],
        ["Stufe 18", "templates/core/governance/rbac/main.bicep",
         "5 Custom RBAC-Rollen + RBAC-Assignments",                             "avm/res/authorization/role-definition"],
    ],
    col_widths=[1.1, 2.8, 2.1, 2.0]
)

doc.add_heading("6.2 Logging & Monitoring Module", level=2)
add_body(doc,
    "Bechtle-Empfehlung: 1 Region (GWC). Alle Logging-Ressourcen werden einmalig "
    "in der Primärregion Germany West Central deployt.", size=10)
add_table(doc,
    ["Ressource", "Name-Konvention", "AVM-Modul", "Region"],
    [
        ["Log Analytics Workspace",    "law-alz-germanywestcentral",    "avm/res/operational-insights/workspace:0.9.0",    "GWC"],
        ["Managed Identity (AMA)",     "mi-alz-germanywestcentral",     "avm/res/managed-identity/user-assigned-identity", "GWC"],
        ["DCR – VM Insights",          "dcr-vmi-alz-germanywestcentral","avm/ptn/alz/ama:0.2.0",                          "GWC"],
        ["DCR – Change Tracking",      "dcr-ct-alz-germanywestcentral", "avm/ptn/alz/ama:0.2.0",                          "GWC"],
        ["DCR – Defender SQL",         "dcr-mdfcsql-alz-germanywestcentral","avm/ptn/alz/ama:0.2.0",                      "GWC"],
        ["Data Collection Endpoint",   "dce-alz-germanywestcentral",    "avm/res/insights/data-collection-endpoint",       "GWC"],
        ["Automation Account",         "aa-alz-germanywestcentral",     "avm/res/automation/automation-account",           "GWC (optional)"],
    ],
    col_widths=[2.2, 2.6, 2.3, 0.9]
)

doc.add_heading("6.3 Netzwerk-Module (Hub – Primärregion GWC)", level=2)
add_body(doc,
    "Bechtle-Empfehlung (Variante A/B, ~€1.050/Monat): 1 Hub in Germany West Central, "
    "Firewall Standard. Kein zweiter Hub – Geo-Redundanz nur bei Variante D (Microsoft Default, ~€5.800/Monat).", size=10)
add_table(doc,
    ["Ressource", "Name-Konvention", "AVM-Modul", "Hub GWC", "Status"],
    [
        ["Virtual Network (Hub)",  "vnet-alz-germanywestcentral",     "avm/res/network/virtual-network",          "10.0.0.0/22",  "Aktiv"],
        ["Azure Firewall Standard","afw-alz-germanywestcentral",      "avm/res/network/azure-firewall",           "10.0.0.0/26",  "Aktiv"],
        ["Firewall Policy",        "afwp-alz-germanywestcentral",     "avm/res/network/firewall-policy",          "Standard",     "Aktiv"],
        ["Azure Bastion Standard", "bas-alz-germanywestcentral",      "avm/res/network/bastion-host",             "10.0.0.64/26", "Aktiv"],
        ["VPN Gateway (VpnGw1AZ)", "vgw-alz-germanywestcentral",     "avm/res/network/virtual-network-gateway",  "10.0.0.128/27","Aktiv"],
        ["DNS Private Resolver",   "dnspr-alz-germanywestcentral",    "avm/res/network/dns-resolver",             "10.0.0.160/28","Aktiv"],
        ["Public IP (Firewall)",   "pip-afw-alz-germanywestcentral",  "avm/res/network/public-ip-address",        "Zone-redundant","Aktiv"],
        ["Public IP (Bastion)",    "pip-bas-alz-germanywestcentral",  "avm/res/network/public-ip-address",        "Zone-redundant","Aktiv"],
        ["Public IP (VPN GW)",     "pip-vpngw-alz-germanywestcentral","avm/res/network/public-ip-address",        "Zone-redundant","Aktiv"],
        ["ExpressRoute Gateway",   "ergw-alz-germanywestcentral",     "avm/res/network/express-route-gateway",    "10.0.0.128/27","Zurückgestellt"],
        ["DDoS Protection Plan",   "ddos-alz",                        "avm/res/network/ddos-protection-plan",     "–",            "Zurückgestellt"],
        ["Firewall Premium",       "(Upgrade, kein Rebuild)",         "avm/res/network/azure-firewall",           "In-Place",     "Zurückgestellt"],
        ["Sekundärer Hub (opt.)",  "vnet-alz-<region>",              "avm/res/network/virtual-network",          "10.1.0.0/22",  "Nur Variante D"],
    ],
    col_widths=[1.8, 2.3, 2.0, 1.1, 1.0]
)

doc.add_heading("6.4 Subscription Vending Module", level=2)
add_body(doc,
    "Neue Workload-Subscriptions werden über das Subscription Vending Template provisioniert, "
    "das automatisch die richtige MG-Zuweisung, Spoke-Netzwerk und RBAC konfiguriert.", size=10
)
add_table(doc,
    ["Template", "Erstellt", "Bicep-Modul"],
    [
        ["templates/core/subscription-vending/main.bicep",
         "Subscription → korrekte MG (corp/online/local); Spoke VNet; VNet-Peering zum Hub; Route Table; NSG; RBAC",
         "avm/ptn/lz/sub-vending:0.5.1"],
        ["templates/networking/spoke/main.bicep",
         "Spoke VNet, Subnetze, NSG, UDR, VNet-Peering, Private DNS Zone Links",
         "avm/res/network/virtual-network"],
        ["templates/networking/spoke/modules/route-table.bicep",
         "Route Table mit Default-Route → Azure Firewall (0.0.0.0/0 → Firewall-IP)",
         "Az. Native"],
        ["templates/networking/spoke/modules/private-dns-zone-link.bicep",
         "Links der 59 Private DNS Zonen aus connectivity in den Spoke",
         "avm/res/network/private-dns-zone"],
    ],
    col_widths=[3.0, 3.2, 1.8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 7. RESSOURCE-NAMING-KONVENTION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("7. Ressource-Naming-Konvention & Ressourcen-Inventar", level=1)
add_body(doc,
    "Alle ALZ-Ressourcen folgen dem Microsoft Cloud Adoption Framework Naming Standard. "
    "Grundmuster: <Dienst-Kürzel>-alz-<Umgebung>-<Zweck>-<Region-Kürzel>. "
    "Region-Kürzel: gwe = Germany West Central, ne = North Europe."
)

doc.add_heading("7.1 Umgebungs-Kürzel", level=2)
add_table(doc,
    ["Kürzel", "Umgebung", "Verwendung", "Beispiel"],
    [
        ["(kein Kürzel)", "Produktion",        "Produktive Landing Zone",               "law-alz-gwe"],
        ["poc",           "Proof of Concept",  "Demo- und Smoke-Run-Ressourcen im Tenant","law-alz-poc-gwe"],
        ["dev",           "Development",       "Entwicklungsumgebung (zukünftig)",       "law-alz-dev-gwe"],
        ["tst",           "Test/Staging",      "Testumgebung (zukünftig)",               "law-alz-tst-gwe"],
    ],
    col_widths=[1.1, 1.5, 2.6, 2.8]
)
add_section_note(doc,
    "PoC-Ressourcen sind zusätzlich mit dem Tag 'Environment: PoC' gekennzeichnet und "
    "damit im Azure Portal per Filter sofort von Produktionsressourcen (Environment: Production) unterscheidbar."
)

doc.add_heading("7.2 Naming-Muster nach Ressourcentyp", level=2)
add_table(doc,
    ["Ressourcentyp", "Muster (Produktion)", "Muster (PoC)"],
    [
        ["Management Group",          "alz[-sub-ebene]",                     "(identisch – umgebungsunabhängig)"],
        ["Resource Group",            "rg-alz-<zweck>-<region>",             "rg-alz-poc-<zweck>-<region>"],
        ["Log Analytics Workspace",   "law-alz-<region>",                    "law-alz-poc-<region>"],
        ["Managed Identity",          "mi-alz-<region>",                     "mi-alz-poc-<region>"],
        ["Data Collection Rule",      "dcr-<typ>-alz-<region>",              "dcr-<typ>-alz-poc-<region>"],
        ["Data Collection Endpoint",  "dce-alz-<region>",                    "dce-alz-poc-<region>"],
        ["Automation Account",        "aa-alz-<region>",                     "aa-alz-poc-<region>"],
        ["Virtual Network (Hub)",     "vnet-alz-hub-<region>",               "vnet-alz-poc-<region>"],
        ["Azure Firewall",            "afw-alz-<region>",                    "(nicht deployed in PoC)"],
        ["Firewall Policy",           "afwp-alz-<region>",                   "(nicht deployed in PoC)"],
        ["Azure Bastion",             "bas-alz-<region>",                    "(nicht deployed in PoC)"],
        ["VPN Gateway",               "vpngw-alz-<region>",                  "(nicht deployed in PoC)"],
        ["ExpressRoute Gateway",      "ergw-alz-<region>",                   "(nicht deployed in PoC)"],
        ["DNS Private Resolver",      "dnspr-alz-<region>",                  "dnspr-alz-poc-<region>"],
        ["Public IP",                 "pip-<dienst>-alz-<region>",           "pip-<dienst>-alz-poc-<region>"],
        ["DDoS Protection Plan",      "ddos-alz",                            "(nicht deployed in PoC)"],
        ["Recovery Services Vault",   "rsv-alz-<region>",                    "rsv-alz-poc-<region>"],
        ["Key Vault",                 "kv-<workload>-<env>-<region>",        "kv-<workload>-poc-<region>"],
        ["Storage Account",           "st<workload>prod<region>",            "st<workload>poc<region>"],
        ["NSG",                       "nsg-<subnetz>-<workload>-<region>",   "nsg-<subnetz>-<workload>-poc-<region>"],
        ["Route Table",               "rt-<zweck>-<workload>-<region>",      "rt-<zweck>-<workload>-poc-<region>"],
        ["Virtual Network (Spoke)",   "vnet-<workload>-prod-<region>",       "vnet-<workload>-poc-<region>"],
    ],
    col_widths=[2.0, 2.5, 3.5]
)

doc.add_heading("7.3 Konkretes Ressourcen-Inventar – Produktion", level=2)
add_body(doc,
    "Vollständige Liste aller tatsächlich deployten Ressourcen mit exakten Namen "
    "für die produktive Landing Zone (Germany West Central als primäre, North Europe als sekundäre Region).",
    size=10
)
add_table(doc,
    ["Ressource (exakter Name)", "Typ", "Region", "Resource Group / Scope"],
    [
        # Management Groups
        ["alz",                               "Management Group", "–",   "Tenant Root"],
        ["alz-platform",                      "Management Group", "–",   "alz"],
        ["alz-platform-connectivity",         "Management Group", "–",   "alz-platform"],
        ["alz-platform-identity",             "Management Group", "–",   "alz-platform"],
        ["alz-platform-management",           "Management Group", "–",   "alz-platform"],
        ["alz-platform-security",             "Management Group", "–",   "alz-platform"],
        ["alz-landingzones",                  "Management Group", "–",   "alz"],
        ["alz-landingzones-corp",             "Management Group", "–",   "alz-landingzones"],
        ["alz-landingzones-online",           "Management Group", "–",   "alz-landingzones"],
        ["alz-landingzones-local",            "Management Group", "–",   "alz-landingzones"],
        ["alz-sandbox",                       "Management Group", "–",   "alz"],
        ["alz-decommissioned",                "Management Group", "–",   "alz"],
        # Resource Groups
        ["rg-alz-logging-germanywestcentral", "Resource Group",   "GWC", "Management Sub"],
        ["rg-alz-logging-northeurope",        "Resource Group",   "NE",  "Management Sub"],
        ["rg-alz-conn-germanywestcentral",    "Resource Group",   "GWC", "Connectivity Sub"],
        ["rg-alz-conn-northeurope",           "Resource Group",   "NE",  "Connectivity Sub"],
        ["rg-alz-dns-germanywestcentral",     "Resource Group",   "GWC", "Connectivity Sub"],
        ["rg-alz-dns-northeurope",            "Resource Group",   "NE",  "Connectivity Sub"],
        # Logging
        ["law-alz-germanywestcentral",        "Log Analytics Workspace", "GWC", "rg-alz-logging-germanywestcentral"],
        ["law-alz-northeurope",               "Log Analytics Workspace", "NE",  "rg-alz-logging-northeurope"],
        ["mi-alz-germanywestcentral",         "Managed Identity",        "GWC", "rg-alz-logging-germanywestcentral"],
        ["mi-alz-northeurope",                "Managed Identity",        "NE",  "rg-alz-logging-northeurope"],
        ["dcr-vmi-alz-germanywestcentral",    "Data Collection Rule",    "GWC", "rg-alz-logging-germanywestcentral"],
        ["dcr-ct-alz-germanywestcentral",     "Data Collection Rule",    "GWC", "rg-alz-logging-germanywestcentral"],
        ["dcr-mdfcsql-alz-germanywestcentral","Data Collection Rule",    "GWC", "rg-alz-logging-germanywestcentral"],
        # Hub Networking GWC
        ["vnet-alz-germanywestcentral",       "Virtual Network (Hub)",   "GWC", "rg-alz-conn-germanywestcentral"],
        ["afw-alz-germanywestcentral",        "Azure Firewall",          "GWC", "rg-alz-conn-germanywestcentral"],
        ["afwp-alz-germanywestcentral",       "Firewall Policy",         "GWC", "rg-alz-conn-germanywestcentral"],
        ["bas-alz-germanywestcentral",        "Azure Bastion",           "GWC", "rg-alz-conn-germanywestcentral"],
        ["vpngw-alz-germanywestcentral",      "VPN Gateway",             "GWC", "rg-alz-conn-germanywestcentral"],
        ["dnspr-alz-germanywestcentral",      "DNS Private Resolver",    "GWC", "rg-alz-conn-germanywestcentral"],
        ["pip-afw-alz-germanywestcentral",    "Public IP (Firewall)",    "GWC", "rg-alz-conn-germanywestcentral"],
        ["pip-afw-mgmt-alz-germanywestcentral","Public IP (FW Mgmt)",   "GWC", "rg-alz-conn-germanywestcentral"],
        ["pip-bas-alz-germanywestcentral",    "Public IP (Bastion)",     "GWC", "rg-alz-conn-germanywestcentral"],
        ["pip-vpngw-alz-germanywestcentral",  "Public IP (VPN GW)",      "GWC", "rg-alz-conn-germanywestcentral"],
        # Hub Networking NE
        ["vnet-alz-northeurope",              "Virtual Network (Hub)",   "NE",  "rg-alz-conn-northeurope"],
        ["afw-alz-northeurope",               "Azure Firewall",          "NE",  "rg-alz-conn-northeurope"],
        ["afwp-alz-northeurope",              "Firewall Policy",         "NE",  "rg-alz-conn-northeurope"],
        ["bas-alz-northeurope",               "Azure Bastion",           "NE",  "rg-alz-conn-northeurope"],
        # DNS
        ["privatelink.blob.core.windows.net (+ 58 weitere Zonen)", "Private DNS Zones (59)", "GWC", "rg-alz-dns-germanywestcentral"],
    ],
    col_widths=[3.0, 1.8, 0.6, 2.6]
)

doc.add_heading("7.4 Konkretes Ressourcen-Inventar – PoC / Demo", level=2)
add_body(doc,
    "Vollständige Liste aller tatsächlich deployten Ressourcen im PoC-Setup "
    "(smokerun/ – Single Subscription, nur Germany West Central, ohne Firewall/Bastion/Gateway).",
    size=10
)
add_table(doc,
    ["Ressource (exakter Name)", "Typ", "Resource Group"],
    [
        # Management Groups – identisch
        ["alz (+ 11 Kind-MGs)",               "Management Groups (12)",           "Tenant Root – identisch zu Prod"],
        # Resource Groups
        ["rg-alz-poc-logging-gwe",            "Resource Group",                   "PoC Subscription"],
        ["rg-alz-poc-conn-germanywestcentral", "Resource Group",                   "PoC Subscription"],
        ["rg-alz-poc-dns-germanywestcentral",  "Resource Group",                   "PoC Subscription"],
        # Logging
        ["law-alz-poc-gwe",                   "Log Analytics Workspace (30 Tage)", "rg-alz-poc-logging-gwe"],
        ["mi-alz-poc-gwe",                    "Managed Identity",                  "rg-alz-poc-logging-gwe"],
        ["dcr-vmi-alz-poc-gwe",               "Data Collection Rule (VM Insights)","rg-alz-poc-logging-gwe"],
        ["dcr-ct-alz-poc-gwe",                "Data Collection Rule (Change Tracking)","rg-alz-poc-logging-gwe"],
        ["dcr-mdfcsql-alz-poc-gwe",           "Data Collection Rule (Defender SQL)","rg-alz-poc-logging-gwe"],
        # Networking
        ["vnet-alz-poc-gwe",                  "Virtual Network (Hub, 10.0.0.0/24)","rg-alz-poc-conn-germanywestcentral"],
        ["privatelink.blob.core.windows.net (+ weitere)", "Private DNS Zones",    "rg-alz-poc-dns-germanywestcentral"],
        # Nicht deployed in PoC
        ["afw-alz-poc-gwe",                   "Azure Firewall → NICHT deployed",  "–"],
        ["bas-alz-poc-gwe",                   "Azure Bastion  → NICHT deployed",  "–"],
        ["vpngw-alz-poc-gwe",                 "VPN Gateway    → NICHT deployed",  "–"],
    ],
    col_widths=[3.2, 2.5, 2.3]
)
add_section_note(doc,
    "Alle PoC-Ressourcen tragen das Tag 'Environment: PoC'. Teardown per: bash smokerun/teardown.sh"
)

# ═══════════════════════════════════════════════════════════════════
# 8. QUELLEN
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("8. Quellen und Referenzen", level=1)
add_table(doc,
    ["Quelle", "URL / Pfad"],
    [
        ["Microsoft ALZ Bicep Accelerator (GitHub)",     "github.com/Azure/ALZ-Bicep"],
        ["Azure Verified Modules Registry",              "aka.ms/avm"],
        ["ALZ Pattern Module (avm/ptn/alz/empty)",       "github.com/Azure/bicep-registry-modules"],
        ["Microsoft Cloud Adoption Framework (CAF)",     "learn.microsoft.com/azure/cloud-adoption-framework"],
        ["ALZ Architektur-Diagramme",                    "learn.microsoft.com/azure/cloud-adoption-framework/ready/landing-zone"],
        ["Azure Policy GitHub Repository",               "github.com/Azure/azure-policy"],
        ["Private DNS Zone Namensliste",                 "learn.microsoft.com/azure/private-link/private-endpoint-dns"],
        ["Lokales Konzept-Dokument",                     "docs/konzept/Word/Azure-Landing-Zone-Konzept.docx"],
        ["Lokale Bicep-Templates",                       "templates/ (dieses Repository)"],
    ],
    col_widths=[3.5, 4.5]
)

add_section_note(doc,
    f"Erstellt: {DATE} | Bechtle AG | Auf Basis: ALZ Bicep Accelerator avm/ptn/alz/empty:0.3.6 | "
    "Dieses Dokument wird durch generate-alz-referenz.py generiert und kann bei ALZ-Versionsupdates aktualisiert werden."
)

os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
