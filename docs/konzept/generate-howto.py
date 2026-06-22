# -*- coding: utf-8 -*-
"""
Generiert die Bechtle-Kundenanleitung "Azure Landing Zone – Nutzungsanleitung" als Word-Dokument.

Output : Azure-Landing-Zone-Nutzungsanleitung.docx

Nutzung:
    python3 docs/konzept/generate-howto.py

Abhängigkeit : python-docx  (pip install python-docx)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(BASE_DIR, "bechtle-brand", "Word", "VORLAGE_Bechtle_Management_Summary.docx")
OUTPUT   = os.path.join(BASE_DIR, "Word", "Azure-Landing-Zone-Nutzungsanleitung.docx")

KUNDE = "<KUNDE>"
DATE  = "22.06.2026"


# ─────────────────────────────────────────────────────────────────────────────
# Hilfsfunktionen (identisch mit generate-konzept.py)
# ─────────────────────────────────────────────────────────────────────────────

def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_body(doc, text: str, size: int = None):
    p = doc.add_paragraph(text, style="Normal")
    p.paragraph_format.space_after = Pt(6)
    if size:
        for run in p.runs:
            run.font.size = Pt(size)
    return p


def add_bullet(doc, text: str, bold_prefix: str = None):
    p = doc.add_paragraph(style="Bechtle Aufzählung")
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.bold = True
    p.add_run(text)
    return p


def add_numbered(doc, text: str, bold_prefix: str = None):
    p = doc.add_paragraph(style="Bechtle Aufzählung")
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.bold = True
    p.add_run(text)
    return p


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    try:
        tbl.style = "Bechtle Tabelle"
    except Exception:
        tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, "053B25")
    for row_data in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return tbl


def add_hint(doc, text: str):
    p = doc.add_paragraph(style="Standard klein")
    p.add_run("Hinweis: ").bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_code(doc, text: str):
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    return p


def update_header(doc, text: str):
    for para in doc.sections[0].header.paragraphs:
        if para.text.strip():
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = text
            else:
                para.add_run(text)
            break


def update_footer(doc, text: str):
    for para in doc.sections[0].footer.paragraphs:
        if para.text.strip():
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = text
            else:
                para.add_run(text)
            break


# ─────────────────────────────────────────────────────────────────────────────
# Dokument
# ─────────────────────────────────────────────────────────────────────────────

def build():
    doc = Document(TEMPLATE)
    clear_body(doc)
    update_header(doc, f"{KUNDE} | Azure Landing Zone – Nutzungsanleitung")
    update_footer(doc, "Bechtle | Azure Landing Zone – Nutzungsanleitung")

    # ═══════════════════════════════════════════════════════════════════
    # DECKBLATT
    # ═══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph(style="Standard klein")
    p.add_run("Bechtle GmbH & Co. KG · Gottlieb-Daimler-Straße 2 · 68165 Mannheim")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    tp = doc.add_paragraph(style="Title")
    tp.add_run("Azure Landing Zone – Nutzungsanleitung")

    sp = doc.add_paragraph(style="Subtitle")
    sp.add_run(
        "Wie Sie Ihre Azure-Infrastruktur nutzen, erweitern und betreiben – "
        "praxisnah und Schritt für Schritt erklärt"
    )
    doc.add_paragraph()
    doc.add_paragraph()

    add_table(doc,
        ["Dokument", "Version", "Stand", "Erstellt durch"],
        [["Azure Landing Zone – Nutzungsanleitung", "1.0", DATE, "Bechtle GmbH & Co. KG"]],
        col_widths=[3.0, 0.8, 1.2, 2.0]
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 1. WAS BRINGT MIR DIE AZURE LANDING ZONE?
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("1. Was bringt mir die Azure Landing Zone?", level=1)
    add_body(doc,
        "Die Azure Landing Zone ist kein einzelnes Azure-Produkt, sondern ein vollständig "
        "vorkonfiguriertes Fundament für Ihren Azure-Tenant. Anstatt jede neue Umgebung "
        "manuell einzurichten, steht Ihnen ab dem ersten Tag eine strukturierte, sichere "
        "und governance-konforme Basis zur Verfügung."
    )

    doc.add_heading("Was ist ab Tag 1 vorhanden?", level=2)
    add_table(doc,
        ["Komponente", "Was das bedeutet für Sie"],
        [
            ["7 Management Groups",
             "Alle Subscriptions sind automatisch strukturiert und erben Richtlinien – kein manuelles Einrichten je Projekt"],
            ["118 Policy Assignments",
             "Sicherheitsregeln und Compliance-Anforderungen gelten automatisch für jede neue Ressource"],
            ["Azure Firewall Standard",
             "Zentraler Sicherheitspunkt für den gesamten ein- und ausgehenden Datenverkehr – kein Perimeter-Loch"],
            ["Private DNS Zones",
             "PaaS-Dienste (z. B. Storage, SQL, Key Vault) sind ohne öffentliche Endpunkte erreichbar"],
            ["Log Analytics Workspace",
             "Alle Logs und Diagnose-Daten fließen zentral zusammen – ein einziger Ort für Monitoring und Fehlersuche"],
            ["Hub-and-Spoke-Netzwerk",
             "Neue Workload-Netzwerke (Spokes) können jederzeit sicher an den Hub angebunden werden"],
        ],
        col_widths=[2.2, 4.8]
    )

    doc.add_heading("Was Sie dadurch NICHT mehr manuell tun müssen", level=2)
    add_bullet(doc, "Sicherheitsrichtlinien je Projekt manuell konfigurieren – Policies greifen automatisch bei jeder neuen Ressource")
    add_bullet(doc, "Netzwerkfirewall je Workload separat einrichten – der Hub übernimmt das zentral")
    add_bullet(doc, "Logging-Infrastruktur aufbauen – der Log Analytics Workspace ist bereits aktiv")
    add_bullet(doc, "DNS für PaaS-Dienste manuell pflegen – Private DNS Zones lösen das automatisch auf")
    add_bullet(doc, "Berechtigungen unsortiert vergeben – RBAC-Rollen sind strukturiert auf MG-Ebene")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 2. IHR SETUP AUF EINEN BLICK
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("2. Ihr Setup auf einen Blick", level=1)

    doc.add_heading("Management-Group-Struktur", level=2)
    add_body(doc,
        "Die 7 Management Groups strukturieren alle Ihre Azure-Subscriptions nach Zweck "
        "und Sicherheitsprofil. Richtlinien werden einmal auf einer übergeordneten Ebene "
        "definiert und vererben sich automatisch nach unten."
    )

    tree_p = doc.add_paragraph(style="Normal")
    tree_r = tree_p.add_run(
        "Tenant Root  (Ihr Azure-Tenant)\n"
        "└─ alz  (Intermediate Root – alle ALZ-Policies)\n"
        "   ├─ alz-platform\n"
        "   │   └─ alz-platform-connectivity   → Connectivity Subscription\n"
        "   │                                    Hub-VNet, Firewall, DNS, Logging\n"
        "   ├─ alz-landingzones\n"
        "   │   └─ alz-landingzones-corp       → Produktion Subscription\n"
        "   │                                    Produktive Workloads, strenge Policies\n"
        "   └─ alz-sandbox                     → Sandbox Subscription\n"
        "                                        Tests, Entwicklung, gelockerte Policies"
    )
    tree_r.font.name = "Courier New"
    tree_r.font.size = Pt(9)
    doc.add_paragraph()

    doc.add_heading("Ihre drei Subscriptions", level=2)
    add_table(doc,
        ["Subscription", "Management Group", "Zweck", "Policies"],
        [
            ["ALZ-Connectivity",
             "alz-platform-connectivity",
             "Hub-Netzwerk, Azure Firewall, Log Analytics, Private DNS",
             "Platform-Policies (Monitoring, Backup, Guardrails)"],
            ["ALZ-Produktion",
             "alz-landingzones-corp",
             "Produktive Workloads – Spoke-VNets, Applikationen, Datenbanken",
             "Strengste Policies: Keine Public Endpoints, kein Public IP, Private DNS Pflicht"],
            ["ALZ-Sandbox",
             "alz-sandbox",
             "Entwicklung, Tests, Experimente",
             "Gelockerte Policies: kein ExpressRoute/VPN, keine Produktions-Ressourcen"],
        ],
        col_widths=[1.6, 2.0, 2.2, 2.2]
    )

    doc.add_heading("Kosten im Überblick", level=2)
    add_table(doc,
        ["Komponente", "Status", "Kosten/Monat"],
        [
            ["Azure Firewall Standard",   "Aktiv",           "ca. €700"],
            ["Log Analytics Workspace",   "Aktiv",           "ca. €50"],
            ["Private DNS Zones",         "Aktiv",           "ca. €15"],
            ["Azure Bastion",             "Zurückgestellt",  "ca. €120 (bei Aktivierung)"],
            ["VPN Gateway (VpnGw1AZ)",   "Zurückgestellt",  "ca. €140 (bei Aktivierung)"],
            ["DNS Private Resolver",      "Zurückgestellt",  "ca. €25 (bei Aktivierung)"],
            ["Gesamt (aktive Dienste)",   "–",               "ca. €765/Monat"],
        ],
        col_widths=[2.5, 1.8, 2.7]
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 3. ERSTE SCHRITTE NACH DEM DEPLOYMENT
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("3. Erste Schritte nach dem Deployment", level=1)
    add_body(doc,
        "Nach dem erfolgreichen Deployment können Sie die Landing Zone im Azure Portal "
        "in Betrieb nehmen. Die folgenden Schritte helfen Ihnen, die Umgebung zu prüfen "
        "und erste Ressourcen zu deployen."
    )

    doc.add_heading("Schritt 1 – Management Groups prüfen", level=2)
    add_bullet(doc, "Azure Portal öffnen: portal.azure.com")
    add_bullet(doc, "Suchfeld: Management Groups eingeben und aufrufen")
    add_bullet(doc, "Sie sehen die vollständige Hierarchie: alz → alz-platform → alz-platform-connectivity usw.")
    add_bullet(doc, "Prüfen: Alle 7 MGs sind sichtbar (Tenant Root + alz + 5 Kind-MGs)")
    add_hint(doc, "Direkt-Link: portal.azure.com/#view/Microsoft_Azure_ManagementGroups")

    doc.add_heading("Schritt 2 – Policy Compliance prüfen", level=2)
    add_bullet(doc, "Suchfeld: Policy eingeben → Azure Policy öffnen")
    add_bullet(doc, "Links: Compliance anklicken")
    add_bullet(doc, "Scope auf die Management Group alz setzen")
    add_bullet(doc, "Sie sehen alle 118 Assignments und deren aktuellen Compliance-Status")
    add_hint(doc,
        "Enforce-Guardrails-*-Initiativen sind initial auf DoNotEnforce gesetzt – "
        "sie evaluieren und melden, blockieren aber noch nicht. Das ermoeglicht ein "
        "schrittweises Onboarding ohne Produktionsunterbrechung."
    )

    doc.add_heading("Schritt 3 – Logging & Monitoring prüfen", level=2)
    add_bullet(doc, "Suchfeld: Log Analytics → Workspace law-alz-germanywestcentral öffnen")
    add_bullet(doc, "Links: Logs → KQL-Abfrage: AzureActivity | take 10")
    add_bullet(doc, "Activity-Logs aus allen Subscriptions sollten hier erscheinen")
    add_hint(doc, "Es dauert ca. 15–30 Minuten nach dem Deployment, bis erste Logs eintreffen.")

    doc.add_heading("Schritt 4 – Subscriptions in die richtigen MGs platzieren", level=2)
    add_body(doc,
        "Ihre Subscriptions müssen einmalig in die zugehörige Management Group verschoben werden. "
        "Das geht im Azure Portal unter Management Groups ohne Downtime oder Datenverlust."
    )
    add_table(doc,
        ["Subscription", "Ziel-Management-Group", "Vorgehen im Portal"],
        [
            ["ALZ-Connectivity", "alz-platform-connectivity",
             "MG öffnen → Subscription hinzufügen → Subscription-ID eintragen"],
            ["ALZ-Produktion",   "alz-landingzones-corp",
             "MG öffnen → Subscription hinzufügen → Subscription-ID eintragen"],
            ["ALZ-Sandbox",      "alz-sandbox",
             "MG öffnen → Subscription hinzufügen → Subscription-ID eintragen"],
        ],
        col_widths=[1.8, 2.2, 3.0]
    )
    add_hint(doc,
        "Sobald eine Subscription in einer MG ist, gelten automatisch alle Policies dieser MG "
        "und aller übergeordneten MGs. Bestehende Ressourcen werden evaluiert aber nicht sofort geblockt "
        "(DoNotEnforce-Modus der Guardrails)."
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 4. NEUE WORKLOADS HINZUFÜGEN
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("4. Neue Workloads hinzufügen", level=1)
    add_body(doc,
        "Jede neue produktive Applikation erhält ein eigenes Spoke-VNet in der "
        "ALZ-Produktion-Subscription. Das Spoke-VNet wird mit dem Hub gepeert und "
        "der gesamte Traffic läuft über die zentrale Azure Firewall."
    )

    doc.add_heading("Spoke-VNet für einen neuen Workload anlegen", level=2)
    add_table(doc,
        ["Schritt", "Aktion", "Detail"],
        [
            ["1", "Adressbereich wählen",
             "Neues /24 aus dem Bereich 10.2.0.0 – 10.255.0.0 wählen (kein Overlap mit 10.0.0.0/22)"],
            ["2", "VNet erstellen",
             "Subscription: ALZ-Produktion | Region: Germany West Central | Adressraum: z. B. 10.2.0.0/24"],
            ["3", "Subnetz anlegen",
             "Mindestens ein Workload-Subnetz (z. B. snet-app, 10.2.0.0/26)"],
            ["4", "Peering zum Hub",
             "VNet → Peerings → Hinzufügen → Remote VNet: vnet-alz-germanywestcentral (Connectivity Sub)"],
            ["5", "Route Table anlegen",
             "UDR mit Default-Route 0.0.0.0/0 → Next Hop: private IP der Azure Firewall"],
            ["6", "Route Table zuweisen",
             "UDR dem Workload-Subnetz zuweisen → Traffic läuft automatisch über Firewall"],
            ["7", "Firewall-Regel erstellen",
             "Im Firewall Policy: Application Rule oder Network Rule für erlaubten Traffic anlegen"],
        ],
        col_widths=[0.4, 2.0, 4.6]
    )
    add_hint(doc,
        "In der Corp-MG ist die Policy Deny-HybridNetworking aktiv – Spoke-VNets "
        "duerfen keine eigenen VPN Gateways haben. Der Weg ins On-Premises-Netz laeuft "
        "ausschliesslich ueber den Hub."
    )

    doc.add_heading("Welche Management Group für welchen Workload?", level=2)
    add_table(doc,
        ["Workload-Typ", "Richtige MG / Subscription", "Begründung"],
        [
            ["Produktive Applikation (intern)",
             "alz-landingzones-corp / ALZ-Produktion",
             "Strengste Policies, keine Public Endpoints, Firewall-Transit Pflicht"],
            ["Entwicklungs- / Testumgebung",
             "alz-sandbox / ALZ-Sandbox",
             "Gelockerte Policies, kein Produktions-Scope, isoliert von Corp"],
            ["Plattform-Dienst (Monitoring, Security-Tooling)",
             "alz-platform-connectivity / ALZ-Connectivity",
             "Infrastrukturelle Dienste, nicht für Workloads"],
        ],
        col_widths=[2.2, 2.3, 2.5]
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 5. GOVERNANCE & COMPLIANCE NUTZEN
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("5. Governance und Compliance im Alltag", level=1)
    add_body(doc,
        "Das Policy-Framework arbeitet im Hintergrund automatisch. Der Policy Compliance "
        "Report zeigt Ihnen jederzeit, welche Ressourcen konform sind und wo Handlungsbedarf besteht."
    )

    doc.add_heading("Policy Compliance Dashboard aufrufen", level=2)
    add_bullet(doc, "Azure Portal → Policy → Compliance")
    add_bullet(doc, "Scope auf 'alz' (Intermediate Root) setzen – zeigt alle MGs und Subscriptions")
    add_bullet(doc, "Filter: Non-compliant → zeigt nur Ressourcen mit Handlungsbedarf")
    add_bullet(doc, "Compliance-Score: Ziel ist >90 %; zu Beginn typisch 60–80 %")

    doc.add_heading("Was bedeuten die Policy-Effekte?", level=2)
    add_table(doc,
        ["Effekt", "Bedeutung", "Beispiel"],
        [
            ["Deny",
             "Ressource wird sofort abgelehnt – Deployment schlägt fehl",
             "Deny-Public-IP-On-NIC: Keine Public IPs an Netzwerkkarten in Corp"],
            ["Audit",
             "Ressource wird erstellt, aber als Non-compliant markiert",
             "Enforce-Guardrails (DoNotEnforce): Zeigt Verstoesse ohne zu blockieren"],
            ["DeployIfNotExists",
             "Azure ergänzt automatisch fehlende Konfiguration",
             "Deploy-VM-Monitoring: AMA Agent wird automatisch auf VMs installiert"],
            ["DoNotEnforce",
             "Policy wertet aus aber blockiert nicht – zum sicheren Onboarding",
             "Enforce-Guardrails-* initial: Audit ohne Produktionsunterbrechung"],
        ],
        col_widths=[1.6, 2.8, 2.6]
    )

    doc.add_heading("Remediation – Bestehende Ressourcen nachträglich konform machen", level=2)
    add_body(doc,
        "DeployIfNotExists-Policies können rückwirkend auf bestehende Ressourcen angewendet "
        "werden. Das ist sinnvoll direkt nach dem Platzieren einer Subscription in eine MG."
    )
    add_bullet(doc, "Policy → Remediation → 'Remediationsaufgabe erstellen'")
    add_bullet(doc, "Policy wählen (z. B. Deploy-VM-Monitoring) → Scope wählen → Starten")
    add_bullet(doc, "Status unter Remediationsaufgaben verfolgen (kann 15–60 Min. dauern)")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 6. MONITORING & LOGGING
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("6. Monitoring und Logging", level=1)
    add_body(doc,
        "Alle Logs laufen zentral im Log Analytics Workspace law-alz-germanywestcentral "
        "in der Connectivity-Subscription zusammen. Der Workspace ist bereits aktiv und "
        "empfaengt Daten aus allen Subscriptions."
    )

    doc.add_heading("Wichtige Abfragen (KQL)", level=2)
    add_body(doc, "Azure Portal → Log Analytics Workspace → law-alz-germanywestcentral → Logs", size=10)

    add_body(doc, "Alle Azure-Aktivitäten der letzten 24 Stunden:", size=10)
    add_code(doc, "AzureActivity | where TimeGenerated > ago(24h) | order by TimeGenerated desc")

    add_body(doc, "Fehlgeschlagene Deployments:", size=10)
    add_code(doc, "AzureActivity | where ActivityStatusValue == 'Failed' | take 50")

    add_body(doc, "Policy-Verstoesse (Deny-Aktionen):", size=10)
    add_code(doc, "AzureActivity | where OperationNameValue contains 'deny' | take 50")

    add_body(doc, "VM-Heartbeat (alle VMs die in den letzten 5 Min. gemeldet haben):", size=10)
    add_code(doc, "Heartbeat | where TimeGenerated > ago(5m) | summarize by Computer")

    doc.add_heading("Azure Monitor Alerts einrichten (Empfehlung)", level=2)
    add_table(doc,
        ["Alert", "Bedingung", "Aktion"],
        [
            ["Firewall-Bedrohung erkannt",
             "Azure Firewall Logs: ThreatIntel-Treffer",
             "E-Mail an Security-Verantwortlichen"],
            ["VM ohne Heartbeat",
             "Heartbeat | summarize LastCall = max(TimeGenerated) by Computer | where LastCall < ago(15m)",
             "E-Mail / Teams-Benachrichtigung"],
            ["Policy Deny ausgelöst",
             "AzureActivity: ActivityStatusValue = 'Failed', OperationName contains 'deny'",
             "E-Mail an zuständigen Admin"],
            ["Hohe Firewall-Bandbreite",
             "AzureFirewallApplicationRule: DataProcessed > Schwellwert",
             "Teams-Benachrichtigung"],
        ],
        col_widths=[1.8, 2.8, 2.4]
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 7. OPTIONALE ERWEITERUNGEN AKTIVIEREN
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("7. Optionale Erweiterungen aktivieren", level=1)
    add_body(doc,
        "Die folgenden Komponenten sind bereits in der Infrastruktur vorbereitet und "
        "können ohne strukturelle Änderungen durch einen einzelnen Parameter aktiviert werden. "
        "Alle Subnetze sind reserviert, die Konfiguration ist vorhanden."
    )

    doc.add_heading("VPN Gateway aktivieren – On-Premises-Anbindung", level=2)
    add_body(doc,
        "Sobald die VPN-Informationen des On-Premises-Standorts vorliegen, kann der "
        "VPN Gateway in wenigen Schritten aktiviert werden."
    )
    add_table(doc,
        ["Schritt", "Aktion"],
        [
            ["1", "kunden-minimal/hubnetworking.bicepparam öffnen"],
            ["2", "deployVpnGateway: false → deployVpnGateway: true setzen"],
            ["3", "Deployment ausführen: deploy-cli.ps1 -DeploymentScope Networking"],
            ["4", "Local Network Gateway anlegen (On-Prem IP + Adressbereiche)"],
            ["5", "VPN Connection anlegen (PSK aus VPN-Informationsanfrage)"],
            ["6", "Verbindungsstatus prüfen: VPN Gateway → Verbindungen → Status 'Verbunden'"],
        ],
        col_widths=[0.4, 6.6]
    )
    add_hint(doc,
        "Benötigte Informationen vorab (Anfrage an Pascal): Öffentliche IP der On-Prem Firewall, "
        "On-Prem Netzwerk-Adressbereiche, VPN-Gerät und Firmware-Version (IKEv2-Kompatibilität), "
        "Pre-Shared Key (PSK). "
        "Wichtig: Kein Adressbereich-Overlap mit dem Azure-Adressraum 10.0.0.0/22."
    )

    doc.add_heading("Azure Bastion aktivieren – Sicherer VM-Zugriff ohne Public IP", level=2)
    add_table(doc,
        ["Schritt", "Aktion"],
        [
            ["1", "kunden-minimal/hubnetworking.bicepparam öffnen"],
            ["2", "deployBastion: false → deployBastion: true setzen"],
            ["3", "Deployment ausführen: deploy-cli.ps1 -DeploymentScope Networking"],
            ["4", "Portal: Virtuelle Maschine → Verbinden → Bastion → Anmeldedaten eingeben"],
        ],
        col_widths=[0.4, 6.6]
    )
    add_hint(doc,
        "Azure Bastion ermöglicht Browser-basierten RDP/SSH-Zugriff auf VMs ohne Public IP. "
        "Empfohlen wenn kein VPN-Tunnel zur Verfügung steht. Mehrkosten: ca. €120/Monat."
    )

    doc.add_heading("DNS Private Resolver aktivieren – Benutzerdefinierte DNS-Auflösung", level=2)
    add_body(doc,
        "Der DNS Private Resolver ist sinnvoll, wenn Sie eigene On-Premises-DNS-Server "
        "haben und Azure-Ressourcen-Namen aus dem On-Premises-Netz auflösen möchten."
    )
    add_table(doc,
        ["Schritt", "Aktion"],
        [
            ["1", "kunden-minimal/hubnetworking.bicepparam öffnen"],
            ["2", "deployDnsPrivateResolver: false → deployDnsPrivateResolver: true setzen"],
            ["3", "Deployment ausführen: deploy-cli.ps1 -DeploymentScope Networking"],
            ["4", "On-Prem DNS-Server: Forwarder für *.privatelink.* auf Resolver Inbound IP eintragen"],
        ],
        col_widths=[0.4, 6.6]
    )

    doc.add_heading("Kosten nach Aktivierung aller optionalen Komponenten", level=2)
    add_table(doc,
        ["Komponente", "Monatliche Kosten"],
        [
            ["Azure Firewall Standard",  "ca. €700"],
            ["Log Analytics Workspace",  "ca. €50"],
            ["Private DNS Zones",        "ca. €15"],
            ["VPN Gateway (VpnGw1AZ)",  "ca. €140"],
            ["Azure Bastion (Standard)", "ca. €120"],
            ["DNS Private Resolver",     "ca. €25"],
            ["Gesamt",                   "ca. €1.050/Monat"],
        ],
        col_widths=[3.5, 3.5]
    )
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 8. SCHNELLREFERENZ
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("8. Schnellreferenz", level=1)

    doc.add_heading("Wichtige Azure Portal-Links", level=2)
    add_table(doc,
        ["Bereich", "Portal-Link"],
        [
            ["Management Groups",      "portal.azure.com/#view/Microsoft_Azure_ManagementGroups"],
            ["Policy Compliance",      "portal.azure.com/#view/Microsoft_Azure_Policy/PolicyMenuBlade/~/Compliance"],
            ["Log Analytics",          "portal.azure.com → Log Analytics Workspaces → law-alz-germanywestcentral"],
            ["Azure Firewall",         "portal.azure.com → Firewalls → afw-alz-germanywestcentral"],
            ["Monitor / Alerts",       "portal.azure.com/#view/Microsoft_Azure_Monitoring"],
            ["Subscription-Übersicht", "portal.azure.com/#view/Microsoft_Azure_Billing/SubscriptionsBladeV2"],
        ],
        col_widths=[2.2, 4.8]
    )

    doc.add_heading("Ressourcen-Namenskonvention", level=2)
    add_table(doc,
        ["Ressource", "Name"],
        [
            ["Log Analytics Workspace", "law-alz-germanywestcentral"],
            ["Managed Identity",        "mi-alz-germanywestcentral"],
            ["Hub VNet",                "vnet-alz-germanywestcentral"],
            ["Azure Firewall",          "afw-alz-germanywestcentral"],
            ["Firewall Public IP",      "pip-afw-alz-germanywestcentral"],
            ["VPN Gateway (bei Bedarf)", "vgw-alz-germanywestcentral"],
            ["Bastion (bei Bedarf)",    "bas-alz-germanywestcentral"],
            ["DNS Resolver (bei Bedarf)", "dnspr-alz-germanywestcentral"],
            ["Conn. Resource Group",    "rg-alz-conn-germanywestcentral"],
            ["Logging Resource Group",  "rg-alz-logging-germanywestcentral"],
            ["DNS Resource Group",      "rg-alz-dns-germanywestcentral"],
        ],
        col_widths=[3.0, 4.0]
    )

    doc.add_heading("Adressraumplan", level=2)
    add_table(doc,
        ["Netz", "Adressraum", "Verwendung"],
        [
            ["Hub VNet",                  "10.0.0.0/22",   "Gesamter Hub-Adressraum Germany West Central"],
            ["AzureFirewallSubnet",       "10.0.0.0/26",   "Azure Firewall"],
            ["AzureBastionSubnet",        "10.0.0.64/26",  "Bastion (reserviert)"],
            ["GatewaySubnet",             "10.0.0.128/27", "VPN Gateway (reserviert)"],
            ["DNS Inbound",               "10.0.0.160/28", "DNS Private Resolver Inbound (reserviert)"],
            ["DNS Outbound",              "10.0.0.176/28", "DNS Private Resolver Outbound (reserviert)"],
            ["AzureFirewallMgmtSubnet",   "10.0.0.192/26", "Firewall Management"],
            ["Spoke Produktion (Bsp.)",   "10.2.0.0/24",   "Erste Workload-Umgebung (erweiterbar)"],
            ["On-Premises (kein Overlap)", "≠ 10.0.0.0/22", "On-Prem-Netze dürfen nicht überlappen"],
        ],
        col_widths=[2.2, 1.8, 3.0]
    )

    doc.add_heading("Ansprechpartner", level=2)
    add_table(doc,
        ["Thema", "Ansprechpartner", "Kontakt"],
        [
            ["Azure Landing Zone, Deployment, Governance",
             "Yannick Beck – Bechtle GmbH & Co. KG",
             "yannick.beck@bechtle.com | +49 621 87503 328"],
            ["VPN-Anbindung On-Premises",
             "Pascal (On-Premises-Netzwerk)",
             "Interne Kontaktdaten"],
        ],
        col_widths=[2.5, 2.5, 2.0]
    )

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"Gespeichert: {OUTPUT}")


if __name__ == "__main__":
    build()
