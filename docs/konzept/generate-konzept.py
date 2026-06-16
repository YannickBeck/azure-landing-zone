# -*- coding: utf-8 -*-
"""
Generiert das Bechtle-Kundenkonzept "Azure Landing Zone" als Word-Dokument.

Basis-Vorlage : bechtle-brand/VORLAGE_Bechtle_Management_Summary.docx
Output        : Azure-Landing-Zone-Konzept.docx

Nutzung:
    python3 docs/konzept/generate-konzept.py

Abhängigkeit  : python-docx  (pip install python-docx)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(BASE_DIR, "bechtle-brand", "Word", "VORLAGE_Bechtle_Management_Summary.docx")
OUTPUT   = os.path.join(BASE_DIR, "Word", "Azure-Landing-Zone-Konzept.docx")

# Platzhalter – ersetzen, sobald Kundenname bekannt ist
KUNDE = "<KUNDE>"
DATE  = "16.06.2026"


# ─────────────────────────────────────────────────────────────────────────────
# Hilfsfunktionen
# ─────────────────────────────────────────────────────────────────────────────

def clear_body(doc):
    """Alle Body-Elemente entfernen; sectPr (Layout/Ränder) bleibt erhalten."""
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_body(doc, text: str):
    p = doc.add_paragraph(text, style="Normal")
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text: str, bold_prefix: str = None):
    p = doc.add_paragraph(style="Bechtle Aufzählung")
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.bold = True
    p.add_run(text)
    return p


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    try:
        tbl.style = "Bechtle Tabelle"
    except Exception:
        tbl.style = "Table Grid"
    # Header row
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, "053B25")
    # Data rows
    for row_data in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    # Widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return tbl


def update_header(doc, text: str):
    for para in doc.sections[0].header.paragraphs:
        if para.text.strip():
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = text
            else:
                para.add_run(text)
            break


def update_footer(doc, text: str):
    for para in doc.sections[0].footer.paragraphs:
        if para.text.strip():
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = text
            else:
                para.add_run(text)
            break


# ─────────────────────────────────────────────────────────────────────────────
# Hauptfunktion
# ─────────────────────────────────────────────────────────────────────────────

def build():
    doc = Document(TEMPLATE)
    clear_body(doc)
    update_header(doc, f"{KUNDE} | Azure Landing Zone")
    update_footer(doc, f"Bechtle | Azure Landing Zone – Konzept")

    # ═════════════════════════════════════════════════════════════════════════
    # 0. DECKBLATT
    # ═════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph(style="Standard klein")
    p.add_run("Bechtle GmbH & Co. KG · Gottlieb-Daimler-Straße 2 · 68165 Mannheim")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    tp = doc.add_paragraph(style="Title")
    tp.add_run("Azure Landing Zone – Konzept und Umsetzungsfahrplan")

    sp = doc.add_paragraph(style="Subtitle")
    sp.add_run(
        "Zielbild, Governance, Netzwerk, Sicherheit, Betrieb und Roadmap "
        "auf Basis des Microsoft ALZ Bicep Accelerator"
    )
    doc.add_paragraph()
    doc.add_paragraph()

    lg = doc.add_paragraph(style="Standard klein")
    lg.add_run(
        "Leistungsgrenze: Dieses Konzept umfasst den Aufbau der Azure Landing Zone "
        "(Grundgerüst, Governance, Netzwerk, Sicherheit, Betrieb) sowie eine initiale "
        "Pilotierung. Produktive Vollmigration, Anwendungs-Onboarding und laufender "
        "Betrieb sind eigenständige Folgeprojekte und nicht Bestandteil dieses Dokuments."
    )
    doc.add_paragraph()

    # ═════════════════════════════════════════════════════════════════════════
    # MANAGEMENT SUMMARY
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("Management Summary", level=1)

    add_body(doc,
        "Eine Azure Landing Zone ist die standardisierte, governance-konforme Grundstruktur "
        "für den Betrieb von Anwendungen und Workloads in Microsoft Azure – von Anfang an "
        "sicher, regelkonform und skalierbar. Sie definiert Management-Gruppen, Zugriffsrechte, "
        "Richtlinien, Netzwerk-Konnektivität und Monitoring als einheitliches Fundament, "
        "auf dem beliebige Fachapplikationen aufsetzen können."
    )
    add_body(doc,
        "Bechtle setzt auf den offiziellen Microsoft ALZ Bicep Accelerator – kein Eigenbau, "
        "sondern das von Microsoft entwickelte und gepflegte Referenz-Framework. Das gewährleistet "
        "Wartbarkeit, Herstellersupport und einfache Übernahme künftiger Updates. Der gesamte "
        "Aufbau erfolgt als Infrastructure as Code (Bicep, Azure Verified Modules) und wird "
        "über generierte CI/CD-Pipelines mit passwortloser Anmeldung (OIDC) reproduzierbar betrieben."
    )
    add_body(doc, "Die vier Kernnutzen im Überblick:")
    add_bullet(doc,
        "149 Policy-Definitionen, 42 Initiativen und 123 konkrete Richtlinien-Zuweisungen "
        "sorgen für automatisierte Leitplanken (Guardrails) auf allen Verwaltungsebenen – "
        "von Sicherheitsanforderungen bis zur Ressourcen-Klassifizierung.",
        "Governance und Policies"
    )
    add_bullet(doc,
        "Ein zentrales Hub-and-Spoke-Netzwerk (Azure Firewall, Bastion, Private DNS, "
        "optionale VPN/ExpressRoute-Anbindung) stellt sichere Konnektivität für alle "
        "Workloads bereit.",
        "Netzwerk-Hub"
    )
    add_bullet(doc,
        "Microsoft Defender for Cloud wird automatisch über Policy-Zuweisungen aktiviert "
        "und liefert eine durchgehende Sicherheitsbaseline für Subscriptions, VMs, "
        "Datenbanken und Containerworkloads.",
        "Sicherheit und Defender"
    )
    add_bullet(doc,
        "18 geordnete Deployment-Stufen, vollständig als Code versioniert und über "
        "What-If-Previews abgesichert, gewährleisten reproduzierbare, nachvollziehbare "
        "und rückrollbare Änderungen.",
        "Automatisierung und IaC"
    )
    add_body(doc,
        "Kostensteuerung: Im Microsoft-Standard entstehen ca. € 5.800/Monat (primär durch "
        "doppelte Firewalls, DDoS Protection und Gateways in zwei Regionen). Bechtle empfiehlt "
        "für den Einstieg einen kostenarmen Pilot-Pfad: Netzwerk-Typ »none« deployt zunächst "
        "nur Management-Gruppen, Policies und Logging (≈ € 0), bevor Netzwerkdienste "
        "schrittweise nach Bedarf aktiviert werden."
    )
    add_body(doc,
        "Empfohlener Einstieg: Pilot mit einer Subscription, network_type: none, "
        "Policies im DoNotEnforce-Modus (audit-only). Kein Kostenrisiko – "
        "volles Governance-Fundament vom ersten Tag an."
    )
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 1. AUSGANGSLAGE UND ZIELSETZUNG
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Ausgangslage und Zielsetzung", level=1)

    doc.add_heading("Ausgangslage", level=2)
    add_body(doc,
        "Der Kunde steht am Beginn seiner Azure-Reise: eine bestehende Subscription, "
        "noch kein etabliertes Cloud-Governance-Modell und kein vorheriger Kickoff. "
        "Es gibt keine strukturierten Leitplanken für Sicherheit, Zugriffsrechte oder "
        "Netzwerk-Konnektivität. Ressourcen werden ad hoc und ohne einheitlichen Rahmen erstellt."
    )
    add_body(doc,
        "Diese Ausgangssituation bietet die Chance, von Anfang an auf einer soliden, "
        "industrieerprobten Grundlage aufzubauen – statt technische Schulden durch "
        "nachträgliche Governance-Nachrüstung zu akkumulieren."
    )

    doc.add_heading("Zielbild", level=2)
    add_body(doc,
        "Das Zielbild ist eine standardisierte, mandantenweite Azure-Grundstruktur "
        "als stabiles Fundament für alle künftigen Workloads:"
    )
    add_bullet(doc, "Management-Group-Hierarchie (12 MGs) als strukturierter Verwaltungsrahmen für alle Subscriptions")
    add_bullet(doc, "Vollständiges ALZ-Policy-Set für automatische Governance-Durchsetzung ohne manuelle Eingriffe je Workload")
    add_bullet(doc, "Zentrales Logging und Monitoring als einheitlicher Sicherheits- und Betriebs-Backbone")
    add_bullet(doc, "Hub-and-Spoke-Netzwerk für sichere, private Konnektivität zwischen Workloads und On-Premises")
    add_bullet(doc, "Reproduzierbares Infrastructure-as-Code-Modell für nachvollziehbare, audit-fähige Änderungen")

    doc.add_heading("Abgrenzung", level=2)
    add_body(doc,
        "Die Azure Landing Zone ist das Fundament – nicht die Anwendung. Migration "
        "bestehender Workloads, Anwendungsentwicklung und laufender Applikationsbetrieb "
        "sind eigenständige Folgeprojekte. Der Scope dieses Konzepts endet mit der "
        "betriebsbereiten Landing Zone."
    )

    doc.add_heading("Erfolgskriterien", level=2)
    add_table(doc,
        ["Erfolgskriterium", "Messgröße / Nachweis"],
        [
            ["Management-Group-Hierarchie steht",
             "12 MGs im Azure Portal sichtbar, Tenant Root → alz → Unterebenen"],
            ["Policies greifen",
             "123 Assignments aktiv; Azure Policy Compliance-Report zeigt Ergebnisse"],
            ["Zentrales Logging aktiv",
             "Log Analytics Workspace vorhanden; Activity Logs aller Subscriptions fließen ein"],
            ["Netzwerk-Hub bereit",
             "Hub-VNet und Firewall-Policy vorhanden; Spoke kann gepeert werden"],
            ["IaC reproduzierbar",
             "Pipeline deployt ohne Fehler; What-If zeigt »keine Änderungen« nach Apply"],
        ],
        col_widths=[2.8, 4.2]
    )

    # ═════════════════════════════════════════════════════════════════════════
    # 2. METHODIK UND VORGEHEN
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Methodik und Vorgehen", level=1)

    doc.add_heading("Microsoft ALZ Bicep Accelerator", level=2)
    add_body(doc,
        "Grundlage ist der offizielle Microsoft ALZ Bicep Accelerator. Dieser Ansatz "
        "statt eines Eigenbaus bietet folgende Vorteile:"
    )
    add_bullet(doc,
        "Microsoft-Standard: dieselbe Basis, die Microsoft für Azure Landing Zones "
        "empfiehlt und für tausende Enterprise-Kunden einsetzt",
        "Standard"
    )
    add_bullet(doc,
        "Wartbarkeit: Updates werden vom Microsoft-Team gepflegt; neue ALZ-Versionen "
        "durch einfache Modulaktualisierung übernehmbar",
        "Wartbarkeit"
    )
    add_bullet(doc,
        "Vollständiges Policy-Set: 149 Definitionen, 42 Initiativen, 123 Assignments – "
        "automatisch aus der öffentlichen Microsoft Container Registry (MCR) bezogen",
        "Policy-Set"
    )
    add_bullet(doc,
        "Azure Verified Modules (AVM): alle IaC-Bausteine sind geprüfte, "
        "Microsoft-zertifizierte Bicep-Module (22 Module im Einsatz)",
        "AVM"
    )

    doc.add_heading("Infrastructure as Code und Pipelines", level=2)
    add_bullet(doc, "Deklarativ: Templates beschreiben den Zielzustand – Azure stellt sicher, dass er erreicht wird")
    add_bullet(doc, "Versioniert: jede Änderung nachvollziehbar, rückrollbar und audit-fähig (Git-History)")
    add_bullet(doc, "Passwortlos: OIDC Federated Identity – keine Passwörter oder Zertifikate in Pipelines")
    add_bullet(doc, "What-If vor Apply: Pipeline zeigt geplante Änderungen zur Freigabe, bevor etwas ausgeführt wird")
    add_bullet(doc, "Approval-Gate: explizite Freigabe erforderlich, bevor die Pipeline Änderungen am Tenant vornimmt")

    doc.add_heading("18 geordnete Deployment-Stufen", level=2)
    add_table(doc,
        ["#", "Stufe", "Scope", "Inhalt"],
        [
            ["1",  "Governance – Intermediate Root",    "Tenant", "alz-Root-MG + volles ALZ-Policy-Set (149 Defs / 42 Initiativen / 17 Assignments)"],
            ["2",  "Governance – Landing Zones",         "Tenant", "landingzones-MG + 53 Assignments"],
            ["3",  "Governance – LZ Corp",               "Tenant", "landingzones-corp + 5 Assignments (Deny Public Endpoints/IP, Hybrid Networking, Private DNS)"],
            ["4",  "Governance – LZ Online",             "Tenant", "landingzones-online"],
            ["5",  "Governance – LZ Local",              "Tenant", "landingzones-local + 1 Assignment (ALDO Services)"],
            ["6",  "Governance – Platform",              "Tenant", "platform-MG + 40 Assignments"],
            ["7",  "Governance – Platform Connectivity", "Tenant", "platform-connectivity + 1 Assignment (DDoS-VNET)"],
            ["8",  "Governance – Platform Identity",     "Tenant", "platform-identity + 4 Assignments"],
            ["9",  "Governance – Platform Management",   "Tenant", "platform-management"],
            ["10", "Governance – Platform Security",     "Tenant", "platform-security"],
            ["11", "Governance – Sandbox",               "Tenant", "sandbox + Enforce-ALZ-Sandbox Guardrails"],
            ["12", "Governance – Decommissioned",        "Tenant", "decommissioned + Enforce-ALZ-Decomm Guardrails"],
            ["13", "RBAC – Platform",                    "Tenant", "Rollen-Zuweisung auf platform-MG"],
            ["14", "RBAC – Platform Connectivity",       "Tenant", "Rollen-Zuweisung auf connectivity-MG"],
            ["15", "RBAC – Landing Zones",               "Tenant", "Rollen-Zuweisung auf landingzones-MG"],
            ["16", "Core Logging",                       "Sub",    "Log Analytics Workspace, 3 DCRs, Managed Identity (avm/ptn/alz/ama:0.2.0)"],
            ["17", "Hub Networking",                     "Sub",    "Hub-VNets, Firewall, Bastion, Gateways, DNS (aktiv bei hubNetworking)"],
            ["18", "Virtual WAN",                        "Sub",    "vWAN, vHub, Firewall (aktiv bei vwanConnectivity; Stufe 17 und 18 exklusiv)"],
        ],
        col_widths=[0.3, 1.9, 0.6, 4.2]
    )

    doc.add_heading("Qualitätssicherung", level=2)
    add_body(doc,
        "Alle 20 Bicep-Templates werden offline gegen die MCR gebaut und validiert – "
        "ohne Azure-Login. Die statische Validierung (bicep build) läuft in der "
        "CI-Pipeline. Status: 20/20 Templates grün, 0 Errors, 0 Warnings (Bicep 0.44.1)."
    )
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 3. ZIELARCHITEKTUR
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Zielarchitektur", level=1)

    doc.add_heading("3.1 Management-Group-Hierarchie", level=2)
    add_body(doc,
        "Die Management-Group-Hierarchie strukturiert alle Azure-Subscriptions des Kunden "
        "nach Zweck und Sicherheitsanforderungen. Policies, RBAC-Zuweisungen und Budgets "
        "werden auf MG-Ebene definiert und an alle untergeordneten Subscriptions vererbt."
    )

    tree_p = doc.add_paragraph(style="Normal")
    tree_r = tree_p.add_run(
        "Tenant Root\n"
        "└─ alz  (Intermediate Root)\n"
        "   ├─ alz-platform\n"
        "   │   ├─ alz-platform-connectivity   → Hub-Netzwerk, Firewall, DNS\n"
        "   │   ├─ alz-platform-identity       → Identity-Dienste\n"
        "   │   ├─ alz-platform-management     → Logging, Monitoring\n"
        "   │   └─ alz-platform-security       → Security-Tooling\n"
        "   ├─ alz-landingzones\n"
        "   │   ├─ alz-landingzones-corp       → intern, keine Public Endpoints\n"
        "   │   ├─ alz-landingzones-online     → internetseitige Workloads\n"
        "   │   └─ alz-landingzones-local      → souverän / vertraulich\n"
        "   ├─ alz-sandbox                     → Experimente (gelockerte Policies)\n"
        "   └─ alz-decommissioned              → Stilllegung (gesperrt)"
    )
    tree_r.font.name = "Courier New"
    tree_r.font.size = Pt(9)
    doc.add_paragraph()

    add_table(doc,
        ["Management Group", "Parent", "Zweck"],
        [
            ["alz (Intermediate Root)", "Tenant Root", "Einstieg für alle ALZ-Policies; alle anderen MGs erben von hier"],
            ["alz-platform", "alz", "Container für alle Plattform-Dienste-Subscriptions"],
            ["alz-platform-connectivity", "alz-platform", "Hub-Netzwerk, Firewall, DNS, Gateway-Subscriptions"],
            ["alz-platform-identity", "alz-platform", "Identity-Dienste (AD DS, Entra Connect – Roadmap)"],
            ["alz-platform-management", "alz-platform", "Logging (Log Analytics), Monitoring, Operations"],
            ["alz-platform-security", "alz-platform", "Security-Tooling, Defender-Konfiguration"],
            ["alz-landingzones", "alz", "Container für alle Workload-Subscriptions"],
            ["alz-landingzones-corp", "alz-landingzones", "Interne Workloads ohne Public Endpoints; strengste Policies"],
            ["alz-landingzones-online", "alz-landingzones", "Internetseitige Workloads (z. B. Public APIs, Websites)"],
            ["alz-landingzones-local", "alz-landingzones", "Souveräne / vertrauliche Workloads (Azure Local)"],
            ["alz-sandbox", "alz", "Experimentier-Bereich mit gelockerten Policies"],
            ["alz-decommissioned", "alz", "Stilllegungszone; Ressourcen-Erstellung gesperrt"],
        ],
        col_widths=[2.3, 1.8, 2.9]
    )

    doc.add_heading("3.2 Subscription-Strategie", level=2)
    add_body(doc,
        "Der Einstieg erfolgt mit einer einzelnen Subscription (Single-Subscription-Setup). "
        "Alle vier Plattform-Rollen (Management, Connectivity, Identity, Security) werden "
        "auf dieselbe Subscription abgebildet. Die Trennung in vier dedizierte "
        "Platform-Subscriptions kann jederzeit nachgeholt werden – die MG-Hierarchie "
        "und alle Policies bleiben dabei unverändert."
    )
    add_table(doc,
        ["Phase", "Subscription-Modell", "Anforderung", "Empfehlung"],
        [
            ["Pilot / Start",        "1 Subscription – alle Plattformrollen",       "EA / MCA / Pay-as-you-go", "Sofort starten"],
            ["Produktionsbetrieb",   "4 dedizierte Platform-Subscriptions",          "EA / MCA",                 "Nach Pilot"],
            ["Workload-Skalierung",  "Je Workload eigene Sub in Corp/Online/Local",  "EA / MCA",                 "Bei Bedarf"],
        ],
        col_widths=[1.4, 2.8, 1.8, 1.2]
    )
    p = doc.add_paragraph(style="Standard klein")
    p.add_run(
        "Hinweis: Kein Free-Tier-Account. "
        "Alle benötigten Resource Provider werden durch den Accelerator automatisch registriert."
    )

    doc.add_heading("3.3 Namens- und Regionskonzept", level=2)
    add_body(doc,
        "Primärregion: Germany West Central (germanywestcentral). "
        "Sekundärregion: North Europe (northeurope). "
        "Die Namensvergabe folgt dem Microsoft-Standard ohne Präfix oder Postfix."
    )
    add_table(doc,
        ["Element", "Konvention / Bereich"],
        [
            ["Management Groups",       "alz, alz-platform, alz-platform-{connectivity,identity,management,security}, alz-landingzones, …"],
            ["Resource Groups",         "rg-alz-<zweck>-<region>"],
            ["Logging",                 "law-alz-<region>, mi-alz-<region>, dcr-*-alz-<region>"],
            ["Netzwerk-Ressourcen",     "vnet-alz-hub-<region>, afw-alz-<region>, bas-alz-<region>"],
            ["Hub-VNet Primär",         "10.0.0.0/22 (Germany West Central)"],
            ["Hub-VNet Sekundär",       "10.1.0.0/22 (North Europe)"],
            ["Spoke-Netz (Beispiel)",   "10.2.0.0/24 (überlappungsfrei zu Hubs und On-Premises)"],
        ],
        col_widths=[2.0, 5.0]
    )
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 4. GOVERNANCE UND POLICIES
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Governance und Policies", level=1)

    add_body(doc,
        "Das Herzstück der Azure Landing Zone ist das vollständige ALZ-Policy-Set, "
        "bezogen über avm/ptn/alz/empty:0.3.6. Statt eines schlanken Eigenbau-Sets "
        "wird das vollständige, von Microsoft gepflegte Policy-Framework ausgerollt."
    )
    add_table(doc,
        ["Komponente", "Anzahl", "Beschreibung"],
        [
            ["Policy-Definitionen (Custom)", "149",
             "Monitoring (55), Network (20), Storage (16), SQL (13), Guardrails je Dienst …"],
            ["Policy-Set-Definitionen (Initiativen)", "42",
             "Deploy-MDFC-Config, Deploy-Private-DNS-Zones, Enforce-Guardrails-* je Dienst, Enforce-ACSB …"],
            ["Policy-Assignments", "123",
             "Direkt zugewiesen auf 9 MG-Ebenen; Kind-MGs erben zusätzlich"],
            ["Custom RBAC-Rollen", "5",
             "Subscription-Owner, Security-Operations, Network-Management, Application-Owners, Network-Subnet-Contributor"],
        ],
        col_widths=[2.4, 0.8, 3.8]
    )

    doc.add_heading("Assignment-Verteilung je Management-Group-Ebene", level=2)
    add_table(doc,
        ["Management Group", "Assignments", "Repräsentative Beispiele"],
        [
            ["alz (Intermediate Root)",   "17", "Deploy-MDFC-Config-H224, Deploy-MDEndpoints, Deploy-AzActivity-Log, Enforce-ACSB"],
            ["alz-landingzones",           "53", "Deny-Storage-http, Deny-MgmtPorts-Internet, Deploy-VM-Monitoring, Enforce-Guardrails-*"],
            ["alz-landingzones-corp",      "5",  "Deny-Public-Endpoints, Deny-Public-IP-On-NIC, Deny-HybridNetworking, Deploy-Private-DNS-Zones"],
            ["alz-landingzones-online",    "0",  "(erbt von landingzones)"],
            ["alz-landingzones-local",     "1",  "Enforce-ALDO-Services"],
            ["alz-platform",               "40", "Deploy-VM-Monitoring, Enforce-Backup, Enforce-Guardrails-*, DenyAction-DeleteUAMIAMA"],
            ["alz-platform-connectivity",  "1",  "Enable-DDoS-VNET"],
            ["alz-platform-identity",      "4",  "Deny-MgmtPorts-Internet, Deny-Public-IP, Deny-Subnet-Without-Nsg, Deploy-VM-Backup"],
            ["alz-platform-management",    "0",  "(erbt von platform)"],
            ["alz-platform-security",      "0",  "(erbt von platform)"],
            ["alz-sandbox",                "1",  "Enforce-ALZ-Sandbox"],
            ["alz-decommissioned",         "1",  "Enforce-ALZ-Decomm"],
            ["Summe",                      "123", ""],
        ],
        col_widths=[2.3, 0.9, 3.8]
    )

    doc.add_heading("DoNotEnforce-Modus für sanftes Onboarding", level=2)
    add_body(doc,
        "Viele Assignments – insbesondere die Enforce-Guardrails-*-Initiativen – "
        "sind im Auslieferungszustand auf DoNotEnforce gesetzt. Policies werden "
        "ausgewertet und im Compliance-Dashboard sichtbar, aber keine Ressource wird "
        "blockiert. Das ermöglicht ein sanftes, nicht-unterbrechendes Onboarding:"
    )
    add_bullet(doc, "Phase 1 – Audit-first: DoNotEnforce für alle Initiativen; Compliance-Dashboard zeigt Handlungsbedarf ohne Produktionsunterbrechung")
    add_bullet(doc, "Phase 2 – Selektives Enforcing: Initiativen je Workload-Typ und Compliance-Anforderung auf »Enabled« schalten")
    add_bullet(doc, "Phase 3 – Vollständiges Enforcement: alle relevanten Guardrails aktiv nach erfolgreicher Remediation")

    doc.add_heading("Was die Policies konkret bewirken", level=2)
    add_table(doc,
        ["Policy-Typ", "Beispiele", "Wirkung"],
        [
            ["Verschlüsselung",        "Deny-Storage-http, Enforce-TLS-SSL",           "Erzwingt HTTPS/TLS für alle Storage-Accounts und Services"],
            ["Netzwerk",               "Deny-Subnet-Without-Nsg, Deny-MgmtPorts",      "Jedes Subnetz braucht NSG; kein RDP/SSH aus Internet"],
            ["Diagnose / Monitoring",  "Deploy-AzActivity-Log, Deploy-Diag-LogsCat",   "Activity Logs und Diagnose automatisch in Log Analytics"],
            ["Defender",               "Deploy-MDFC-Config-H224, Deploy-MDEndpoints",  "Defender for Cloud automatisch per DeployIfNotExists aktiviert"],
            ["Backup",                 "Deploy-VM-Backup",                              "VMs ohne Backup-Tag automatisch in Recovery Vault gesichert"],
            ["Corp – Public Endpoints","Deny-Public-Endpoints, Deny-Public-IP-On-NIC", "Keine öffentlichen IPs oder PaaS-Endpunkte in corp-Zone"],
        ],
        col_widths=[1.8, 2.5, 2.7]
    )
    p = doc.add_paragraph(style="Standard klein")
    p.add_run(
        "Hinweis: Die Assignments Deploy-MDFC-Config-H224 und Deploy-MDEndpoints auf alz-Root-Ebene "
        "aktivieren Microsoft Defender for Cloud auf Standard-Tier – das verursacht laufende "
        "Kosten je Resource. Tiering ist über die Konfiguration steuerbar."
    )
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 5. NETZWERK-ARCHITEKTUR
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Netzwerk-Architektur", level=1)

    doc.add_heading("5.1 Hub-and-Spoke-Topologie", level=2)
    add_body(doc,
        "Die Netzwerk-Architektur folgt dem Hub-and-Spoke-Muster: Ein zentrales "
        "Hub-Netzwerk stellt Shared Services (Firewall, Bastion, DNS, Gateways) für "
        "alle angeschlossenen Workload-Netzwerke (Spokes) bereit. Der gesamte "
        "ausgehende und eingehende Datenverkehr der Spokes läuft über die Azure "
        "Firewall im Hub."
    )
    add_body(doc,
        "Es werden zwei Hub-VNets deployed – eines je Region – die bidirektional "
        "gepeert sind. Spokes werden pro Region an den jeweiligen Hub angebunden."
    )

    doc.add_heading("5.2 Subnetze und Dienste", level=2)
    add_table(doc,
        ["Subnetz", "Hub 1 Primär (GWC)", "Hub 2 Sekundär (NE)", "Zweck"],
        [
            ["AzureFirewallSubnet",               "10.0.0.0/26",     "10.1.0.0/26",     "Azure Firewall (Pflichtsubnetz)"],
            ["AzureFirewallManagementSubnet",      "10.0.0.192/26",   "10.1.0.192/26",   "Firewall Management Traffic"],
            ["AzureBastionSubnet",                 "10.0.0.64/26",    "10.1.0.64/26",    "Azure Bastion (sicheres RDP/SSH)"],
            ["GatewaySubnet",                      "10.0.0.128/27",   "10.1.0.128/27",   "VPN-/ExpressRoute-Gateways"],
            ["DNSPrivateResolverInboundSubnet",    "10.0.0.160/28",   "10.1.0.160/28",   "DNS Private Resolver Inbound"],
            ["DNSPrivateResolverOutboundSubnet",   "10.0.0.176/28",   "10.1.0.176/28",   "DNS Private Resolver Outbound"],
        ],
        col_widths=[2.4, 1.5, 1.5, 1.6]
    )

    add_table(doc,
        ["Dienst", "Default", "Kosten/Monat (ca.)", "Schalter"],
        [
            ["Azure Firewall Premium (je Hub)",          "AN (beide)",  "~€1.100/Hub",  "deployAzureFirewall"],
            ["Azure Bastion (Standard)",                  "AN (beide)",  "~€120/Hub",    "deployBastion"],
            ["VPN Gateway (VpnGw1AZ, BGP, ASN 65515)",   "AN (beide)",  "~€140/Hub",    "deployVpnGateway"],
            ["ExpressRoute Gateway",                      "AN (beide)",  "~€280/Hub",    "deployExpressRouteGateway"],
            ["DDoS Network Protection",                   "AN (Hub 1)",  "~€2.500",      "deployDdosProtectionPlan"],
            ["Private DNS Zones (alle Azure-Zonen)",      "AN (beide)",  "~€15",         "deployPrivateDnsZones"],
            ["DNS Private Resolver",                      "AN (beide)",  "~€25/Hub",     "deployDnsPrivateResolver"],
            ["Gesamt Default",                            "–",           "~€5.800",      "network_type: none → €0"],
        ],
        col_widths=[2.5, 0.9, 1.4, 2.2]
    )
    p = doc.add_paragraph(style="Standard klein")
    p.add_run(
        "Alle Dienste sind im Microsoft-Standard aktiviert. "
        "Für kostenarme Rollouts: network_type: none (nur MGs + Policies + Logging, ≈ €0) oder "
        "individuelle Schalter in der Bicep-Parameter-Datei auf false setzen."
    )

    doc.add_heading("5.3 DNS-Konzept", level=2)
    add_body(doc,
        "Das DNS-Konzept basiert auf Azure Private DNS Zones für alle PaaS-Dienste "
        "(Private Endpoints) und dem DNS Private Resolver für hybride Namensauflösung."
    )
    add_bullet(doc, "Private DNS Zones: zentral im Hub-VNet, VNet-Links in Spoke-Netzwerke")
    add_bullet(doc, "DNS Private Resolver: Inbound für On-Premises → Azure; Outbound für Weiterleitung zu On-Premises")
    add_bullet(doc, "Policy-Enforcement: Initiative Deploy-Private-DNS-Zones (59 Policies) stellt sicher, dass PaaS-Dienste mit Private Endpoints automatisch korrekt in DNS eingetragen werden")

    doc.add_heading("5.4 Alternative: Virtual WAN", level=2)
    add_body(doc,
        "Als Alternative kann Virtual WAN (vwanConnectivity) gewählt werden. "
        "Geeignet für Szenarien mit vielen Standorten oder bestehenden WAN-Strukturen."
    )
    add_table(doc,
        ["Aspekt", "Hub-and-Spoke", "Virtual WAN"],
        [
            ["Komplexität",           "Mittel",                      "Höher"],
            ["Kosten",                "Flexibel (je Dienst steuerbar)", "Höher (vHub-Gebühren)"],
            ["Eignung",               "Standardfall, 1–3 Standorte", "Viele Standorte, globales Routing"],
            ["Implementierungsstatus","Umgesetzt (default)",          "Vorbereitet / inaktive Alternative"],
        ],
        col_widths=[1.8, 2.2, 3.0]
    )
    add_bullet(doc, "Zu prüfen: Anzahl der Standorte, bestehende WAN-Infrastruktur, Routing-Anforderungen")
    add_bullet(doc, "Mögliche Folgestufe: Migration zu Virtual WAN bei Bedarf (Template liegt vor)")
    add_bullet(doc, "Abgrenzung: ExpressRoute/VPN-Konfiguration und IP-Adresskonzept On-Premises sind nicht Teil der Landing Zone")

    doc.add_heading("5.5 Bechtle-Empfehlung: ~€1.050/Monat-Konfiguration", level=2)
    add_body(doc,
        "Durch drei gezielte Anpassungen reduziert sich der monatliche Basispreis von "
        "~€5.800 auf ~€1.050 – ohne den Funktionsumfang der Landing Zone einzuschränken. "
        "Alle Governance-Objekte (MGs, Policies, RBAC), das Logging-Framework und "
        "die vollständige Hub-and-Spoke-Topologie mit zentraler Firewall bleiben erhalten."
    )
    add_table(doc,
        ["Anpassung", "Microsoft-Default", "Bechtle-Empfehlung", "Ersparnis"],
        [
            ["Azure Firewall SKU",
             "Premium × 2 Regionen (~€2.200)",
             "Standard × 1 Region (~€700)",
             "~€1.500"],
            ["DDoS Network Protection",
             "AN (~€2.500/Monat)",
             "AUS – bei internet-facing Workloads aktivieren",
             "~€2.500"],
            ["ExpressRoute Gateway",
             "AN × 2 Regionen (~€560)",
             "AUS – bei ER-Leitung aktivieren",
             "~€560"],
            ["VPN Gateway",
             "AN × 2 Regionen (~€280)",
             "AN × 1 Region (~€140)",
             "~€140"],
            ["Azure Bastion",
             "AN × 2 Regionen (~€240)",
             "AN × 1 Region (~€120)",
             "~€120"],
            ["DNS Private Resolver",
             "AN × 2 Regionen (~€50)",
             "AN × 1 Region (~€25)",
             "~€25"],
            ["Netzwerk gesamt",
             "~€5.800",
             "~€1.050",
             "~€4.750"],
        ],
        col_widths=[2.0, 1.9, 2.1, 1.0]
    )
    add_body(doc,
        "Was bleibt vollständig erhalten:"
    )
    add_bullet(doc, "Alle 149 Policy-Definitionen, 42 Initiativen, 123 Assignments – voller Governance-Umfang")
    add_bullet(doc, "Hub-and-Spoke-Topologie mit zentraler Azure Firewall (Datenverkehr-Inspektion, FQDN-Filterung, Threat Intelligence)")
    add_bullet(doc, "Azure Bastion für sicheres RDP/SSH ohne öffentliche IPs")
    add_bullet(doc, "VPN Gateway für On-Premises-Konnektivität (BGP, active-active, ASN 65515)")
    add_bullet(doc, "Private DNS Zones + DNS Private Resolver für hybride Namensauflösung")
    add_bullet(doc, "Zentrales Logging (Log Analytics, DCRs, DINE-Policies)")
    add_bullet(doc, "Alle 5 Custom RBAC-Rollen und 18 Deploy-Stufen")

    add_body(doc, "Was verschoben wird (deferred, nicht gestrichen):")
    add_bullet(doc,
        "Sekundäre Region (North Europe): Deployment auf Primärregion beschränkt; "
        "Ausbau auf zwei Regionen bei Bedarf jederzeit möglich",
        "Zweite Region"
    )
    add_bullet(doc,
        "Azure Firewall Premium: Standard → Premium-Upgrade ist unterbrechungsfrei möglich; "
        "Premium-Features (IDPS, TLS-Inspektion, URL-Filterung) bei konkretem Bedarf aktivieren",
        "Firewall Premium"
    )
    add_bullet(doc,
        "DDoS Network Protection: erst sinnvoll bei internet-facing Workloads mit öffentlichen IPs; "
        "kann per einzelnem Schalter (deployDdosProtectionPlan: true) nachträglich aktiviert werden",
        "DDoS Protection"
    )
    add_bullet(doc,
        "ExpressRoute Gateway: erst bei bestellter ER-Leitung relevant; "
        "VPN Gateway übernimmt in der Zwischenzeit die On-Prem-Anbindung",
        "ExpressRoute"
    )

    add_body(doc, "Bicep-Parameter für die Bechtle-Empfehlung (in templates/networking/hubnetworking/main.bicepparam):")
    code_p = doc.add_paragraph(style="Normal")
    code_r = code_p.add_run(
        "azureFirewallSettings:        { deployAzureFirewall: true, azureFirewallTier: 'Standard' }\n"
        "bastionHostSettings:          { deployBastion: true }\n"
        "vpnGatewaySettings:           { deployVpnGateway: true }\n"
        "expressRouteGatewaySettings:  { deployExpressRouteGateway: false }  // bei Bedarf: true\n"
        "ddosProtectionPlanSettings:   { deployDdosProtectionPlan: false }   // bei Bedarf: true\n"
        "privateDnsSettings:           { deployPrivateDnsZones: true, deployDnsPrivateResolver: true }\n"
        "// Sekundär-Hub (northeurope): Deployment-Stufe 17 für zweite Region nicht ausführen"
    )
    code_r.font.name = "Courier New"
    code_r.font.size = Pt(8)
    doc.add_paragraph()

    p = doc.add_paragraph(style="Standard klein")
    p.add_run(
        "Hinweis: Der Wechsel von Firewall Standard auf Premium ist ein In-Place-Upgrade "
        "ohne Neukonfiguration der Firewall-Regeln. Der Wechsel von einer auf zwei Regionen "
        "erfordert einen weiteren Bootstrap-Lauf mit aktualisierter Regionskonfiguration."
    )
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 6. SICHERHEIT
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Sicherheit", level=1)

    add_body(doc,
        "Die Sicherheitsarchitektur ist mehrschichtig: präventive Kontrollen "
        "(Netzwerksegmentierung, Policy-Enforcing), detective Controls "
        "(Logging, Monitoring, Defender for Cloud) und automatische Remediation."
    )

    doc.add_heading("Microsoft Defender for Cloud", level=2)
    add_body(doc,
        "Microsoft Defender for Cloud wird nicht manuell konfiguriert, sondern über "
        "Policy-Assignments auf alz-Root-Ebene automatisch aktiviert (DeployIfNotExists). "
        "Jede neue Subscription in der MG-Hierarchie erhält automatisch die Defender-Konfiguration."
    )
    add_table(doc,
        ["Assignment", "Effekt", "Kosten"],
        [
            ["Deploy-MDFC-Config-H224",     "Aktiviert Defender for Cloud Konfiguration + Security Contacts", "Basis kostenlos; Pläne je Ressource"],
            ["Deploy-MDEndpoints",          "Defender for Endpoint Agent auf VMs",                           "Im Server-Plan enthalten"],
            ["Deploy-MDEndpointsAMA",       "Defender for Endpoint Integration mit MDE",                      "Im Server-Plan enthalten"],
            ["Deploy-MDFC-OssDb",           "Advanced Threat Protection für Open-Source-DBs",                "~€13/Server/Monat"],
            ["Deploy-MDFC-SqlAtp",          "Defender für SQL Server und Managed Instances",                  "~€13/Server/Monat"],
            ["Deploy-ASC-Monitoring",       "Microsoft Cloud Security Benchmark v1",                         "Kostenlos (Audit)"],
            ["Deploy-MCSB2-Monitoring",     "Microsoft Cloud Security Benchmark v2",                         "Kostenlos (Audit)"],
            ["Enforce-ACSB",                "Azure Compute Security Baseline Compliance-Auditing",           "Kostenlos (Audit)"],
        ],
        col_widths=[2.4, 2.8, 1.8]
    )

    doc.add_heading("Zentrales Logging als Sicherheits-Backbone", level=2)
    add_body(doc,
        "Der Log Analytics Workspace (law-alz-<region>, 365 Tage Retention) ist das "
        "zentrale Sicherheits-Repository. Policy-Assignments speisen automatisch ein:"
    )
    add_bullet(doc, "Activity Logs aller Subscriptions (Deploy-AzActivity-Log)")
    add_bullet(doc, "Diagnose-Kategorien aller unterstützten Ressourcen (Deploy-Diag-LogsCat)")
    add_bullet(doc, "VM Insights, Change Tracking und Defender-for-SQL-Daten (über Data Collection Rules)")
    add_bullet(doc, "Service Health Alerts (Deploy-SvcHealth-BuiltIn)")

    doc.add_heading("Optional: Microsoft Sentinel (Roadmap)", level=2)
    add_bullet(doc, "Zu prüfen: SIEM-Anforderungen, Compliance-Verpflichtungen, Security-Operations-Kapazitäten")
    add_bullet(doc, "Mögliche Folgestufe: Sentinel-Onboarding mit Data Connectors (Azure AD, Defender for Cloud, Office 365) und Analyseregeln")
    add_bullet(doc, "Abgrenzung: Sentinel-Konfiguration, Playbooks und Incident-Handling sind nicht Teil der Landing Zone")
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 7. MONITORING UND ZENTRALES LOGGING
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("7. Monitoring und zentrales Logging", level=1)

    add_body(doc,
        "Das zentrale Logging und Monitoring wird über das AVM-Pattern (avm/ptn/alz/ama:0.2.0) "
        "bereitgestellt. Alle benötigten Data Collection Rules werden automatisch erstellt "
        "und über Policy-Assignments mit Ressourcen verbunden."
    )
    add_table(doc,
        ["Ressource", "Name / Konvention", "AVM-Modul", "Kosten"],
        [
            ["Resource Group",               "rg-alz-logging-<region>",           "avm/res/resources/resource-group:0.4.3",          "Kostenlos"],
            ["Log Analytics Workspace",      "law-alz-<region> (PerGB2018, 365 Tage)", "avm/res/operational-insights/workspace:0.14.2", "5 GB/Mon. frei; dann ~€2/GB"],
            ["User-Assigned Managed Identity","mi-alz-<region>",                  "avm/ptn/alz/ama:0.2.0",                           "Kostenlos"],
            ["DCR – VM Insights",            "dcr-vmi-alz-<region>",              "avm/ptn/alz/ama:0.2.0",                           "Im LAW enthalten"],
            ["DCR – Change Tracking",        "dcr-ct-alz-<region>",               "avm/ptn/alz/ama:0.2.0",                           "Im LAW enthalten"],
            ["DCR – Defender SQL",           "dcr-mdfcsql-alz-<region>",          "avm/ptn/alz/ama:0.2.0",                           "Im LAW enthalten"],
            ["LAW Solution",                 "ChangeTracking",                     "(im Workspace)",                                  "Im LAW enthalten"],
            ["Automation Account (optional)","aa-alz-<region> (Basic)",           "avm/res/automation/automation-account:0.17.1",    "Default aus; bei Bedarf"],
        ],
        col_widths=[1.8, 2.2, 2.2, 1.3]
    )
    add_body(doc,
        "DINE-Policies (Deploy-VM-Monitoring, Deploy-VM-ChangeTrack u. a.) auf den "
        "landingzones- und platform-MG-Ebenen stellen sicher, dass jede neue VM oder "
        "Scale Set automatisch mit dem Log Analytics Workspace und den DCRs verbunden wird – "
        "ohne manuelle Konfiguration."
    )

    # ═════════════════════════════════════════════════════════════════════════
    # 8. IDENTITY UND RBAC
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("8. Identity und RBAC", level=1)

    doc.add_heading("5 Custom RBAC-Rollen", level=2)
    add_table(doc,
        ["Rolle", "Beschreibung", "Typische Zuweisung"],
        [
            ["Subscription-Owner (alz)",         "Delegierter Owner für Subscriptions – eingeschränkt, ohne volle Owner-Rechte (1 Action)", "Plattform-Admins"],
            ["Security-Operations (alz)",         "Horizontale Sicherheitssicht über alle Subscriptions (12 Actions)",                       "Security-Team"],
            ["Network-Management (alz)",          "Plattformweites Netzwerk-Management: VNets, UDRs, NSGs, NVAs (4 Actions)",               "Netzwerk-Team"],
            ["Application-Owners (alz)",          "Contributor-Rechte auf Resource-Group-Ebene für Applikations-/Ops-Teams (1 Action)",      "Applikations-Teams"],
            ["Network-Subnet-Contributor (alz)",  "Vollständiger Zugriff auf Subnetz-Management (8 Actions)",                               "Netzwerk-Ops"],
        ],
        col_widths=[2.2, 3.0, 1.8]
    )

    doc.add_heading("RBAC-Zuweisungsmodell", level=2)
    add_body(doc,
        "RBAC-Zuweisungen werden als Entra-ID-Gruppen-Object-IDs in den Bicep-Parameter-"
        "Dateien gepflegt und pro MG zugewiesen. Ein leeres Array ist ein No-Op – "
        "der Schritt läuft gefahrlos ohne Zuweisungen. Object-IDs werden nach dem "
        "Kickoff durch den Kunden befüllt."
    )

    doc.add_heading("Identity-Domäne (Roadmap)", level=2)
    add_bullet(doc, "Zu prüfen: Entra-ID-Struktur, PIM-Anforderungen, Conditional Access Policies, hybride Identitäten")
    add_bullet(doc, "Mögliche Folgestufe: Identity-Baseline mit Entra ID Diagnostik, PIM-Konfiguration und Conditional Access")
    add_bullet(doc, "Abgrenzung: Entra ID Tenant-Konfiguration, MFA-Rollout und Identity-Governance sind nicht Bestandteil der Landing Zone")
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 9. AUTOMATISIERUNG, CI/CD UND BETRIEB
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("9. Automatisierung, CI/CD und Betrieb", level=1)

    doc.add_heading("Bootstrap – Phase 0", level=2)
    add_body(doc,
        "Bevor die eigentliche Landing Zone deployed wird, richtet der Accelerator "
        "in Phase 0 alle notwendigen Infrastrukturen für Pipelines und Identitäten ein. "
        "Der Bootstrap ist Terraform-basiert – auch wenn das IaC-Modell Bicep ist."
    )
    add_body(doc, "Was der Bootstrap erstellt:")
    add_bullet(doc, "Bootstrap-Resource-Group + Storage Account (Terraform State)")
    add_bullet(doc, "Managed Identity mit Federated Credentials (OIDC) – keine Passwörter, keine Zertifikate")
    add_bullet(doc, "GitHub-Repository mit dem gesamten Starter-Modul und allen Deployment-Templates")
    add_bullet(doc, "GitHub Environments mit Approval-Gates (apply_approvers: Freigabe-E-Mail)")
    add_bullet(doc, "Generierte Deployment-Pipelines (GitHub Actions Workflows) für alle 18 Stufen")

    add_table(doc,
        ["Voraussetzung", "Detail"],
        [
            ["PowerShell",           "≥ 7.4"],
            ["Azure CLI",            "≥ 2.60, angemeldet via az login"],
            ["ALZ PS-Modul",         "Install-Module ALZ -Scope CurrentUser"],
            ["Azure-Berechtigungen", "Owner auf den Subscriptions + erhöhter Zugriff am Tenant Root (MG-Erstellung)"],
            ["GitHub PAT (Classic)", "Scopes: repo, workflow, admin:org, delete_repo"],
            ["Erhöhter Zugriff",     "Entra ID → Properties → Access management for Azure resources → Ein (einmalig; danach deaktivieren)"],
        ],
        col_widths=[1.8, 5.2]
    )

    doc.add_heading("Pipeline-Workflow", level=2)
    add_table(doc,
        ["Schritt", "Beschreibung", "Azure-Wirkung"],
        [
            ["1. Trigger",    "Push oder Pull Request löst Pipeline aus",                    "Keine"],
            ["2. Validation", "bicep build aller Templates (statische Validierung)",          "Keine"],
            ["3. What-If",    "Vorschau aller geplanten Änderungen je Deploy-Stufe",         "Keine"],
            ["4. Approval",   "Freigabe durch apply_approvers in GitHub Environment",        "Keine"],
            ["5. Apply",      "Sequentielle Ausführung der 18 Deploy-Stufen",               "Erzeugt/aktualisiert Azure-Ressourcen"],
        ],
        col_widths=[1.3, 3.2, 2.5]
    )

    doc.add_heading("GitOps-Betriebsmodell", level=2)
    add_bullet(doc, "Nachvollziehbarkeit: jede Änderung mit Autor, Timestamp und Begründung im Git-Log")
    add_bullet(doc, "Rückrollbarkeit: jede Änderung kann durch neuen Commit rückgängig gemacht werden")
    add_bullet(doc, "Audit-Fähigkeit: Git-History + Azure Activity Log = vollständiges Änderungsprotokoll")
    add_bullet(doc, "Kollaboration: mehrere Teammitglieder können parallel arbeiten, Code Reviews sind Standard")
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 10. KOSTEN UND KOSTENSTEUERUNG
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("10. Kosten und Kostensteuerung", level=1)

    add_body(doc,
        "Die Kosten der Azure Landing Zone werden primär durch die Netzwerk-Dienste "
        "bestimmt. Governance-Objekte (MGs, Policies, RBAC), Logging und "
        "Identity-Ressourcen sind nahezu kostenlos. Bechtle empfiehlt eine "
        "zielgerichtete Konfiguration mit ~€1.050/Monat, die alle wesentlichen "
        "Funktionen enthält und durch nachträgliche Schalter auf den vollen "
        "Microsoft-Standard (~€5.800/Monat) erweiterbar ist."
    )

    doc.add_heading("Kostenvergleich: Varianten im Überblick", level=2)
    add_table(doc,
        ["Dienst / Komponente", "Microsoft-Default", "Bechtle-Empfehlung (~€1.050)", "Pilot (€0)"],
        [
            ["Azure Firewall",              "2x Premium ~€2.200", "1x Standard ~€700",   "–"],
            ["DDoS Network Protection",     "~€2.500",            "–",                    "–"],
            ["ExpressRoute Gateway",        "2x ~€560",           "–",                    "–"],
            ["VPN Gateway",                 "2x ~€280",           "1x ~€140",             "–"],
            ["Azure Bastion",               "2x ~€240",           "1x ~€120",             "–"],
            ["DNS Private Resolver",        "2x ~€50",            "1x ~€25",              "–"],
            ["Private DNS Zones",           "~€15",               "~€15",                 "–"],
            ["Log Analytics Workspace",     "~€50",               "~€50",                 "~€50"],
            ["MGs / Policies / RBAC",       "€0",                 "€0",                   "€0"],
            ["Gesamt",                      "~€5.800",            "~€1.050",              "ca. €50"],
            ["Regionen",                    "GWC + NE (beide)",   "GWC (Primär)",         "GWC"],
            ["Netzwerk-Typ",                "hubNetworking",      "hubNetworking",        "none"],
        ],
        col_widths=[2.1, 1.6, 2.0, 1.3]
    )

    doc.add_heading("Bechtle-Empfehlung: ~€1.050/Monat", level=2)
    add_body(doc,
        "Die Bechtle-Konfiguration stellt alle produktionsrelevanten Netzwerk- und "
        "Sicherheitsdienste in der Primärregion (Germany West Central) bereit. "
        "Kein Funktionsverlust gegenueber dem Microsoft-Default – lediglich "
        "Dienste, die erst bei konkretem Bedarf relevant werden, sind initial deaktiviert."
    )
    add_table(doc,
        ["Funktion", "Status", "Aktivierbar durch"],
        [
            ["Management Groups (12 MGs)",              "Aktiv", "–"],
            ["Policy-Set (149 Defs / 123 Assignments)", "Aktiv", "–"],
            ["Zentrales Logging (LAW + 3 DCRs)",        "Aktiv", "–"],
            ["Hub-and-Spoke-Topologie",                 "Aktiv", "–"],
            ["Azure Firewall (Standard, 1 Region)",     "Aktiv", "–"],
            ["Azure Bastion (1 Region)",                "Aktiv", "–"],
            ["VPN Gateway (1 Region)",                  "Aktiv", "–"],
            ["Private DNS Zones",                       "Aktiv", "–"],
            ["DNS Private Resolver (1 Region)",         "Aktiv", "–"],
            ["Firewall Premium (IDPS, TLS)",      "Deferred", "azureFirewallTier: Premium (In-Place-Upgrade, kein Rebuild)"],
            ["DDoS Network Protection",           "Deferred", "deployDdosProtectionPlan: true (bei internet-facing Workloads)"],
            ["ExpressRoute Gateway",              "Deferred", "deployExpressRouteGateway: true (bei ER-Leitungsbestellung)"],
            ["Sekundaere Region (North Europe)",  "Deferred", "Zweiten Hub-Deployment-Lauf ausfuehren"],
        ],
        col_widths=[2.6, 0.8, 3.6]
    )

    doc.add_heading("Gestufter Ausbau", level=2)
    add_table(doc,
        ["Stufe", "Konfiguration", "Kosten/Monat", "Ausloser"],
        [
            ["1 – Governance-Pilot",
             "network_type: none",
             "ca. €50 (nur LAW)",
             "Sofort – MGs, Policies, Logging, kein Kostenrisiko"],
            ["2 – Bechtle-Empfehlung",
             "Firewall Standard, 1 Region, kein DDoS/ER",
             "~€1.050",
             "Erste Workloads; On-Prem-Anbindung geplant"],
            ["3 – Geo-Redundanz",
             "+ Sekundar-Hub North Europe",
             "~€1.800",
             "SLA-Anforderungen erfordern zweite Region"],
            ["4 – Firewall Premium",
             "azureFirewallTier: Premium",
             "~€2.400",
             "IDPS / TLS-Inspektion / URL-Filterung benoetigt"],
            ["5 – DDoS Protection",
             "deployDdosProtectionPlan: true",
             "~€4.900",
             "Internet-facing Workloads mit DDoS-Risiko"],
            ["6 – Microsoft-Default",
             "Alle Dienste, beide Regionen",
             "~€5.800",
             "Vollstaendiger Enterprise-Standard"],
        ],
        col_widths=[1.5, 2.3, 1.0, 2.2]
    )
    p = doc.add_paragraph(style="Standard klein")
    p.add_run(
        "Alle Stufen sind additive Aenderungen an bestehenden Parameter-Dateien. "
        "Kein Rueckbau von bereits deployten Ressourcen erforderlich. "
        "Der Standard-zu-Premium-Upgrade der Azure Firewall erfolgt ohne Unterbrechung "
        "und ohne Neukonfiguration der Firewall-Regeln."
    )

    doc.add_heading("Konfigurationsvergleich: Vor- und Nachteile", level=2)
    add_body(doc,
        "Die folgende Tabelle bewertet jede Ausbaustufe anhand von Vorteilen, Nachteilen "
        "und typischen Einsatzszenarien – als Entscheidungshilfe für das initiale Sizing."
    )
    add_table(doc,
        ["Konfiguration", "Kosten/Mon.", "Vorteile", "Nachteile", "Empfohlen für"],
        [
            ["Pilot\n(network_type:\nnone)",
             "~€50",
             "+ Nullrisiko\n+ Vollständige Governance (MGs, Policies, Logging)\n+ Ideal für PoC und Compliance-Audit",
             "– Keine Netzwerkkonnektivität\n– Workloads können nicht deployt werden\n– Kein Firewall-/Bastion-Schutz",
             "Einstieg, PoC, Compliance-Audit ohne Netzwerkbedarf"],
            ["Bechtle-Empfehlung\n(Standard, 1 Region)",
             "~€1.050",
             "+ Voller Funktionsumfang\n+ Firewall, Bastion, VPN, DNS aktiv\n+ 4× günstiger als Microsoft-Default\n+ In-Place-Upgrade auf Premium möglich",
             "– Keine zweite Region (kein Geo-Failover)\n– Kein IDPS / TLS-Inspektion\n– Kein DDoS-Schutz",
             "Standard-Produktionsumgebungen, erste Workloads"],
            ["Geo-Redundanz\n(+ Hub North Europe)",
             "~€1.800",
             "+ Regionale Ausfallsicherheit\n+ Zwei unabhängige Netzwerkhubs\n+ Automatisches Routing-Failover",
             "– Doppelte Netzwerkkosten\n– Erhöhte Betriebskomplexität\n– Kein IDPS / kein DDoS",
             "SLA-kritische Produktionssysteme, Multi-Region-Anforderungen"],
            ["Firewall Premium\n(+ IDPS / TLS)",
             "~€2.400",
             "+ IDPS (Intrusion Detection & Prevention)\n+ TLS-Inspektion verschlüsselter Verbindungen\n+ URL-Filterung und erweiterte Bedrohungsabwehr\n+ In-Place-Upgrade ohne Rebuild",
             "– +€1.350 ggü. Bechtle-Empfehlung\n– TLS-Inspektion erfordert Zertifikat-Rollout\n– Erhöhter Konfigurations-aufwand",
             "Hohe Sicherheits-/Compliance-Anforderungen, regulierte Branchen"],
            ["DDoS Protection\n(Network Plan)",
             "~€4.900",
             "+ Vollständiger DDoS-Schutz für alle öffentlichen IPs\n+ Automatische Traffic-Analyse und Mitigation\n+ SLA-Garantie bei DDoS-Ereignissen",
             "– ~€2.500/Monat Fixkosten (Plan Fee) unabhängig von Angriffen\n– Nur relevant bei internet-facing Workloads\n– Kein Schutz für interne Kommunikation",
             "Internet-facing Workloads mit realen DDoS-Risiken"],
            ["Microsoft-Default\n(alle Dienste,\n2 Regionen)",
             "~€5.800",
             "+ Maximale Redundanz und Sicherheit\n+ Enterprise-Standard out-of-the-box\n+ Beide Regionen + Premium FW + DDoS\n+ Herstellerempfehlung für Enterprise",
             "– Höchste Kosten aller Varianten\n– Viele Komponenten initial oft ungenutzt\n– Lange Deployment-Zeit für alle Stufen\n– Overkill für die meisten Startszenarien",
             "Unternehmenskritische Systeme mit strengsten Compliance-Anforderungen"],
        ],
        col_widths=[1.5, 0.8, 2.3, 2.3, 1.7]
    )

    doc.add_page_break()
    # 11. ROADMAP, PHASEN UND GATES
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("11. Roadmap, Phasen und Gates", level=1)

    add_table(doc,
        ["Phase", "Inhalt", "Gate / Ergebnis", "Kosten"],
        [
            ["1 – Beratung & Kickoff",
             "Discovery-Workshop: Anforderungen, Subscription-Strategie, Netzwerk-Scope, Compliance",
             "Entscheidungsprotokoll: Region, network_type, DDoS, Enforcement-Zeitpunkt",
             "≈ €0"],
            ["2 – Bootstrap & Grundgerüst",
             "Deploy-Accelerator ausführen; MGs + volles Policy-Set im DoNotEnforce; Logging",
             "Governance steht; Compliance-Dashboard zeigt Ist-Zustand",
             "≈ €0"],
            ["3 – Netzwerk-Pilot",
             "Hub-VNets + Private DNS Zones; erste Spoke-VNets; ggf. minimale Firewall",
             "Workloads können in Landing Zones platziert werden",
             "€15–€1.300"],
            ["4 – Sicherheit & Enforcement",
             "Defender for Cloud konfigurieren; Guardrail-Initiativen scharfschalten; Remediation",
             "Policy-Compliance > 80 %; Defender aktiviert",
             "+Defender-Kosten"],
            ["5 – Workload-Onboarding",
             "Erste Workloads in Landing Zones; Spoke-Netzwerke; RBAC-Gruppen befüllen",
             "Erste Applikation läuft in ALZ-Scope; Policies greifen",
             "Workload-abhängig"],
            ["6 – Betrieb & Optimierung",
             "Cost Management, Tagging-Pflicht, Monitoring-Dashboards, Policy-Weiterentwicklung",
             "Laufender ALZ-Betrieb; regelmäßige Review-Zyklen",
             "Laufend"],
        ],
        col_widths=[1.5, 2.5, 2.0, 1.0]
    )
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 12. RISIKEN, OFFENE PUNKTE UND ENTSCHEIDUNGSBEDARF
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("12. Risiken, offene Punkte und Entscheidungsbedarf", level=1)

    doc.add_heading("Risikomatrix", level=2)
    add_table(doc,
        ["Risiko", "Auswirkung", "WSK", "Maßnahme"],
        [
            ["Subscription-Typ ungültig (Free Tier)",
             "Hoch – Bootstrap schlägt fehl", "Mittel",
             "EA/MCA/PAYG sicherstellen vor Bootstrap"],
            ["IP-Adresskonflikt mit On-Premises",
             "Mittel – Hub-VNet nicht routbar zu On-Prem", "Mittel",
             "IP-Adressplan vor Netzwerk-Deployment abstimmen"],
            ["DDoS-Kosten (Default: an)",
             "Hoch – €2.500/Monat unerwartet", "Hoch",
             "deployDdosProtectionPlan: false für Start"],
            ["Defender-Kosten (Policy aktiviert auto.)",
             "Mittel – Defender-Pläne ungewollt kostenpflichtig", "Mittel",
             "Defender-Tier-Konfiguration vor Apply reviewen"],
            ["Tenant Root Rechte fehlen",
             "Hoch – MG-Erstellung schlägt fehl", "Mittel",
             "Erhöhten Zugriff vor Bootstrap aktivieren"],
            ["Guardrails blockieren bestehende Ressourcen",
             "Mittel – Compliance-Fehler, ggf. Policy-Konflikte", "Gering",
             "DoNotEnforce-Modus für Start; Remediation geplant"],
        ],
        col_widths=[2.0, 1.7, 0.6, 2.7]
    )

    doc.add_heading("Offene Punkte", level=2)
    add_table(doc,
        ["ID", "Offener Punkt", "Aufwand", "Priorität"],
        [
            ["A1", "DNS Private Resolver: Flag deployDnsPrivateResolver vorhanden, aber kein AVM-DNS-Resolver-Modul verdrahtet", "Klein", "Mittel"],
            ["A2", "Decommissioned-MG: Header verspricht Deny-Policies, keine sind implementiert", "Klein", "Niedrig"],
            ["A3", "Sandbox-MG: Policy-Exemption für gelockerte Policies fehlt noch", "Klein", "Niedrig"],
            ["B",  "Smoke Run: Kunden-Tenant + GitHub PAT benötigt für Ende-zu-Ende-Test", "Mittel", "Hoch"],
            ["C",  "Security-Baseline: Defender-Pricings explizit, Security Contacts, Sentinel-Onboarding", "Mittel", "Mittel"],
        ],
        col_widths=[0.4, 3.2, 0.8, 0.8]
    )

    doc.add_heading("Entscheidungsbedarf Kunde", level=2)
    add_table(doc,
        ["Entscheidung", "Optionen", "Bechtle-Empfehlung"],
        [
            ["Subscription-Modell",       "1 Sub vs. 4 Platform-Subs",                    "1 Sub für Start; Trennung später"],
            ["Primärregion",              "Germany West Central bestätigen oder ändern",   "GWC (Datenresidenz Deutschland)"],
            ["Netzwerk-Typ für Pilot",    "none (€0) vs. hubNetworking (~€5.800)",         "none für Pilot"],
            ["DDoS Protection",           "An (~€2.500/Monat) vs. Aus",                    "Aus für Start; nach Internet-Workloads prüfen"],
            ["Enforcement-Zeitpunkt",     "Wann von DoNotEnforce auf Enabled",             "Nach 4–8 Wochen Audit-Phase"],
            ["On-Prem-Anbindung",         "VPN vs. ExpressRoute vs. keins",               "Nach Bandbreiten-/SLA-Anforderung"],
            ["Defender-Tier",             "Standard vs. Free je Plan",                     "Standard für Produktions-Workloads"],
            ["Entra ID / Identity",       "Entra-only vs. hybrid vs. AD DS",              "Anforderungsabhängig"],
        ],
        col_widths=[1.8, 2.3, 2.9]
    )
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 13. ERGÄNZENDE RANDTHEMEN UND ABHÄNGIGKEITEN
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("13. Ergänzende Randthemen und Abhängigkeiten", level=1)

    add_body(doc,
        "Die folgenden Themen sind eng mit der Azure Landing Zone verknüpft, "
        "aber außerhalb des Landing-Zone-Scopes. Sie werden als Abhängigkeiten "
        "und mögliche Folgebausteine eingeordnet."
    )

    doc.add_heading("Entra ID und Identity", level=2)
    add_bullet(doc, "Zu prüfen: Entra-ID-Tenant-Struktur, vorhandene Gruppen und Rollen, MFA-Status, Conditional Access Policies, PIM-Nutzung")
    add_bullet(doc, "Mögliche Folgestufe: Identity-Baseline (Entra Diagnostik → LAW, Conditional Access Grundkonfiguration, PIM für privilegierte Rollen, hybride Identitäten)")
    add_bullet(doc, "Abgrenzung: Entra-ID-Konfiguration und MFA-Rollout sind nicht Bestandteil der Landing Zone; die MG alz-platform-identity und die Identity-Subscription sind vorgesehen")

    doc.add_heading("On-Premises-Anbindung", level=2)
    add_bullet(doc, "Zu prüfen: Bandbreitenanforderungen, Latenz-SLA, bestehende Leitungsinfrastruktur, IP-Adressplan und Überlappungsrisiken mit Hub-VNets (10.0.0.0/22, 10.1.0.0/22)")
    add_bullet(doc, "Mögliche Folgestufe: VPN-Gateway-Konfiguration (VpnGw1AZ, active-active BGP, ASN 65515 – bereits deployt) oder ExpressRoute-Gateway mit Provider-Anbindung")
    add_bullet(doc, "Abgrenzung: Leitungsbestellung, Router-Konfiguration On-Premises und BGP-Peering mit Provider sind außerhalb des Landing-Zone-Scopes")

    doc.add_heading("Backup und Disaster Recovery", level=2)
    add_bullet(doc, "Zu prüfen: RPO/RTO-Anforderungen je Workload, vorhandene Backup-Lösungen, Geo-Redundanz-Bedarf")
    add_bullet(doc, "Mögliche Folgestufe: Recovery Services Vault je Landing Zone, Azure Backup Policies, Azure Site Recovery für kritische VMs")
    add_bullet(doc, "Abgrenzung: Die Policy Deploy-VM-Backup sichert VMs ohne Backup-Tag automatisch; Recovery-Vault-Konfiguration und DR-Tests sind eigene Projektschritte")

    doc.add_heading("Kostenmanagement", level=2)
    add_bullet(doc, "Zu prüfen: Budget-Limits je Subscription / Abteilung, Tagging-Strategie für Kostenzuordnung, Cost Alert-Schwellwerte")
    add_bullet(doc, "Mögliche Folgestufe: Azure Cost Management Dashboards, Budget-Alerts je MG/Sub, Tagging-Enforcement via Policy (Audit-Tags-Mandatory-Rg ist im Policy-Set enthalten)")
    add_bullet(doc, "Abgrenzung: Chargeback-Modelle, FinOps-Prozesse und Lizenz-Optimierung sind nicht Teil der Landing Zone")

    doc.add_heading("Compliance und Regulatorik", level=2)
    add_bullet(doc, "Zu prüfen: branchenspezifische Anforderungen (BSI, ISO 27001, DSGVO), Datenresidenz-Vorgaben, regulatorische Berichtspflichten")
    add_bullet(doc, "Mögliche Folgestufe: Compliance-spezifische Policy-Initiativen, Defender for Cloud Regulatory Compliance Dashboard, Microsoft Purview")
    add_bullet(doc, "Abgrenzung: rechtliche Compliance-Bewertung und Zertifizierungsberatung sind außerhalb des technischen Landing-Zone-Scopes")

    doc.add_heading("Workload-Migration", level=2)
    add_bullet(doc, "Zu prüfen: Inventar bestehender Workloads, Abhängigkeiten, Migrations-Reihenfolge, Downtime-Toleranz")
    add_bullet(doc, "Mögliche Folgestufe: Azure Migrate Assessment, Rehosting (Lift & Shift), Refactoring oder Rebuild je Workload")
    add_bullet(doc, "Abgrenzung: Workload-Migration ist ein eigenständiges Folgeprojekt; die Landing Zone ist das Ziel, nicht der Migrations-Prozess")
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 14. NÄCHSTE SCHRITTE UND EMPFEHLUNG
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("14. Nächste Schritte und Empfehlung", level=1)

    doc.add_heading("Sofort-Maßnahmen (ohne Azure-Kosten)", level=2)
    add_table(doc,
        ["#", "Maßnahme", "Verantwortlich", "Aufwand"],
        [
            ["1", "Kickoff-Termin vereinbaren; Entscheidungsprotokoll erstellen",                           "Beide",         "1 Tag"],
            ["2", "Azure Subscription bereitstellen (EA/MCA/PAYG, kein Free Tier)",                         "Kunde",         "1 Stunde"],
            ["3", "GitHub PAT (Classic) mit Scopes repo, workflow, admin:org erstellen",                     "Kunde",         "30 Min."],
            ["4", "Erhöhten Zugriff am Tenant Root aktivieren (Entra ID → Properties)",                     "Kunde / Admin", "10 Min."],
            ["5", "IP-Adressplan festlegen (Hub-Bereiche 10.0.0.0/22, 10.1.0.0/22 prüfen / anpassen)",    "Beide",         "½–1 Tag"],
            ["6", "Entscheidungsprotokoll ausfüllen: Region, network_type, DDoS, Enforcement-Zeitpunkt",   "Beide",         "½ Tag"],
        ],
        col_widths=[0.3, 3.5, 1.2, 0.8]
    )

    doc.add_heading("Empfohlener Pilot-Pfad", level=2)
    add_bullet(doc, "Schritt 1: Subscription-ID und GitHub PAT in config/inputs-github.yaml eintragen")
    add_bullet(doc, "Schritt 2: Deploy-Accelerator lokal ausführen – Bootstrap erstellt Repo, OIDC-Identität und Pipelines")
    add_bullet(doc, "Schritt 3: Pipeline mit What-If ausführen und Review: welche Ressourcen werden erstellt?")
    add_bullet(doc, "Schritt 4: Apply genehmigen – mit network_type: none (€0) für risikofreien Start")
    add_bullet(doc, "Schritt 5: Policy-Compliance-Dashboard reviewen: welche Guardrails greifen, welche nicht?")
    add_bullet(doc, "Schritt 6: Nach Analyse schrittweise Enforcement und Netzwerk-Dienste aktivieren")
    add_body(doc,
        "Das vollständige Runbook mit allen Details steht in docs/ACCELERATOR-BOOTSTRAP.md."
    )

    doc.add_heading("Bechtle-Begleitung", level=2)
    add_body(doc,
        "Bechtle begleitet den Kunden über alle Phasen: "
        "von der Kickoff-Workshop-Vorbereitung über den Bootstrap bis zum ersten "
        "Workload-Onboarding. Das Angebot umfasst:"
    )
    add_bullet(doc, "Kickoff-Workshop (Discovery, Architektur-Entscheidungen, Entscheidungsprotokoll)")
    add_bullet(doc, "Bootstrap-Begleitung (technische Durchführung, Troubleshooting, Abnahme)")
    add_bullet(doc, "Pilot-Pilotierung (erste Landing Zone, Spoke-Netzwerk, Workload-Onboarding)")
    add_bullet(doc, "Optionaler laufender Betrieb und Weiterentwicklung nach Abschluss der Grundimplementierung")
    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # ANHANG
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_heading("Anhang", level=1)

    # A – 18 Deployment-Stufen
    doc.add_heading("A – 18 Deployment-Stufen (vollständig)", level=2)
    add_table(doc,
        ["#", "Name (.config/ALZ-Powershell.config.json)", "Scope", "Template-Pfad"],
        [
            ["1",  "mgmt_groups_int_root",                "managementGroup", "core/governance/mgmt-groups/int-root/main.bicep"],
            ["2",  "mgmt_groups_landing_zones",           "managementGroup", "core/governance/mgmt-groups/landingzones/main.bicep"],
            ["3",  "mgmt_groups_landing_zones_corp",      "managementGroup", "core/governance/mgmt-groups/landingzones/landingzones-corp/main.bicep"],
            ["4",  "mgmt_groups_landing_zones_online",    "managementGroup", "core/governance/mgmt-groups/landingzones/landingzones-online/main.bicep"],
            ["5",  "mgmt_groups_landing_zones_local",     "managementGroup", "core/governance/mgmt-groups/landingzones/landingzones-local/main.bicep"],
            ["6",  "mgmt_groups_platform",                "managementGroup", "core/governance/mgmt-groups/platform/main.bicep"],
            ["7",  "mgmt_groups_platform_connectivity",   "managementGroup", "core/governance/mgmt-groups/platform/platform-connectivity/main.bicep"],
            ["8",  "mgmt_groups_platform_identity",       "managementGroup", "core/governance/mgmt-groups/platform/platform-identity/main.bicep"],
            ["9",  "mgmt_groups_platform_management",     "managementGroup", "core/governance/mgmt-groups/platform/platform-management/main.bicep"],
            ["10", "mgmt_groups_platform_security",       "managementGroup", "core/governance/mgmt-groups/platform/platform-security/main.bicep"],
            ["11", "mgmt_groups_sandbox",                 "managementGroup", "core/governance/mgmt-groups/sandbox/main.bicep"],
            ["12", "mgmt_groups_decommissioned",          "managementGroup", "core/governance/mgmt-groups/decommissioned/main.bicep"],
            ["13", "rbac_platform",                       "managementGroup", "core/governance/mgmt-groups/platform/main-rbac.bicep"],
            ["14", "rbac_platform_connectivity",          "managementGroup", "core/governance/mgmt-groups/platform/platform-connectivity/main-rbac.bicep"],
            ["15", "rbac_landingzones",                   "managementGroup", "core/governance/mgmt-groups/landingzones/main-rbac.bicep"],
            ["16", "logging",                             "subscription",    "core/logging/main.bicep"],
            ["17", "hub_networking",                      "subscription",    "networking/hubnetworking/main.bicep (aktiv bei hubNetworking)"],
            ["18", "virtual_wan",                         "subscription",    "networking/virtualwan/main.bicep (aktiv bei vwanConnectivity)"],
        ],
        col_widths=[0.3, 2.5, 1.1, 3.1]
    )

    # B – AVM-Module
    doc.add_heading("B – Azure Verified Modules (22 Module)", level=2)
    add_table(doc,
        ["Modul", "Version", "Verw.", "Zweck"],
        [
            ["avm/ptn/alz/empty",                               "0.3.6",  "12", "MG-Hierarchie + volles ALZ-Policy-Set (Governance)"],
            ["avm/ptn/authorization/role-assignment",           "0.2.4",  "10", "RBAC-Zuweisung auf MG-Ebene"],
            ["avm/res/resources/resource-group",                "0.4.3",   "7", "Resource Groups"],
            ["avm/res/network/virtual-network",                 "0.7.2",   "3", "Hub-VNets"],
            ["avm/res/network/virtual-network-gateway",         "0.10.1",  "2", "VPN- und ExpressRoute-Gateways"],
            ["avm/res/network/route-table",                     "0.5.0",   "2", "Route Tables für Spoke-Netzwerke"],
            ["avm/res/network/firewall-policy",                 "0.3.4",   "2", "Azure Firewall Policy"],
            ["avm/res/network/dns-resolver",                    "0.5.6",   "2", "DNS Private Resolver"],
            ["avm/res/network/ddos-protection-plan",            "0.3.2",   "2", "DDoS Network Protection Plan"],
            ["avm/res/network/bastion-host",                    "0.8.2",   "2", "Azure Bastion"],
            ["avm/res/network/azure-firewall",                  "0.9.2",   "2", "Azure Firewall"],
            ["avm/ptn/network/private-link-private-dns-zones",  "0.7.2",   "2", "Private DNS Zones für Private Link"],
            ["avm/res/operational-insights/workspace",          "0.14.2",  "1", "Log Analytics Workspace"],
            ["avm/res/network/vpn-server-configuration",        "0.1.2",   "1", "VPN Server-Konfiguration"],
            ["avm/res/network/vpn-gateway",                     "0.2.2",   "1", "VPN Gateway (Virtual WAN)"],
            ["avm/res/network/virtual-wan",                     "0.4.3",   "1", "Virtual WAN"],
            ["avm/res/network/virtual-hub",                     "0.4.3",   "1", "Virtual Hub (vWAN)"],
            ["avm/res/network/public-ip-address",               "0.12.0",  "1", "Public IP (Bastion, Gateway)"],
            ["avm/res/network/network-security-group",          "0.5.2",   "1", "Network Security Group"],
            ["avm/res/network/express-route-gateway",           "0.8.0",   "1", "ExpressRoute Gateway"],
            ["avm/res/automation/automation-account",           "0.17.1",  "1", "Automation Account (optional)"],
            ["avm/ptn/alz/ama",                                 "0.2.0",   "1", "Log Analytics + DCRs + Managed Identity (Logging)"],
        ],
        col_widths=[3.2, 0.7, 0.5, 2.6]
    )

    # C – Glossar
    doc.add_heading("C – Glossar", level=2)
    add_table(doc,
        ["Begriff", "Bedeutung"],
        [
            ["ALZ",           "Azure Landing Zone – standardisiertes, governance-konformes Cloud-Fundament für Workloads in Azure"],
            ["AVM",           "Azure Verified Modules – von Microsoft geprüfte Bicep-Module aus der öffentlichen Microsoft Container Registry (MCR)"],
            ["MG",            "Management Group – hierarchische Verwaltungseinheit für Azure-Subscriptions; Policies und RBAC werden an alle untergeordneten Subscriptions vererbt"],
            ["Hub-and-Spoke", "Netzwerktopologie: zentraler Hub (Shared Services: Firewall, DNS, Gateways) + angebundene Spoke-Netzwerke (Workloads)"],
            ["OIDC",          "OpenID Connect – passwortlose Pipeline-Anmeldung an Azure über Federated Identity (keine gespeicherten Secrets)"],
            ["DCR",           "Data Collection Rule – Azure Monitor Konfigurationsobjekt für Telemetrie-Erfassung"],
            ["DoNotEnforce",  "Policy-Modus: Policies werden ausgewertet (Audit), aber keine Ressource wird blockiert – ideal für Onboarding-Phasen"],
            ["DINE",          "DeployIfNotExists – Policy-Effekt: Azure deployed automatisch fehlende Konfiguration (z. B. Diagnose-Einstellungen)"],
            ["Guardrail",     "Automatische Sicherheitsleitplanke per Policy; blockiert unerwünschte Konfigurationen oder korrigiert sie automatisch"],
            ["Bootstrap",     "Phase 0 des ALZ-Deployments: Terraform-basierte Einrichtung von GitHub-Repository, OIDC-Identität und Deployment-Pipelines"],
            ["What-If",       "Vorschau-Modus von ARM: zeigt alle geplanten Änderungen ohne Ausführung"],
            ["Vending",       "Automatisierte Bereitstellung und Platzierung von Subscriptions in die Landing-Zone-Hierarchie"],
            ["Spoke",         "Workload-Netzwerk, das über VNet-Peering an den Hub angebunden ist und Datenverkehr über die Hub-Firewall leitet"],
        ],
        col_widths=[1.8, 5.2]
    )

    # D – Referenzen
    doc.add_heading("D – Referenzen", level=2)
    add_table(doc,
        ["Quelle", "Beschreibung"],
        [
            ["aka.ms/alz",                                           "Microsoft Azure Landing Zones – offizielle Dokumentation und Design-Prinzipien"],
            ["aka.ms/alz/acc",                                       "ALZ Bicep Accelerator – offizielles GitHub Repository und Dokumentation"],
            ["azure.github.io/Azure-Verified-Modules/",             "Azure Verified Modules Registry – alle AVM-Module mit Versionsverlauf"],
            ["learn.microsoft.com/azure/azure-resource-manager/bicep/", "Microsoft Bicep Dokumentation"],
            ["learn.microsoft.com/azure/governance/policy/",        "Azure Policy Dokumentation"],
            ["learn.microsoft.com/azure/defender-for-cloud/",       "Microsoft Defender for Cloud Dokumentation"],
            ["docs/ACCELERATOR-BOOTSTRAP.md",                       "Lokales Runbook: Bootstrap- und Deployment-Schritte (Schritt-für-Schritt)"],
            ["docs/TECHNICAL-REFERENCE.md",                         "Lokale Technische Referenz: alle 22 AVM-Module, 123 Assignments, 149 Policy-Definitionen"],
        ],
        col_widths=[2.8, 4.2]
    )

    # Speichern
    doc.save(OUTPUT)
    print(f"Konzept gespeichert: {OUTPUT}")
    return OUTPUT


if __name__ == "__main__":
    build()
